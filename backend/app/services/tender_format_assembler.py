"""格式填充装配器：把推导目录 + 抽取模板 + 项目/企业事实 装配成一份可导出的投标文件 docx。

定位（B 路径，独立装配器）：
  · 不改动现有 context_pack / business_draft 生成管线，快速产出与样张一致的真实 docx。
  · 复用 tender_directory（目录推导）、tender_outline（节点→章节映射）、
    tender_format_templates（模板抽取+填充）三个已验证模块。

装配规则（按推导出的章节顺序）：
  · 命中格式模板（声明/委托书/身份证明）→ 渲染填充后的正文，留盖章/签字位。
  · 附件/证照章（资格证明资料等）→ 标题 + 附件清单占位（待企业材料库插入扫描件）。
  · 结构化表格章（偏离表/报价表）→ 标题 + 表头占位。
  · 其余 → 人工撰写占位。
"""

from __future__ import annotations

import hashlib
from io import BytesIO
import re
from typing import Any

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.services.tender_compliance_coverage import (
    CoverLayer,
    CoverStatus,
    CoverageItem,
    CoverageReport,
    _classify_attachment,
    compute_coverage,
)
from app.services.tender_format_templates import (
    FormTemplate,
    _title_to_key,
    extract_format_templates,
    fill_template,
)
from app.services.tender_pricing import PricingValidationReport, build_pricing_report

# ----------------------------------------------------------------------------
# 事实适配器：现有 _project_facts + 招标信号 → 模板引擎 field_keys
# ----------------------------------------------------------------------------

# 引擎 field_key ← _project_facts 键
_FACT_MAP: dict[str, str] = {
    "purchaser": "tenderer_name",
    "project_name": "project_name",
    "agent_code": "tender_project_no",
    "supplier_name": "bidder_name",
    "legal_representative": "legal_representative_name",
    "agent_person": "authorized_agent_name",
    "social_credit_code": "unified_social_credit_code",
    "address": "bidder_address",
    "date": "bid_date",
    "construction_period_days": "construction_period_days",
    "quality_standard": "quality_standard",
    "warranty_period": "warranty_period",
    "project_scope": "project_scope",
    "bid_price_amount": "bid_price_amount",
}

_CN_DIGITS = {"零": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4, "五": 5,
              "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}


def _parse_copies_duplicate(copies: Any) -> str | None:
    """从『一正两副』『一正贰副』『1正2副』提取副本份数（原样保留中文数字）。"""
    if not copies:
        return None
    s = str(copies)
    m = re.search(r"([0-9一二两三四五六七八九十]+)\s*副", s)
    return m.group(1) if m else None


def build_form_facts(
    project_facts: dict[str, Any],
    signals: dict[str, Any] | None = None,
    overrides: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """构造模板引擎所需的填充事实字典；未知字段缺省，引擎将留占位。"""
    signals = signals or {}
    facts: dict[str, Any] = {}
    for fkey, src in _FACT_MAP.items():
        val = project_facts.get(src)
        if val not in (None, ""):
            facts[fkey] = val

    # 签字代表：优先授权代理人，回退法定代表人
    if "agent_person" in facts:
        facts.setdefault("signatory", facts["agent_person"])
    elif "legal_representative" in facts:
        facts.setdefault("signatory", facts["legal_representative"])

    # 招标信号驱动
    dup = _parse_copies_duplicate(signals.get("copies"))
    if dup:
        facts["copies_duplicate"] = dup
    if signals.get("bid_validity_days"):
        facts["bid_validity_days"] = str(signals["bid_validity_days"])

    if overrides:
        for k, v in overrides.items():
            if v not in (None, ""):
                facts[k] = v
    return facts


# ----------------------------------------------------------------------------
# docx 样式（与 business_draft 导出一致：宋体正文 / 黑体居中标题）
# ----------------------------------------------------------------------------

_MARGINS_CM = {"top": 2.54, "bottom": 2.54, "left": 3.18, "right": 3.18}
_SEAL_HINTS = ("盖单位公章", "盖单位章", "盖公章", "签字或印章", "签字或盖章")
_QUALIFICATION_MATERIAL_TYPES = {
    "license",
    "qualification",
    "personnel",
    "performance",
    "commitment",
    "test_report",
    "other",
}


def _set_font(run: Any, name: str = "宋体", size: Pt | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def _configure(document: Any) -> None:
    for section in document.sections:
        section.top_margin = Cm(_MARGINS_CM["top"])
        section.bottom_margin = Cm(_MARGINS_CM["bottom"])
        section.left_margin = Cm(_MARGINS_CM["left"])
        section.right_margin = Cm(_MARGINS_CM["right"])
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    h1 = document.styles["Heading 1"]
    h1.font.name = "黑体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


_DEVIATION_HEADERS = ["序号", "磋商文件条目号", "采购规格/商务条款", "响应文件的规格/商务条款", "响应与偏离", "说明"]
_PRICE_HEADERS = ["序号", "项目名称", "单位", "数量", "综合单价(元)", "合价(元)", "备注"]
_PRICE_REVIEW_HEADERS = [*_PRICE_HEADERS, "校验状态"]


def _render_form(document: Any, filled_lines: list[str]) -> None:
    for line in filled_lines:
        text = line.strip()
        if not text:
            continue
        para = document.add_paragraph()
        run = para.add_run(text)
        _set_font(run)
        # 盖章/签字行右对齐，便于人工盖章
        if any(h in text for h in _SEAL_HINTS):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _render_attachment_chapter(
    document: Any,
    attachments: list[str],
    status_by_title: dict[str, CoverageItem] | None = None,
    *,
    chapter_title: str = "",
    material_assets: list[dict[str, Any]] | None = None,
    include_review_status: bool = True,
    page_ref_registry: _PageRefRegistry | None = None,
) -> list[dict[str, Any]]:
    note = document.add_paragraph()
    note_text = (
        "［附：以下证照/材料扫描件由企业材料库插入，正式投标前需替换为加盖公章的复印件］"
        if include_review_status
        else "以下证明材料按招标文件要求随本章提交。"
    )
    r = note.add_run(note_text)
    _set_font(r, size=Pt(10.5))
    status_by_title = status_by_title or {}
    material_assets = material_assets or []
    embedded: list[dict[str, Any]] = []
    used_material_ids: set[str] = set()
    render_attachments = attachments or _fallback_material_attachment_titles(
        chapter_title,
        material_assets,
    )
    for att in render_attachments:
        p = document.add_paragraph(style="List Bullet")
        if page_ref_registry is not None:
            page_ref_registry.add_to_paragraph(p, kind="attachment", label=att)
        _set_font(p.add_run(att))
        item = status_by_title.get(att.strip())
        if include_review_status and item is not None:
            tag = f"  ［{item.status_label}{'·废标风险' if item.disqualifying and item.status in (CoverStatus.NEEDS_MATERIAL, CoverStatus.MISSING) else ''}］"
            run = p.add_run(tag)
            _set_font(run, size=Pt(9))
        asset = _best_material_asset_for_attachment(
            att,
            material_assets,
            used_material_ids,
            require_embeddable=True,
        )
        if asset is not None:
            result = _render_material_asset(
                document,
                asset,
                page_ref_registry=page_ref_registry,
            )
            if result is not None:
                embedded.append({"attachment": att, **result})
                if result.get("material_id"):
                    used_material_ids.add(str(result["material_id"]))
    if not render_attachments:
        _set_font(document.add_paragraph("（按招标文件第三章资格要求逐项附证明材料）").runs[0])
    return embedded


def _fallback_material_attachment_titles(
    chapter_title: str,
    material_assets: list[dict[str, Any]],
) -> list[str]:
    """When a qualification chapter has no explicit attachment list, embed usable proofs.

    Many tenders only name a broad "资格证明资料" chapter and describe required
    files elsewhere. The fallback stays conservative: it only renders materials
    that are already confirmed, allowed, and have actual embeddable images.
    """
    if not any(hint in chapter_title for hint in _ATTACH_SECTION_HINTS):
        return []
    titles: list[str] = []
    for asset in material_assets:
        if not isinstance(asset, dict):
            continue
        if not _material_asset_confirmed(asset) or not _material_asset_allowed(asset):
            continue
        if not _material_has_embeddable_images(asset):
            continue
        material_type = str(asset.get("material_type") or "")
        if material_type not in _QUALIFICATION_MATERIAL_TYPES:
            continue
        title = str(asset.get("material_name") or asset.get("name") or "").strip()
        if title:
            titles.append(title)
    return titles


def _material_asset_text(asset: dict[str, Any]) -> str:
    fields = [
        asset.get("material_name"),
        asset.get("name"),
        asset.get("material_type"),
        asset.get("category"),
        asset.get("certificate_no"),
        asset.get("evidence_text"),
        asset.get("file_name"),
    ]
    return " ".join(str(value).strip() for value in fields if value not in (None, ""))


def _attachment_match_keywords(attachment_title: str) -> tuple[str, ...]:
    _, _, keywords = _classify_attachment(attachment_title)
    if keywords:
        return keywords
    return tuple(part for part in re.split(r"[，、\s/（）()]+", attachment_title) if len(part) >= 2)


def _material_matches_attachment(
    attachment_title: str,
    asset: dict[str, Any],
) -> bool:
    keywords = _attachment_match_keywords(attachment_title)
    if not keywords:
        return False
    text = _material_asset_text(asset)
    return any(keyword in text for keyword in keywords)


def _material_asset_confirmed(asset: dict[str, Any]) -> bool:
    status = str(
        asset.get("verification_status")
        or asset.get("status")
        or asset.get("confirmation_status")
        or ""
    ).strip()
    return status in {"confirmed", "verified", "valid", "active"}


def _material_asset_allowed(asset: dict[str, Any]) -> bool:
    data_level = str(asset.get("data_level") or "internal").strip()
    return data_level in {"public", "internal"}


def _material_embeddable_images(asset: dict[str, Any]) -> list[dict[str, Any]]:
    return [
        img
        for img in asset.get("embedded_images") or []
        if isinstance(img, dict) and img.get("data")
    ]


def _material_has_embeddable_images(asset: dict[str, Any]) -> bool:
    return bool(_material_embeddable_images(asset))


def _best_material_asset_for_attachment(
    attachment_title: str,
    material_assets: list[dict[str, Any]],
    used_material_ids: set[str],
    *,
    require_embeddable: bool = False,
) -> dict[str, Any] | None:
    for asset in material_assets:
        if not isinstance(asset, dict):
            continue
        material_id = str(asset.get("material_id") or asset.get("id") or "")
        if material_id and material_id in used_material_ids:
            continue
        if not _material_asset_confirmed(asset) or not _material_asset_allowed(asset):
            continue
        if require_embeddable and not _material_has_embeddable_images(asset):
            continue
        if _material_matches_attachment(attachment_title, asset):
            return asset
    return None


def _render_material_asset(
    document: Any,
    asset: dict[str, Any],
    *,
    page_ref_registry: _PageRefRegistry | None = None,
) -> dict[str, Any] | None:
    images = _material_embeddable_images(asset)
    if not images:
        return None

    material_name = str(asset.get("material_name") or asset.get("name") or "企业证明材料")
    title = document.add_paragraph()
    if page_ref_registry is not None:
        page_ref_registry.add_to_paragraph(title, kind="material", label=material_name)
    _set_font(title.add_run(f"材料：{material_name}"), bold=True, size=Pt(10.5))
    inserted = 0
    for index, image in enumerate(images, start=1):
        try:
            document.add_picture(BytesIO(image["data"]), width=Cm(14.0))
        except Exception:
            continue
        inserted += 1
        caption_text = image.get("caption") or f"{material_name} 第{index}页"
        caption = document.add_paragraph()
        _set_font(caption.add_run(str(caption_text)), size=Pt(9))
    if inserted == 0:
        return None
    return {
        "material_id": str(asset.get("material_id") or asset.get("id") or ""),
        "material_name": material_name,
        "image_count": inserted,
    }


_CHECKLIST_HEADERS = ["层级", "条目", "状态", "说明"]


def _render_compliance_checklist(document: Any, report: CoverageReport) -> None:
    """渲染分层合规自检清单：废标风险红牌 + 按 L1/L2/L3 分组的状态表。"""
    document.add_heading("合规自检清单", level=1)

    disq = report.disqualifying_gaps
    banner = document.add_paragraph()
    if disq:
        names = "；".join(i.title for i in disq[:6])
        run = banner.add_run(f"⚠ 废标风险：{len(disq)} 项核心资格尚未满足 —— {names}")
    else:
        run = banner.add_run("✓ 暂未发现废标风险（核心资格项均有事实或材料覆盖）")
    _set_font(run, bold=True, size=Pt(11))

    gap_total = len(report.gaps)
    sub = document.add_paragraph()
    _set_font(sub.add_run(f"待处理缺口合计 {gap_total} 项（含需补字段 / 需上传材料 / 待逐条应答）。"), size=Pt(10.5))

    table = document.add_table(rows=1, cols=len(_CHECKLIST_HEADERS))
    table.style = "Table Grid"
    for i, h in enumerate(_CHECKLIST_HEADERS):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            _set_font(run, bold=True, size=Pt(10.5))

    for layer in (CoverLayer.L1_FORMAT, CoverLayer.L2_SCORING, CoverLayer.L3_RESPONSE):
        layer_items = report.for_layer(layer)
        if not layer_items:
            continue
        for it in layer_items:
            cells = table.add_row().cells
            mark = "★" if it.disqualifying else ""
            values = [layer.value, f"{mark}{it.title}", it.status_label, it.detail]
            for cell, val in zip(cells, values):
                cell.text = val
                for run in cell.paragraphs[0].runs:
                    _set_font(run, size=Pt(9.5), bold=bool(mark))

    legend = document.add_paragraph()
    _set_font(
        legend.add_run("说明：★为核心资格/废标项；L2 评分点、L3 偏离响应当前列出待应答项，逐条比对将在后续接入评分办法与技术参数解析。"),
        size=Pt(9),
    )


def _render_table_chapter(document: Any, headers: list[str], note: str) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            _set_font(run, bold=True, size=Pt(10.5))
    # 预留一空行供填写
    for cell in table.add_row().cells:
        cell.text = ""
    hint = document.add_paragraph()
    _set_font(hint.add_run(note), size=Pt(10.5))


# ----------------------------------------------------------------------------
# 装配
# ----------------------------------------------------------------------------

# section_type / mapped key → 渲染策略（按标题关键词路由，先于模板表单）
_TABLE_SECTION_HINTS = ("偏离", "deviation")
_PRICE_SECTION_HINTS = ("报价", "价格", "pricing", "price")
_INDEX_SECTION_HINTS = ("评分索引", "索引")
_ATTACH_SECTION_HINTS = ("资格", "证明", "qualification", "认为需提供", "资料")
_INDEX_HEADERS = ["序号", "评分项", "对应章节/附件", "页码"]
_INDEX_REVIEW_HEADERS = [*_INDEX_HEADERS, "覆盖状态"]
# 真正走表单填充的语义模板键（声明/委托书/身份证明）
_FORM_TEMPLATE_KEYS = {"bid_commitment", "authorization_letter", "legal_representative_identity"}
_L4_SECTION_HINTS = (
    "技术方案",
    "施工组织",
    "项目理解",
    "招标要求响应",
    "EPC",
    "设计施工协同",
    "限额设计",
    "预算控制",
    "质量控制",
    "材料采购",
    "管材",
    "阀门",
    "沟槽",
    "管道连接",
    "防腐",
    "回填",
    "路面恢复",
    "压力试验",
    "吹扫",
    "置换",
    "验收",
    "安全文明",
    "交通组织",
    "进度计划",
    "资源配置",
    "缺陷责任",
    "服务承诺",
    "竣工资料",
)
_GAS_SECTION_HINTS = ("燃气", "管道", "GB1", "中压", "沟槽", "压力试验", "吹扫", "置换")
_L4_FORBIDDEN_SPEC_HINTS = ("DN", "MPa", "km", "PE", "钢管", "管径", "压力等级", "工程量")
_SCORING_TARGET_KEYWORDS = (
    "技术",
    "方案",
    "施工",
    "组织",
    "质量",
    "安全",
    "进度",
    "服务",
    "承诺",
    "业绩",
    "报价",
    "价格",
    "偏离",
    "响应",
    "商务",
    "资格",
)


def _canonical_ref_text(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "").strip())


class _PageRefRegistry:
    """Track Word bookmarks used by scoring-index PAGEREF fields.

    Page numbers are renderer-dependent, so the assembler writes native Word
    fields instead of guessing static numbers. Word/LibreOffice updates these
    fields from the real pagination when the document is opened or refreshed.
    """

    def __init__(self) -> None:
        self._name_by_key: dict[tuple[str, str], str] = {}
        self._inserted_names: set[str] = set()
        self._next_id = 1

    def reserve(self, kind: str, label: Any) -> str | None:
        canonical = _canonical_ref_text(label)
        if not canonical:
            return None
        key = (kind, canonical)
        if key not in self._name_by_key:
            digest = hashlib.sha1(f"{kind}:{canonical}".encode("utf-8")).hexdigest()[:16]
            self._name_by_key[key] = f"bm_{digest}"
        return self._name_by_key[key]

    def add_to_paragraph(self, paragraph: Any, *, kind: str, label: Any) -> str | None:
        name = self.reserve(kind, label)
        if not name or name in self._inserted_names:
            return name
        bookmark_id = str(self._next_id)
        self._next_id += 1

        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bookmark_id)
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark_id)

        paragraph._p.insert(0, start)
        paragraph._p.append(end)
        self._inserted_names.add(name)
        return name


def _enable_field_updates_on_open(document: Any) -> None:
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _add_pageref_field(paragraph: Any, bookmark_name: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)
    _set_font(begin_run, size=Pt(9.5))

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGEREF {bookmark_name} \\h "
    instr_run._r.append(instr)
    _set_font(instr_run, size=Pt(9.5))

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    _set_font(separate_run, size=Pt(9.5))

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    _set_font(end_run, size=Pt(9.5))


def _active_scoring_bindings(item: dict[str, Any]) -> list[dict[str, Any]]:
    bindings = item.get("evidence_bindings")
    if not isinstance(bindings, list):
        return []
    return [
        binding
        for binding in bindings
        if isinstance(binding, dict) and str(binding.get("status") or "active") == "active"
    ]


def _scoring_index_rows(compliance_items: list[Any] | None) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for raw in compliance_items or []:
        if not isinstance(raw, dict):
            continue
        if raw.get("item_type") != "scoring" or raw.get("status") in {"rejected", "superseded"}:
            continue
        requirement = str(raw.get("requirement_text") or raw.get("title") or "").strip()
        if not requirement:
            continue
        bindings = _active_scoring_bindings(raw)
        material_names = [
            str(binding.get("material_name") or binding.get("enterprise_material_name") or "").strip()
            for binding in bindings
        ]
        material_names = [name for name in material_names if name]
        rows.append(
            {
                "requirement_text": requirement,
                "target": "、".join(material_names),
                "page": "",
                "status": "已绑定证据" if bindings else "待应答",
            }
        )
    return rows


def _technical_response_rows(compliance_items: list[Any] | None) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for raw in compliance_items or []:
        if not isinstance(raw, dict):
            continue
        if raw.get("item_type") != "technical_response":
            continue
        if raw.get("status") in {"rejected", "superseded"}:
            continue
        requirement = str(raw.get("requirement_text") or raw.get("title") or "").strip()
        if not requirement:
            continue
        bindings = _active_scoring_bindings(raw)
        material_names = [
            str(binding.get("material_name") or binding.get("enterprise_material_name") or "").strip()
            for binding in bindings
        ]
        rows.append(
            {
                "requirement_text": requirement,
                "response_text": "",
                "deviation": "",
                "note": "、".join(name for name in material_names if name),
            }
        )
    return rows


def _attachment_render_plan(
    chapters: list[dict[str, Any]],
    material_assets: list[dict[str, Any]] | None,
) -> tuple[set[str], set[str]]:
    material_assets = material_assets or []
    attachment_titles: set[str] = set()
    material_names: set[str] = set()
    used_material_ids: set[str] = set()
    for chapter in chapters:
        title = str(chapter.get("title") or "")
        if not any(h in title for h in _ATTACH_SECTION_HINTS):
            continue
        attachments = chapter.get("attachments") or []
        render_attachments = attachments or _fallback_material_attachment_titles(title, material_assets)
        for attachment in render_attachments:
            attachment_text = str(attachment or "").strip()
            if attachment_text:
                attachment_titles.add(_canonical_ref_text(attachment_text))
            asset = _best_material_asset_for_attachment(
                attachment_text,
                material_assets,
                used_material_ids,
                require_embeddable=True,
            )
            if asset is None:
                continue
            material_name = str(asset.get("material_name") or asset.get("name") or "").strip()
            if material_name:
                material_names.add(_canonical_ref_text(material_name))
            material_id = str(asset.get("material_id") or asset.get("id") or "")
            if material_id:
                used_material_ids.add(material_id)
    return attachment_titles, material_names


def material_render_candidate_ids(
    chapters: list[dict[str, Any]],
    material_assets: list[Any] | None,
    compliance_items: list[Any] | None = None,
) -> set[str]:
    """Return material ids that may need image bytes during attachment rendering.

    The export service uses this before fetching MinIO/PDF bytes. It intentionally
    works from metadata only; the assembler still decides final rendering after
    images are loaded.
    """
    assets = [asset for asset in material_assets or [] if isinstance(asset, dict)]
    candidate_ids: set[str] = set()

    def material_id_for(asset: dict[str, Any]) -> str:
        return str(asset.get("material_id") or asset.get("id") or "").strip()

    def add_asset(asset: dict[str, Any]) -> None:
        material_id = material_id_for(asset)
        if material_id:
            candidate_ids.add(material_id)

    for chapter in chapters:
        title = str(chapter.get("title") or "")
        if not any(h in title for h in _ATTACH_SECTION_HINTS):
            continue
        attachments = [str(attachment or "").strip() for attachment in chapter.get("attachments") or []]
        attachments = [attachment for attachment in attachments if attachment]
        if attachments:
            for attachment in attachments:
                for asset in assets:
                    if not _material_asset_confirmed(asset) or not _material_asset_allowed(asset):
                        continue
                    if _material_matches_attachment(attachment, asset):
                        add_asset(asset)
            continue

        # Without explicit attachment titles, qualification chapters fall back
        # to rendering usable qualification proofs. Fetch only those types.
        for asset in assets:
            if not _material_asset_confirmed(asset) or not _material_asset_allowed(asset):
                continue
            if str(asset.get("material_type") or "") in _QUALIFICATION_MATERIAL_TYPES:
                add_asset(asset)

    assets_by_name: dict[str, list[dict[str, Any]]] = {}
    for asset in assets:
        for field in (asset.get("material_name"), asset.get("name")):
            canonical = _canonical_ref_text(field)
            if canonical:
                assets_by_name.setdefault(canonical, []).append(asset)

    for raw in compliance_items or []:
        if not isinstance(raw, dict):
            continue
        for binding in _active_scoring_bindings(raw):
            material_id = str(
                binding.get("enterprise_material_id")
                or binding.get("material_id")
                or binding.get("id")
                or ""
            ).strip()
            if material_id:
                candidate_ids.add(material_id)
            material_name = _canonical_ref_text(
                binding.get("material_name") or binding.get("enterprise_material_name")
            )
            for asset in assets_by_name.get(material_name, []):
                add_asset(asset)

    return candidate_ids


def _split_target_names(value: Any) -> list[str]:
    return [
        part.strip()
        for part in re.split(r"[、,，;；]+", str(value or ""))
        if part.strip()
    ]


def _best_scoring_target_chapter(
    requirement_text: str,
    chapters: list[dict[str, Any]],
) -> str | None:
    requirement = str(requirement_text or "")
    if not requirement:
        return None
    best_title = None
    best_score = 0
    for chapter in chapters:
        title = str(chapter.get("title") or "")
        if not title or any(h in title for h in _INDEX_SECTION_HINTS):
            continue
        score = 0
        if title in requirement:
            score += 6
        if requirement in title:
            score += 4
        score += sum(1 for keyword in _SCORING_TARGET_KEYWORDS if keyword in requirement and keyword in title)
        if score > best_score:
            best_score = score
            best_title = title
    return best_title if best_score > 0 else None


def _attach_scoring_page_refs(
    rows: list[dict[str, Any]],
    chapters: list[dict[str, Any]],
    material_assets: list[dict[str, Any]] | None,
    page_ref_registry: _PageRefRegistry,
) -> None:
    attachment_titles, material_names = _attachment_render_plan(chapters, material_assets)
    chapter_titles = {
        _canonical_ref_text(chapter.get("title")): str(chapter.get("title") or "").strip()
        for chapter in chapters
        if str(chapter.get("title") or "").strip()
    }
    for title in chapter_titles.values():
        page_ref_registry.reserve("chapter", title)

    for row in rows:
        page_refs: list[dict[str, str]] = []
        for target in _split_target_names(row.get("target")):
            canonical = _canonical_ref_text(target)
            if canonical in material_names:
                bookmark = page_ref_registry.reserve("material", target)
                if bookmark:
                    page_refs.append({"kind": "material", "target": target, "bookmark": bookmark})
            elif canonical in attachment_titles:
                bookmark = page_ref_registry.reserve("attachment", target)
                if bookmark:
                    page_refs.append({"kind": "attachment", "target": target, "bookmark": bookmark})
            elif canonical in chapter_titles:
                chapter_title = chapter_titles[canonical]
                bookmark = page_ref_registry.reserve("chapter", chapter_title)
                if bookmark:
                    page_refs.append({"kind": "chapter", "target": chapter_title, "bookmark": bookmark})
        if not page_refs and not row.get("target"):
            chapter_title = _best_scoring_target_chapter(str(row.get("requirement_text") or ""), chapters)
            if chapter_title:
                bookmark = page_ref_registry.reserve("chapter", chapter_title)
                if bookmark:
                    row["target"] = chapter_title
                    page_refs.append({"kind": "chapter", "target": chapter_title, "bookmark": bookmark})
        row["page_refs"] = page_refs


def _ensure_scoring_index_chapter(
    chapters: list[dict[str, Any]],
    scoring_rows: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    render_chapters = [dict(ch) for ch in chapters]
    has_index = any(
        any(h in ch.get("title", "") for h in _INDEX_SECTION_HINTS)
        for ch in render_chapters
    )
    if scoring_rows and not has_index:
        render_chapters.append(
            {
                "section_type": "scoring_index",
                "title": "评分索引",
                "custom": True,
                "mapped_from": "generated_scoring_index",
                "attachments": [],
            }
        )
    return render_chapters


def _render_scoring_index_chapter(
    document: Any,
    rows: list[dict[str, Any]],
    *,
    include_review_status: bool,
) -> None:
    headers = _INDEX_REVIEW_HEADERS if include_review_status else _INDEX_HEADERS
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            _set_font(run, bold=True, size=Pt(10.5))
    if not rows:
        rows = [{"requirement_text": "", "target": "", "page": "", "status": "待应答"}]
    for index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        values = [
            str(index),
            str(row.get("requirement_text") or ""),
            str(row.get("target") or ""),
        ]
        if include_review_status:
            values.append(str(row.get("status") or "待应答"))
        for cell, value in zip(cells[:3], values[:3]):
            cell.text = value
            for run in cell.paragraphs[0].runs:
                _set_font(run, size=Pt(9.5))
        page_cell = cells[3]
        page_refs = [
            ref
            for ref in row.get("page_refs") or []
            if isinstance(ref, dict) and ref.get("bookmark")
        ]
        if page_refs:
            page_cell.text = ""
            paragraph = page_cell.paragraphs[0]
            for ref_index, ref in enumerate(page_refs):
                if ref_index:
                    sep = paragraph.add_run("、")
                    _set_font(sep, size=Pt(9.5))
                _add_pageref_field(paragraph, str(ref["bookmark"]))
        else:
            page_value = str(row.get("page") or "")
            if include_review_status and not page_value:
                page_value = "待定位"
            page_cell.text = page_value
            for run in page_cell.paragraphs[0].runs:
                _set_font(run, size=Pt(9.5))
        if include_review_status:
            status_cell = cells[4]
            status_cell.text = str(row.get("status") or "待应答")
            for run in status_cell.paragraphs[0].runs:
                _set_font(run, size=Pt(9.5))


def _render_deviation_chapter(
    document: Any,
    rows: list[dict[str, Any]],
    *,
    include_review_status: bool,
) -> None:
    table = document.add_table(rows=1, cols=len(_DEVIATION_HEADERS))
    table.style = "Table Grid"
    for i, h in enumerate(_DEVIATION_HEADERS):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            _set_font(run, bold=True, size=Pt(10.5))
    if not rows:
        for cell in table.add_row().cells:
            cell.text = ""
        if include_review_status:
            hint = document.add_paragraph()
            _set_font(
                hint.add_run("说明：『响应与偏离』应注明『响应』或『偏离』；逐条对应招标技术/商务条款填写。"),
                size=Pt(10.5),
            )
        return
    for index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        response_text = str(row.get("response_text") or "")
        deviation = str(row.get("deviation") or "")
        if include_review_status:
            response_text = response_text or "［待填:我方响应］"
            deviation = deviation or "［待确认］"
        values = [
            str(index),
            "",
            str(row.get("requirement_text") or ""),
            response_text,
            deviation,
            str(row.get("note") or ""),
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for run in cell.paragraphs[0].runs:
                _set_font(run, size=Pt(9.5))


def _format_decimal(value: Any) -> str:
    if value in (None, ""):
        return ""
    return str(value)


def _pricing_status_label(status: str) -> str:
    return {
        "matched": "合价匹配",
        "computed_total": "已按数量×单价计算",
        "missing_unit_price": "缺综合单价",
        "missing_quantity": "缺工程数量",
        "total_mismatch": "合价不一致",
        "pending_price": "待报价",
    }.get(status, status)


def _render_price_chapter(
    document: Any,
    report: PricingValidationReport,
    *,
    include_review_status: bool,
) -> None:
    headers = _PRICE_REVIEW_HEADERS if include_review_status else _PRICE_HEADERS
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            _set_font(run, bold=True, size=Pt(10.5))

    if not report.rows:
        for cell in table.add_row().cells:
            cell.text = ""
    else:
        for row in report.rows:
            cells = table.add_row().cells
            unit_price = _format_decimal(row.unit_price)
            line_total = _format_decimal(row.line_total)
            note = ""
            if include_review_status:
                unit_price = unit_price or "［待填:综合单价］"
                line_total = line_total or "［待计算］"
                if row.expected_total is not None and row.total_status == "total_mismatch":
                    note = f"应为 {row.expected_total}"
            values = [
                row.item_no,
                row.item_name,
                row.unit,
                _format_decimal(row.quantity),
                unit_price,
                line_total,
                note,
            ]
            if include_review_status:
                values.append(_pricing_status_label(row.total_status))
            for cell, value in zip(cells, values):
                cell.text = value
                for run in cell.paragraphs[0].runs:
                    _set_font(run, size=Pt(9.5))

    if include_review_status:
        summary = document.add_paragraph()
        if report.total_amount is not None:
            total_text = f"投标总价校验：{report.total_amount}"
            if report.budget_amount is not None:
                total_text += f"；预算/最高限价：{report.budget_amount}；状态：{report.budget_status}"
        else:
            total_text = "投标总价校验：待补齐综合单价/合价后计算"
        _set_font(summary.add_run(total_text), size=Pt(10.5), bold=True)
        for issue in report.issues:
            p = document.add_paragraph(style="List Bullet")
            _set_font(p.add_run(issue), size=Pt(9.5))


def _is_l4_narrative_chapter(chapter: dict[str, Any]) -> bool:
    section_type = str(chapter.get("section_type") or "")
    title = str(chapter.get("title") or "")
    if section_type.startswith("gas_"):
        return True
    return any(hint in title for hint in _L4_SECTION_HINTS)


def _fact_value(facts: dict[str, Any], key: str) -> str | None:
    value = facts.get(key)
    if value in (None, ""):
        return None
    return str(value)


def _known_fact_bullets(facts: dict[str, Any], pricing_report: PricingValidationReport) -> list[str]:
    pairs = [
        ("项目名称", _fact_value(facts, "project_name")),
        ("采购人/招标人", _fact_value(facts, "purchaser")),
        ("供应商", _fact_value(facts, "supplier_name")),
        ("计划工期", _fact_value(facts, "construction_period_days")),
        ("质量标准", _fact_value(facts, "quality_standard")),
        ("缺陷责任/保修期", _fact_value(facts, "warranty_period")),
        ("项目范围", _fact_value(facts, "project_scope")),
    ]
    bullets = [f"{label}：{value}" for label, value in pairs if value]
    if pricing_report.total_amount is not None:
        bullets.append(f"报价校验总额：{pricing_report.total_amount}")
    elif pricing_report.rows:
        bullets.append("报价清单：已识别工程量行，综合单价/合价仍需补齐或复核")
    return bullets


def _l4_boundary_lines(*, include_review_status: bool) -> list[str]:
    if include_review_status:
        return [
            "本章仅基于已解析招标文件、已确认企业资料、评分/技术响应条目和报价校验结果形成审阅稿。",
            "未在当前资料中明确的管径、压力等级、线路长度、管材型号、工程量、施工图、现场踏勘结论和设备配置，均不得由系统补写，应由人工依据招标人最终资料确认。",
        ]
    return [
        "本章内容以招标文件、合同文件、设计文件、施工图、工程量清单及经确认的施工组织设计为依据。",
        "未在投标文件中列明的技术参数、工程数量、设备配置和现场条件，均以招标人最终发布或确认的资料为准。",
    ]


def _l4_method_lines(title: str, *, gas_project: bool) -> list[str]:
    lines: list[str] = []
    if "限额设计" in title or "预算" in title:
        lines.extend(
            [
                "建立设计成果、施工图预算和工程量清单之间的复核链条，投标报价与后续实施范围保持一致。",
                "对暂未明确的清单项目、综合单价和暂列金额，不在本章推定，按招标文件和经确认的报价文件执行。",
            ]
        )
    if "进度" in title or "资源" in title:
        lines.extend(
            [
                "进度安排以招标工期、现场移交条件、材料到货和关键工序验收为约束，形成可调整的施工节点计划。",
                "资源配置依据施工图、工程量清单和审批后的施工组织设计确定，不预设未确认的队伍数量、设备型号或分段工期。",
            ]
        )
    if "质量" in title or "材料" in title or "管材" in title or "阀门" in title:
        lines.extend(
            [
                "材料进场前核对规格型号、质量证明文件、检测报告和外观质量，未经确认不得投入使用。",
                "关键材料、设备和工序验收资料随施工过程同步归档，作为竣工移交和质量追溯依据。",
            ]
        )
    if "沟槽" in title or "回填" in title or "路面" in title:
        lines.extend(
            [
                "沟槽开挖、支护、管道安装、接口处理、回填压实和路面恢复按施工图、审批方案及现场条件组织。",
                "地下既有管线、交通导改和周边设施保护措施须在施工前完成核查和交底。",
            ]
        )
    if "压力试验" in title or "吹扫" in title or "置换" in title or "验收" in title:
        lines.extend(
            [
                "压力试验、吹扫、置换和验收移交按招标文件、设计文件和现行管理要求组织，记录、签证和检测资料同步形成。",
                "未明确试验压力、介质、稳压时间等参数时，不在本章推定，按经确认的专项方案执行。",
            ]
        )
    if "安全" in title or "交通" in title or "文明" in title:
        lines.extend(
            [
                "施工前完成安全技术交底、风险识别和现场围护，涉及交通组织、既有管线保护的内容按审批方案实施。",
                "文明施工、环境保护和临时设施布置与现场条件匹配，过程问题闭环整改并留存记录。",
            ]
        )
    if "服务" in title or "缺陷责任" in title or "竣工资料" in title:
        lines.extend(
            [
                "缺陷责任期内建立响应、排查、处置和回访闭环，服务范围和期限以招标文件及合同约定为准。",
                "竣工资料按验收、移交、运维需要分类整理，确保资料真实、完整、可追溯。",
            ]
        )
    if "EPC" in title or "协同" in title:
        lines.extend(
            [
                "建立设计、采购、施工之间的接口管理机制，设计变更、材料采购和现场实施保持同源复核。",
                "施工图预算、工程量清单和现场签证纳入统一台账，避免范围、数量和价格口径不一致。",
            ]
        )
    if not lines:
        if gas_project:
            lines.extend(
                [
                    "燃气管道实施以测量放线、既有管线核查、沟槽开挖与支护、管道安装、接口质量控制、试验验收、回填恢复和资料移交为主线。",
                    "管材、连接方式、防腐、压力等级和分段长度以施工图、设计说明和工程量清单为准，当前资料未明确时不作推定。",
                ]
            )
        else:
            lines.extend(
                [
                    "本章围绕招标要求、施工准备、过程控制、验收移交和资料归档展开，确保技术响应与商务承诺保持一致。",
                    "涉及具体工程数量、材料规格、设备配置和现场条件的内容，以已确认资料和人工复核结果为准。",
                ]
            )
    return lines


def _render_l4_narrative_chapter(
    document: Any,
    *,
    title: str,
    facts: dict[str, Any],
    technical_rows: list[dict[str, Any]],
    scoring_rows: list[dict[str, Any]],
    pricing_report: PricingValidationReport,
    include_review_status: bool,
) -> None:
    gas_project = any(hint in title for hint in _GAS_SECTION_HINTS)
    intro = document.add_paragraph()
    project_name = _fact_value(facts, "project_name") or "本项目"
    _set_font(
        intro.add_run(f"本章围绕《{project_name}》的招标要求编制，内容以已确认事实和可追溯条款为边界。"),
        size=Pt(10.5),
    )

    document.add_paragraph("一、已知事实与编制边界")
    for line in _known_fact_bullets(facts, pricing_report) or ["当前资料未形成可自动填入的项目事实，需人工补充确认。"]:
        p = document.add_paragraph(style="List Bullet")
        _set_font(p.add_run(line), size=Pt(10.5))
    for line in _l4_boundary_lines(include_review_status=include_review_status):
        p = document.add_paragraph(style="List Bullet")
        _set_font(p.add_run(line), size=Pt(10.5))

    document.add_paragraph("二、招标条款响应关注点")
    focus_rows = [row for row in technical_rows[:6] if row.get("requirement_text")]
    if focus_rows:
        for row in focus_rows:
            p = document.add_paragraph(style="List Bullet")
            _set_font(p.add_run(str(row["requirement_text"])), size=Pt(10.5))
    else:
        p = document.add_paragraph(style="List Bullet")
        _set_font(p.add_run("当前未接入可逐条引用的技术响应条款，本章仅形成通用组织与控制框架。"), size=Pt(10.5))

    if scoring_rows:
        document.add_paragraph("三、评分点覆盖提示")
        for row in scoring_rows[:6]:
            text = str(row.get("requirement_text") or "")
            if row.get("target"):
                text += f"；已绑定证据：{row['target']}"
            p = document.add_paragraph(style="List Bullet")
            _set_font(p.add_run(text), size=Pt(10.5))

    document.add_paragraph("四、实施与控制措施")
    for line in _l4_method_lines(title, gas_project=gas_project):
        p = document.add_paragraph(style="List Bullet")
        _set_font(p.add_run(line), size=Pt(10.5))


def assemble_format_docx(
    *,
    text: str,
    chapters: list[dict[str, Any]],
    facts: dict[str, Any],
    project_name: str | None = None,
    supplier_name: str | None = None,
    available_materials: list[Any] | None = None,
    compliance_items: list[Any] | None = None,
    pricing_rows: list[Any] | None = None,
    budget_amount: Any = None,
    export_mode: str = "review",
) -> tuple[bytes, dict[str, Any]]:
    """装配投标文件 docx。chapters 为 map_nodes_to_chapters 输出（含 attachments）。

    available_materials：企业材料库可用材料名/类别（用于合规自检判定附件是否已备）。
    export_mode：
        review     审阅版，插入合规自检清单和附件状态标注；
        submission 正式版，不输出内部自检/风险/待办状态。
    """
    if export_mode not in {"review", "submission"}:
        raise ValueError("export_mode must be 'review' or 'submission'")
    include_review_status = export_mode == "review"
    scoring_rows = _scoring_index_rows(compliance_items)
    technical_rows = _technical_response_rows(compliance_items)
    pricing_report = build_pricing_report(
        text=text,
        pricing_rows=pricing_rows,
        budget_amount=budget_amount,
    )
    render_chapters = _ensure_scoring_index_chapter(chapters, scoring_rows)
    page_ref_registry = _PageRefRegistry()
    _attach_scoring_page_refs(
        scoring_rows,
        render_chapters,
        available_materials,
        page_ref_registry,
    )
    top_titles = [c["title"] for c in render_chapters]
    templates = extract_format_templates(text, top_titles)
    by_key: dict[str, FormTemplate] = {t.key: t for t in templates.values()}

    # 合规覆盖：分层自检（L1 完整 / L2、L3 骨架）
    coverage = compute_coverage(
        text=text, templates=templates, chapters=render_chapters,
        facts=facts, available_materials=available_materials,
        compliance_items=compliance_items,
    )
    status_by_title = {
        it.title: it for it in coverage.for_layer(CoverLayer.L1_FORMAT)
    }

    def form_for(title: str) -> FormTemplate | None:
        """按语义键匹配可填表单模板（声明/委托书/身份证明）。"""
        key = _title_to_key(title)
        if key in _FORM_TEMPLATE_KEYS and key in by_key:
            return by_key[key]
        return None

    document = WordDocument()
    _configure(document)
    _enable_field_updates_on_open(document)

    document.add_heading("响应文件组成", level=1)
    for idx, ch in enumerate(render_chapters, start=1):
        cn = "一二三四五六七八九十"[idx - 1] if idx <= 10 else str(idx)
        _set_font(document.add_paragraph(f"{cn}、{ch['title']}").runs[0])
    document.add_page_break()

    if include_review_status:
        _render_compliance_checklist(document, coverage)
        document.add_page_break()

    diag: dict[str, Any] = {
        "rendered": [],
        "templates_hit": list(templates.keys()),
        "coverage": coverage.summary(),
        "disqualifying_gaps": [i.title for i in coverage.disqualifying_gaps],
        "embedded_materials": [],
        "scoring_index_rows": scoring_rows,
        "technical_response_rows": technical_rows,
        "pricing": {
            "row_count": len(pricing_report.rows),
            "total_amount": str(pricing_report.total_amount) if pricing_report.total_amount is not None else None,
            "budget_amount": str(pricing_report.budget_amount) if pricing_report.budget_amount is not None else None,
            "budget_status": pricing_report.budget_status,
            "issues": pricing_report.issues,
        },
        "export_mode": export_mode,
        "review_checklist_included": include_review_status,
    }

    for idx, ch in enumerate(render_chapters, start=1):
        cn = "一二三四五六七八九十"[idx - 1] if idx <= 10 else str(idx)
        title = ch["title"]
        heading = document.add_heading(f"{cn}、{title}", level=1)
        page_ref_registry.add_to_paragraph(heading, kind="chapter", label=title)
        attachments = ch.get("attachments") or []

        # 标题路由优先（偏离/报价/索引/资格），表单仅留给声明类
        if any(h in title for h in _TABLE_SECTION_HINTS):
            _render_deviation_chapter(
                document,
                technical_rows,
                include_review_status=include_review_status,
            )
            rendered_as = "table_deviation"
        elif any(h in title for h in _PRICE_SECTION_HINTS):
            _render_price_chapter(
                document,
                pricing_report,
                include_review_status=include_review_status,
            )
            rendered_as = "table_price"
        elif any(h in title for h in _INDEX_SECTION_HINTS):
            _render_scoring_index_chapter(
                document,
                scoring_rows,
                include_review_status=include_review_status,
            )
            rendered_as = "table_index"
        elif (tpl := form_for(title)) is not None:
            _render_form(document, fill_template(tpl, facts).lines)
            rendered_as = "form"
        elif _is_l4_narrative_chapter(ch):
            _render_l4_narrative_chapter(
                document,
                title=title,
                facts=facts,
                technical_rows=technical_rows,
                scoring_rows=scoring_rows,
                pricing_report=pricing_report,
                include_review_status=include_review_status,
            )
            rendered_as = "l4_narrative"
        elif any(h in title for h in _ATTACH_SECTION_HINTS):
            embedded = _render_attachment_chapter(
                document,
                attachments,
                status_by_title,
                chapter_title=title,
                material_assets=available_materials,
                include_review_status=include_review_status,
                page_ref_registry=page_ref_registry,
            )
            diag["embedded_materials"].extend(embedded)
            rendered_as = "attachment"
            attachments = []  # 已在附件章内渲染
        else:
            _set_font(document.add_paragraph(f"［请人工撰写：{title}］").runs[0])
            rendered_as = "manual"

        # 附件：可填表单（身份证明/委托书）→ 填充；其余 → 扫描件占位
        for att in attachments:
            att_tpl = None
            akey = _title_to_key(att)
            if akey in _FORM_TEMPLATE_KEYS and akey in by_key:
                att_tpl = by_key[akey]
            if att_tpl is not None:
                document.add_heading(att, level=2)
                _render_form(document, fill_template(att_tpl, facts).lines)
            else:
                p = document.add_paragraph(style="List Bullet")
                suffix = "（扫描件待插入）" if include_review_status else ""
                _set_font(p.add_run(f"附：{att}{suffix}"), size=Pt(10.5))

        document.add_page_break()
        diag["rendered"].append({"title": title, "as": rendered_as})

    buf = BytesIO()
    document.save(buf)
    return buf.getvalue(), diag
