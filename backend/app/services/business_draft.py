from __future__ import annotations

import hashlib
import json
import re
import uuid
from datetime import UTC, datetime
from io import BytesIO
from typing import Any

from docx import Document as WordDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pydantic import BaseModel, Field, ValidationError
from sqlalchemy import delete, select
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models import (
    AuditLog,
    BidSection,
    BusinessDraftChapter,
    BusinessDraftEvidenceRef,
    ComplianceEvidenceBinding,
    ComplianceItem,
    DraftContextPack,
    DraftFactCheck,
    ExportFile,
    Project,
)
from app.services.evidence_policy import requires_enterprise_evidence
from app.services.llm_gateway import LLMGatewayError, chat_completion
from app.services.storage import put_object_bytes


class BusinessDraftError(Exception):
    pass


class LLMBusinessDraftResponse(BaseModel):
    content_text: str = Field(min_length=1)


_EXPORT_DROP_TERMS = (
    "字段填充草稿",
    "响应草稿",
    "已绑定证据",
    "事实性校验",
    "ContextPack",
    "MVP1.3",
)

_TABLE_CHAPTER_TITLES = {
    "投标函附录",
    "投标总价封面",
    "资格业绩汇总表",
    "投标人基本情况表",
    "近年财务状况表",
    "评分业绩汇总表",
    "项目管理班子配备情况表",
    "项目负责人简历表",
}

_SIGNATURE_PREFIXES = (
    "投标人：",
    "法定代表人",
    "日期：",
)

BID_DOC_MARGINS_CM = {
    "top": 2.5,
    "bottom": 2.5,
    "left": 3.0,
    "right": 2.5,
}


def _sanitize_export_paragraph(text: str) -> str | None:
    """Remove review-only wording before writing the bid Word body."""
    value = text.strip()
    if not value:
        return None
    if any(term in value for term in _EXPORT_DROP_TERMS):
        return None
    if re.match(r"^\d+[.、]\s*招标要求[:：]", value):
        return None
    value = re.sub(r"^我方响应[:：]\s*", "", value)
    value = re.sub(r"（证据[:：][^）]*）", "", value)
    value = re.sub(r"\(证据[:：][^)]*\)", "", value)
    value = re.sub(r"燃气项目模拟-[^，。；）\]\s]+", "相关证明材料", value)
    value = re.sub(r"（招标项目编号[:：]\s*section-\d+）", "", value)
    value = re.sub(r"招标项目编号[:：]\s*section-\d+[，,]?", "", value)
    value = value.replace("投标响应：", "").replace("我方响应：", "")
    value = value.strip()
    return value or None


def _chinese_order(index: int) -> str:
    digits = "零一二三四五六七八九"
    if index <= 0:
        return str(index)
    if index < 10:
        return digits[index]
    if index == 10:
        return "十"
    if index < 20:
        return f"十{digits[index - 10]}"
    if index < 100:
        ten, one = divmod(index, 10)
        return f"{digits[ten]}十{digits[one] if one else ''}"
    return str(index)


def _directory_lines(chapters: list[BusinessDraftChapter]) -> list[str]:
    titles = [
        chapter.title
        for chapter in chapters
        if chapter.title not in {"商务标封面", "投标文件封面", "商务标目录", "目录"}
    ]
    return [f"{_chinese_order(index)}、{title}" for index, title in enumerate(titles, start=1)]


def _set_run_font(run: Any, *, name: str = "宋体", size: Pt | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def _configure_bid_document(document: Any) -> None:
    for section in document.sections:
        section.top_margin = Cm(BID_DOC_MARGINS_CM["top"])
        section.bottom_margin = Cm(BID_DOC_MARGINS_CM["bottom"])
        section.left_margin = Cm(BID_DOC_MARGINS_CM["left"])
        section.right_margin = Cm(BID_DOC_MARGINS_CM["right"])

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    heading1 = document.styles["Heading 1"]
    heading1.font.name = "黑体"
    heading1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(12)


def _chapter_export_lines(chapter: BusinessDraftChapter) -> list[str]:
    lines: list[str] = []
    for paragraph in chapter.content_text.splitlines():
        text = _sanitize_export_paragraph(paragraph)
        if text and text != chapter.title:
            lines.append(text)
    return lines


def _split_label_value(text: str) -> tuple[str, str] | None:
    if re.match(r"^\d+[.、]", text):
        return None
    if "：" not in text:
        return None
    label, value = text.split("：", 1)
    label = label.strip()
    value = value.strip()
    if not label or not value or len(label) > 24:
        return None
    return label, value


def _chapter_label_rows(lines: list[str]) -> tuple[list[tuple[str, str]], list[str]]:
    rows: list[tuple[str, str]] = []
    remaining: list[str] = []
    for line in lines:
        parsed = _split_label_value(line)
        if parsed:
            rows.append(parsed)
        else:
            remaining.append(line)
    return rows, remaining


def _add_label_value_table(document: Any, rows: list[tuple[str, str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(4.2)
        cells[1].width = Cm(11.5)
        cells[0].text = label
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            _set_run_font(run, bold=True)
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run)


def _add_export_paragraph(document: Any, text: str) -> None:
    paragraph = document.add_paragraph(text)
    if text.startswith(_SIGNATURE_PREFIXES):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _render_cover_chapter(document: Any, chapter: BusinessDraftChapter) -> None:
    lines = _chapter_export_lines(chapter)
    title = "投 标 文 件"
    if lines and lines[0].replace(" ", "") == "投标文件":
        lines = lines[1:]
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(title)
    _set_run_font(title_run, name="黑体", size=Pt(22), bold=True)
    rows, remaining = _chapter_label_rows(lines)
    _add_label_value_table(document, rows)
    for line in remaining:
        _add_export_paragraph(document, line)
    document.add_page_break()


def _render_table_chapter(document: Any, chapter: BusinessDraftChapter) -> None:
    document.add_heading(chapter.title, level=1)
    lines = [line for line in _chapter_export_lines(chapter) if line != chapter.title]
    rows, remaining = _chapter_label_rows(lines)
    if rows:
        _add_label_value_table(document, rows)
    for line in remaining:
        _add_export_paragraph(document, line)


def _render_text_chapter(document: Any, chapter: BusinessDraftChapter) -> None:
    document.add_heading(chapter.title, level=1)
    for line in _chapter_export_lines(chapter):
        _add_export_paragraph(document, line)


def _item_type_label(value: str) -> str:
    return {
        "qualification": "资格要求",
        "mandatory_response": "强制响应",
        "format": "格式要求",
        "deadline": "截止时间",
        "scoring": "评分办法",
        "reference_info": "参考信息",
        "technical_response": "技术响应",
        "other": "其他",
    }.get(value, value)


def _risk_label(value: str) -> str:
    return {"high": "高", "medium": "中", "low": "低"}.get(value, value)


def _binding_snapshot(binding: ComplianceEvidenceBinding) -> dict[str, Any]:
    snapshot = binding.material_snapshot or {}
    return {
        "binding_id": str(binding.id),
        "enterprise_material_id": str(binding.enterprise_material_id),
        "material_name": snapshot.get("name"),
        "material_type": snapshot.get("material_type"),
        "verification_status": snapshot.get("verification_status"),
        "evidence_text": binding.evidence_text,
        "confidence_score": str(binding.confidence_score) if binding.confidence_score is not None else None,
    }


def _chapter_groups(items: list[ComplianceItem]) -> list[tuple[str, str, list[ComplianceItem]]]:
    qualification = [item for item in items if item.item_type == "qualification"]
    business = [
        item
        for item in items
        if item.item_type in {"mandatory_response", "deadline", "format", "scoring"}
    ]
    reference = [item for item in items if item.item_type in {"reference_info", "other"}]
    groups = [
        ("qualification_response", "资格响应", qualification),
        ("business_response", "商务响应与实质性条款", business),
        ("commitment", "承诺、偏离与参考事项", reference),
    ]
    return [(chapter_type, title, group_items) for chapter_type, title, group_items in groups if group_items]


def _legacy_bid_sentence(item: ComplianceItem) -> str:
    suggestion = (item.response_suggestion or "").strip()
    generic_markers = [
        "请在投标文件中逐项响应",
        "绑定对应证据来源",
        "绑定证明材料",
        "企业资质库",
        "人员库",
        "联合体资料",
        "已绑定证据",
        "响应草稿",
        "合规矩阵",
    ]
    defaults = {
        "qualification": "我方具备招标文件规定的相应资格条件，相关证明材料随本投标文件提交。",
        "mandatory_response": "我方承诺全面响应招标文件的实质性要求，并按招标文件及合同约定履行义务。",
        "format": "我方将按招标文件规定的格式、签章、加密和递交流程编制并提交投标文件。",
        "deadline": "我方承诺遵守招标文件规定的投标截止时间、投标有效期及相关时间安排。",
        "scoring": "我方已按评审办法要求编制相关响应资料，具体内容详见本投标文件相应章节。",
        "reference_info": "我方已充分理解招标文件相关说明，并将在履约过程中按招标人要求执行。",
        "other": "我方承诺按招标文件及相关补充文件要求提交本项目所需资料。",
    }
    if not suggestion or any(marker in suggestion for marker in generic_markers):
        suggestion = defaults.get(
            item.item_type,
            f"我方将按{_item_type_label(item.item_type)}要求提交资料并履行相应义务。",
        )
    suggestion = re.sub(r"\s+", " ", suggestion).strip()
    suggestion = suggestion.replace(" [请人工补充证据]", "").replace("[请人工补充证据]", "")
    if suggestion and suggestion[-1] not in "。；;":
        suggestion += "。"
    return suggestion


def _build_chapter_content(
    *,
    title: str,
    items: list[ComplianceItem],
    bindings_by_item: dict[uuid.UUID, list[ComplianceEvidenceBinding]],
) -> tuple[str, list[dict[str, Any]]]:
    lines = [f"{title}", ""]
    refs: list[dict[str, Any]] = []
    rendered_paragraphs: set[str] = set()
    for item in items:
        bindings = bindings_by_item.get(item.id, [])
        base_response = _legacy_bid_sentence(item)
        response = base_response
        if bindings:
            proof_sentence = "相关证明材料详见本投标文件资格审查资料或附件。"
            if proof_sentence not in response:
                response = f"{response}{proof_sentence}"
        elif requires_enterprise_evidence(item):
            response = "本项涉及需由企业资料证明的响应内容，当前未绑定可核验材料，暂不写成已满足结论。[请人工补充相关证明材料]"
        else:
            response = f"{response}[请人工补充相关证明材料]"
        if base_response not in rendered_paragraphs:
            rendered_paragraphs.add(base_response)
            lines.append(f"{len(rendered_paragraphs)}. {response}")
            lines.append("")
        if bindings:
            for binding in bindings:
                snapshot = _binding_snapshot(binding)
                refs.append(
                    {
                        "compliance_item_id": str(item.id),
                        "evidence_binding_id": str(binding.id),
                        "enterprise_material_id": str(binding.enterprise_material_id),
                        "source_type": "enterprise_material",
                        "source_snapshot": snapshot,
                        "quote_text": binding.evidence_text,
                    }
                )
        else:
            refs.append(
                {
                    "compliance_item_id": str(item.id),
                    "evidence_binding_id": None,
                    "enterprise_material_id": None,
                    "source_type": "compliance_item",
                    "source_snapshot": {
                        "requirement_text": item.requirement_text,
                        "item_type": item.item_type,
                        "risk_level": item.risk_level,
                        "source_page_no": item.source_page_no,
                    },
                    "quote_text": item.evidence_text or item.requirement_text,
                }
            )
    return "\n".join(lines).strip(), refs


def _json_from_model_text(content: str) -> dict[str, Any]:
    content = content.strip()
    fenced = re.search(r"```(?:json)?\s*(.*?)```", content, flags=re.DOTALL)
    if fenced:
        content = fenced.group(1).strip()
    return json.loads(content)


def _directive_prompt_block(directives: list[dict[str, Any]] | None) -> str:
    """Render author directives for the LLM prompt with a hard fact boundary.

    Style / emphasis steer wording and focus; mandatory_text must appear
    verbatim. None of them may introduce facts that lack evidence — the boundary
    statement makes that explicit and the downstream fact-check enforces it.
    """
    if not directives:
        return ""
    label = {"style": "风格", "emphasis": "内容侧重", "mandatory_text": "强制措辞（须原样保留）"}
    lines = ["人工生成指令（仅影响表达与侧重，禁止据此新增任何无证据事实）："]
    for directive in directives:
        directive_type = directive.get("directive_type")
        lines.append(f"- [{label.get(directive_type, directive_type)}] {directive.get('text')}")
    return "\n".join(lines)


def _chapter_prompt(
    *,
    project: Project,
    title: str,
    template_content: str,
    refs: list[dict[str, Any]],
    directives: list[dict[str, Any]] | None = None,
) -> list[dict[str, str]]:
    evidence_payload = [
        {
            "source_type": ref.get("source_type"),
            "quote_text": ref.get("quote_text"),
            "source_snapshot": ref.get("source_snapshot"),
        }
        for ref in refs[:80]
    ]
    directive_block = _directive_prompt_block(directives)
    directive_section = f"\n\n{directive_block}" if directive_block else ""
    return [
        {
            "role": "system",
            "content": (
                "你是商务标书草稿助手。只输出 JSON，不要输出解释。"
                "你只能基于输入的合规矩阵草稿和证据改写商务标章节，禁止编造项目名称、证书编号、人员、金额、日期、"
                "工程参数、工程量、管径、压力等级、路段长度、材料型号、项目编号或现场踏勘结论。"
                "无法从证据确认的事实必须写成 [请人工确认]。"
                "人工生成指令只能调整表达、侧重或保留指定措辞，不得作为新增事实的依据。"
            ),
        },
        {
            "role": "user",
            "content": (
                f"项目名称：{project.name}\n"
                f"章节标题：{title}\n\n"
                "请将下面的模板草稿整理为可人工复核的商务标章节正文，保留逐项响应关系，"
                "语言可以更正式，但不得新增没有证据的事实。\n\n"
                "输出 JSON 格式：{\"content_text\":\"章节正文\"}\n\n"
                f"模板草稿：\n{template_content}{directive_section}\n\n"
                f"证据：\n{json.dumps(evidence_payload, ensure_ascii=False)}"
            ),
        },
    ]


def _generate_chapter_content_with_llm(
    db: Session,
    *,
    tenant_id: uuid.UUID,
    project: Project,
    section_id: uuid.UUID,
    actor_user_id: uuid.UUID,
    title: str,
    template_content: str,
    refs: list[dict[str, Any]],
    item_count: int,
    directives: list[dict[str, Any]] | None = None,
) -> tuple[str, str, uuid.UUID | None]:
    messages = _chapter_prompt(
        project=project,
        title=title,
        template_content=template_content,
        refs=refs,
        directives=directives,
    )
    try:
        result = chat_completion(
            db,
            tenant_id=tenant_id,
            project_id=project.id,
            section_id=section_id,
            actor_user_id=actor_user_id,
            actor_type="user",
            task_type="business_draft_generation",
            prompt_version="business-draft-chapter@2026-05-17",
            messages=messages,
            complexity="complex" if item_count >= 12 or len(template_content) >= 5000 else "simple",
            temperature=0.1,
            response_format={"type": "json_object"},
            evidence_refs={
                "compliance_item_ids": [
                    ref["compliance_item_id"]
                    for ref in refs
                    if ref.get("compliance_item_id")
                ],
                "evidence_binding_ids": [
                    ref["evidence_binding_id"]
                    for ref in refs
                    if ref.get("evidence_binding_id")
                ],
            },
        )
        parsed = LLMBusinessDraftResponse.model_validate(_json_from_model_text(result.content))
        return parsed.content_text.strip(), f"{result.provider}:{result.model_name}", result.log_id
    except (LLMGatewayError, ValidationError, json.JSONDecodeError, KeyError, TypeError, ValueError):
        return template_content, "template", None


def _evidence_corpus(db: Session, chapter: BusinessDraftChapter) -> str:
    refs = db.scalars(
        select(BusinessDraftEvidenceRef).where(
            BusinessDraftEvidenceRef.tenant_id == chapter.tenant_id,
            BusinessDraftEvidenceRef.chapter_id == chapter.id,
        )
    ).all()
    values: list[str] = []
    for ref in refs:
        values.append(ref.quote_text or "")
        values.append(str(ref.source_snapshot or ""))
    generated = chapter.generated_from_json or {}
    form_facts = generated.get("verified_form_facts") or {}
    if isinstance(form_facts, dict):
        values.extend(str(value) for value in form_facts.values() if value not in (None, ""))
    return "\n".join(values)


_PERSON_ROLE_PATTERN = re.compile(
    r"(?:项目经理|项目负责人|技术负责人|项目总工|总监理工程师|注册建造师|建造师|造价工程师|项目总监)"
    r"(?:拟派|拟任|为|由|是|：|:|，|,|\s)+"
    r"([\u4e00-\u9fa5]{2,4})"
)
_ID_CARD_PATTERN = re.compile(r"(?<![0-9])[0-9]{17}[0-9Xx](?![0-9])")
_BUILDER_CERT_PATTERN = re.compile(
    r"(?:建造师|证书|执业)\s*(?:编号|证号)[：:，,为是\s]*([A-Za-z0-9]{6,32})"
)
_PERFORMANCE_AMOUNT_PATTERN = re.compile(
    r"(?:合同金额|中标金额|合同额|工程造价|合同价)[：:，,\s]*(\d+(?:\.\d+)?\s*(?:万元|亿元|元))"
)
_PERFORMANCE_CONTRACT_PATTERN = re.compile(
    r"(?:业绩|类似工程|类似项目|合同名称|工程名称)[：:，,\s]*《([^》]{2,60})》"
)
_ENGINEERING_PARAMETER_PATTERN = re.compile(
    r"(?<![A-Za-z0-9])(?:DN\s*\d{2,4}|\d+(?:\.\d+)?\s*(?:km|KM|公里|MPa|Mpa|kPa|KPa|mm|毫米|米))(?![A-Za-z0-9])"
)
_ENGINEERING_STANDARD_PATTERN = re.compile(
    r"(?<![A-Za-z0-9])(?:GB\s*\d{4,5}(?:-\d{4})?|GB1\s*级?|中压\s*[AB]?|高压|低压)(?![A-Za-z0-9])"
)
_PERSON_NAME_STOPWORDS = {"我方", "详见", "见附", "拟派", "拟任", "暂定", "待定", "无", "略"}


def _personnel_fact_candidates(text: str) -> list[tuple[str, str]]:
    """Detect personnel facts (project manager / certified builder names + ids).

    MVP1.3 #7 requires personnel facts to be fact-checked. Names are only taken
    when they follow an explicit role keyword + separator so we do not flag
    arbitrary two-character phrases, and obvious filler words are dropped.
    """
    out: list[tuple[str, str]] = []
    for name in _PERSON_ROLE_PATTERN.findall(text):
        if name and name not in _PERSON_NAME_STOPWORDS:
            out.append(("person_name", name))
    for id_no in _ID_CARD_PATTERN.findall(text):
        out.append(("person_name", id_no))
    for cert_no in _BUILDER_CERT_PATTERN.findall(text):
        out.append(("certificate_no", cert_no))
    return out


def _performance_fact_candidates(text: str) -> list[tuple[str, str]]:
    """Detect performance facts (contract amounts and similar-project names).

    MVP1.3 #7 requires 业绩 (track-record) facts to be fact-checked so they are
    not fabricated. Contract amounts map to ``amount`` and contract/project names
    to ``other`` so they verify against bound evidence like any other fact.
    """
    out: list[tuple[str, str]] = []
    for amount in _PERFORMANCE_AMOUNT_PATTERN.findall(text):
        out.append(("amount", amount.strip()))
    for contract in _PERFORMANCE_CONTRACT_PATTERN.findall(text):
        out.append(("other", contract.strip()))
    return out


def _engineering_fact_candidates(text: str) -> list[tuple[str, str]]:
    """Detect engineering parameters that must stay evidence-backed.

    Competitor-style drafts often infer route splits, pressure levels or
    standards from a thin tender notice. Treat these as checkable facts so
    unsupported values such as ``5.23km`` or ``高压`` are converted to human
    confirmation placeholders before export.
    """
    out: list[tuple[str, str]] = []
    for value in _ENGINEERING_PARAMETER_PATTERN.findall(text):
        out.append(("number", re.sub(r"\s+", "", value)))
    for value in _ENGINEERING_STANDARD_PATTERN.findall(text):
        out.append(("other", re.sub(r"\s+", "", value)))
    return out


def _fact_supported_by_corpus(fact_type: str, fact_text: str, corpus: str) -> bool:
    if fact_text in corpus or fact_type == "project_name":
        return True
    if fact_type in {"number", "other"}:
        normalized_fact = re.sub(r"\s+", "", fact_text).upper()
        normalized_corpus = re.sub(r"\s+", "", corpus).upper()
        return normalized_fact in normalized_corpus
    return False


def recompose_chapter_text_from_blocks(db: Session, chapter: BusinessDraftChapter) -> str:
    """Rebuild a chapter's ``content_text`` from its current draft blocks.

    With MVP1.3 per-clause blocks a single block no longer represents the whole
    chapter, so editing one block must not overwrite the entire chapter body.
    We concatenate the chapter's blocks in sort order; the caller persists the
    result and re-runs fact checks on the recomposed text.
    """
    from app.models import DraftBlock

    blocks = db.scalars(
        select(DraftBlock)
        .where(
            DraftBlock.tenant_id == chapter.tenant_id,
            DraftBlock.chapter_id == chapter.id,
        )
        .order_by(DraftBlock.sort_order.asc(), DraftBlock.created_at.asc())
    ).all()
    segments = [
        block.content_text.strip()
        for block in blocks
        if block.block_type != "heading" and block.content_text.strip()
    ]
    if not segments:
        return chapter.content_text
    return "\n\n".join(segments)


def run_fact_checks(
    db: Session,
    *,
    chapter: BusinessDraftChapter,
    project: Project,
    actor_user_id: uuid.UUID,
) -> list[DraftFactCheck]:
    db.execute(
        delete(DraftFactCheck).where(
            DraftFactCheck.tenant_id == chapter.tenant_id,
            DraftFactCheck.chapter_id == chapter.id,
        )
    )
    corpus = _evidence_corpus(db, chapter)
    checks: list[DraftFactCheck] = []

    candidates: list[tuple[str, str]] = []
    if project.name in chapter.content_text:
        candidates.append(("project_name", project.name))
    for value in sorted(set(re.findall(r"[A-Z0-9]{8,32}", chapter.content_text))):
        candidates.append(("certificate_no", value))
    for value in sorted(set(re.findall(r"\d+(?:\.\d+)?\s*(?:万元|元|个月|天|日历日|年)", chapter.content_text))):
        candidates.append(("number", value))
    for value in sorted(set(re.findall(r"\d{4}[-年]\d{1,2}[-月]\d{1,2}日?", chapter.content_text))):
        candidates.append(("date", value))
    candidates.extend(_personnel_fact_candidates(chapter.content_text))
    candidates.extend(_performance_fact_candidates(chapter.content_text))
    candidates.extend(_engineering_fact_candidates(chapter.content_text))

    seen: set[tuple[str, str]] = set()
    for fact_type, fact_text in candidates:
        key = (fact_type, fact_text)
        if key in seen:
            continue
        seen.add(key)
        verified = _fact_supported_by_corpus(fact_type, fact_text, corpus)
        check = DraftFactCheck(
            tenant_id=chapter.tenant_id,
            project_id=chapter.project_id,
            section_id=chapter.section_id,
            chapter_id=chapter.id,
            fact_type=fact_type,
            fact_text=fact_text,
            check_status="verified" if verified else "unverified",
            risk_level="low" if verified else "high",
            evidence_text=fact_text if verified else None,
            detail="已在项目基础信息或绑定证据中找到一致事实。"
            if verified
            else "未在当前章节绑定证据中找到一致事实，正式使用前需人工确认。",
            created_by=actor_user_id,
        )
        db.add(check)
        checks.append(check)

    if not checks:
        check = DraftFactCheck(
            tenant_id=chapter.tenant_id,
            project_id=chapter.project_id,
            section_id=chapter.section_id,
            chapter_id=chapter.id,
            fact_type="other",
            fact_text="未发现需自动校验的关键事实",
            check_status="verified",
            risk_level="low",
            evidence_text=None,
            detail="当前章节没有识别到证书编号、金额、日期等高风险事实。",
            created_by=actor_user_id,
        )
        db.add(check)
        checks.append(check)

    unverified_values = [
        item.fact_text for item in checks if item.check_status == "unverified" and item.fact_text
    ]
    updated_content = chapter.content_text
    for value in sorted(unverified_values, key=len, reverse=True):
        updated_content = updated_content.replace(value, "[请人工确认]")
    if updated_content != chapter.content_text:
        chapter.content_text = updated_content
    chapter.fact_check_status = (
        "unverified" if any(item.check_status == "unverified" for item in checks) else "verified"
    )
    return checks


def generate_business_draft_chapters(
    db: Session,
    *,
    tenant_id: uuid.UUID,
    project_id: uuid.UUID,
    section_id: uuid.UUID,
    actor_user_id: uuid.UUID,
) -> list[BusinessDraftChapter]:
    project = db.get(Project, project_id)
    if project is None:
        raise BusinessDraftError("项目不存在")

    items = db.scalars(
        select(ComplianceItem)
        .where(
            ComplianceItem.tenant_id == tenant_id,
            ComplianceItem.project_id == project_id,
            ComplianceItem.section_id == section_id,
            ComplianceItem.deleted_at.is_(None),
        )
        .order_by(ComplianceItem.item_type.asc(), ComplianceItem.risk_level.desc(), ComplianceItem.created_at.asc())
    ).all()
    if not items:
        raise BusinessDraftError("当前标段没有可生成草稿的合规矩阵项")

    bindings = db.scalars(
        select(ComplianceEvidenceBinding).where(
            ComplianceEvidenceBinding.tenant_id == tenant_id,
            ComplianceEvidenceBinding.project_id == project_id,
            ComplianceEvidenceBinding.section_id == section_id,
            ComplianceEvidenceBinding.status == "active",
        )
    ).all()
    bindings_by_item: dict[uuid.UUID, list[ComplianceEvidenceBinding]] = {}
    for binding in bindings:
        bindings_by_item.setdefault(binding.compliance_item_id, []).append(binding)

    existing = db.scalars(
        select(BusinessDraftChapter).where(
            BusinessDraftChapter.tenant_id == tenant_id,
            BusinessDraftChapter.project_id == project_id,
            BusinessDraftChapter.section_id == section_id,
            BusinessDraftChapter.status != "superseded",
        )
    ).all()
    for chapter in existing:
        chapter.status = "superseded"

    latest_pack = db.scalar(
        select(DraftContextPack)
        .where(
            DraftContextPack.tenant_id == tenant_id,
            DraftContextPack.project_id == project_id,
            DraftContextPack.section_id == section_id,
            DraftContextPack.status != "superseded",
        )
        .order_by(DraftContextPack.created_at.desc())
    )
    pack_directives = [
        directive
        for directive in ((latest_pack.context_json if latest_pack else {}) or {}).get(
            "author_directives"
        )
        or []
        if directive.get("scope") == "pack"
    ]

    chapters: list[BusinessDraftChapter] = []
    for sort_order, (chapter_type, title, group_items) in enumerate(_chapter_groups(items), start=1):
        template_content, refs = _build_chapter_content(
            title=title,
            items=group_items,
            bindings_by_item=bindings_by_item,
        )
        content_text, generator, invocation_log_id = _generate_chapter_content_with_llm(
            db,
            tenant_id=tenant_id,
            project=project,
            section_id=section_id,
            actor_user_id=actor_user_id,
            title=title,
            template_content=template_content,
            refs=refs,
            item_count=len(group_items),
            directives=pack_directives,
        )
        chapter = BusinessDraftChapter(
            tenant_id=tenant_id,
            project_id=project_id,
            section_id=section_id,
            chapter_type=chapter_type,
            title=title,
            sort_order=sort_order,
            content_text=content_text,
            outline_json={"item_count": len(group_items), "item_types": sorted({item.item_type for item in group_items})},
            evidence_summary_json={
                "bound_evidence_count": sum(
                    len(bindings_by_item.get(item.id, [])) for item in group_items
                ),
                "unbound_item_count": sum(1 for item in group_items if not bindings_by_item.get(item.id)),
            },
            fact_check_status="pending",
            status="pending_review",
            version_no=1,
            generated_from_json={
                "compliance_item_ids": [str(item.id) for item in group_items],
                "generated_at": datetime.now(UTC).isoformat(),
                "generator": generator,
                "model_invocation_log_id": str(invocation_log_id) if invocation_log_id else None,
            },
            created_by=actor_user_id,
            updated_by=actor_user_id,
        )
        db.add(chapter)
        db.flush()
        for ref in refs:
            db.add(
                BusinessDraftEvidenceRef(
                    tenant_id=tenant_id,
                    project_id=project_id,
                    section_id=section_id,
                    chapter_id=chapter.id,
                    compliance_item_id=uuid.UUID(ref["compliance_item_id"])
                    if ref.get("compliance_item_id")
                    else None,
                    evidence_binding_id=uuid.UUID(ref["evidence_binding_id"])
                    if ref.get("evidence_binding_id")
                    else None,
                    enterprise_material_id=uuid.UUID(ref["enterprise_material_id"])
                    if ref.get("enterprise_material_id")
                    else None,
                    source_type=ref["source_type"],
                    source_snapshot=ref["source_snapshot"],
                    quote_text=ref.get("quote_text"),
                )
            )
        db.flush()
        run_fact_checks(db, chapter=chapter, project=project, actor_user_id=actor_user_id)
        chapters.append(chapter)

    db.add(
        AuditLog(
            tenant_id=tenant_id,
            project_id=project_id,
            section_id=section_id,
            actor_user_id=actor_user_id,
            actor_type="user",
            action="business_draft.generated",
            object_type="business_draft_chapter",
            object_id=None,
            after_json={"chapter_count": len(chapters), "chapter_ids": [str(chapter.id) for chapter in chapters]},
            reason="基于合规矩阵和企业证据生成商务标章节草稿",
            severity="info",
        )
    )
    return chapters


def export_business_draft_word(
    db: Session,
    *,
    tenant_id: uuid.UUID,
    project_id: uuid.UUID,
    section_id: uuid.UUID,
    actor_user_id: uuid.UUID,
    extra_snapshot: dict[str, object] | None = None,
) -> ExportFile:
    project = db.get(Project, project_id)
    section = db.get(BidSection, section_id)
    chapters = db.scalars(
        select(BusinessDraftChapter)
        .where(
            BusinessDraftChapter.tenant_id == tenant_id,
            BusinessDraftChapter.project_id == project_id,
            BusinessDraftChapter.section_id == section_id,
            BusinessDraftChapter.status != "superseded",
        )
        .order_by(BusinessDraftChapter.sort_order.asc(), BusinessDraftChapter.created_at.asc())
    ).all()
    if not chapters:
        raise BusinessDraftError("尚无可导出的商务标章节草稿")

    document = WordDocument()
    _configure_bid_document(document)
    has_cover_chapter = any(chapter.title in {"商务标封面", "投标文件封面"} for chapter in chapters)
    if not has_cover_chapter:
        document.add_heading("投 标 文 件", level=0)
        if project:
            document.add_paragraph(f"项目名称：{project.name}")
        document.add_paragraph(f"标段名称：{section.name if section else '未识别标段'}")
        document.add_paragraph("文件内容：商务标")
        document.add_page_break()
    for chapter in chapters:
        is_cover_chapter = chapter.title in {"商务标封面", "投标文件封面"}
        is_directory_chapter = chapter.title in {"商务标目录", "目录"}
        if is_cover_chapter:
            _render_cover_chapter(document, chapter)
            continue
        if is_directory_chapter:
            document.add_heading("目录", level=1)
            for line in _directory_lines(chapters):
                document.add_paragraph(line)
            document.add_page_break()
            continue
        if chapter.title in _TABLE_CHAPTER_TITLES:
            _render_table_chapter(document, chapter)
            continue
        _render_text_chapter(document, chapter)

    buffer = BytesIO()
    document.save(buffer)
    data = buffer.getvalue()
    content_hash = hashlib.sha256(data).hexdigest()
    export_id = uuid.uuid4()
    now = datetime.now(UTC)
    file_name = f"商务标章节草稿-{now.strftime('%Y%m%d%H%M%S')}.docx"
    object_key = (
        f"tenant/{tenant_id}/project/{project_id}/section/{section_id}/"
        f"exports/{export_id}/{file_name}"
    )
    put_object_bytes(
        bucket=settings.minio_bucket,
        object_key=object_key,
        data=data,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    source_snapshot = {
        "chapter_ids": [str(chapter.id) for chapter in chapters],
        "chapter_count": len(chapters),
        "exported_at": now.isoformat(),
        "snapshot_note": "商务标章节草稿，需人工复核后使用",
    }
    if extra_snapshot:
        source_snapshot.update(extra_snapshot)

    export_file = ExportFile(
        id=export_id,
        tenant_id=tenant_id,
        project_id=project_id,
        section_id=section_id,
        task_id=None,
        export_type="business_draft_word",
        file_name=file_name,
        bucket=settings.minio_bucket,
        object_key=object_key,
        sha256=content_hash,
        filter_json=None,
        source_snapshot_json=source_snapshot,
        status="available",
        created_by=actor_user_id,
    )
    db.add(export_file)
    db.add(
        AuditLog(
            tenant_id=tenant_id,
            project_id=project_id,
            section_id=section_id,
            actor_user_id=actor_user_id,
            actor_type="user",
            action="business_draft.word_exported",
            object_type="export_file",
            object_id=export_file.id,
            after_json={"file_name": file_name, "sha256": content_hash, "source_snapshot": source_snapshot},
            reason="导出商务标章节草稿 Word 文件",
            severity="warning" if source_snapshot.get("preflight_status") == "block" else "info",
        )
    )
    return export_file
