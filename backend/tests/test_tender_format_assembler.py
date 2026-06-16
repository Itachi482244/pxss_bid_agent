from __future__ import annotations

import base64
from io import BytesIO
from typing import Any
from zipfile import ZipFile

from docx import Document as WordDocument
from docx.shared import Cm

from app.services.tender_format_assembler import assemble_format_docx


SAMPLE_TEXT = """第五章 响应文件组成
供应商的响应文件应包含以下部分：
一、磋商响应声明
二、供应商的资格证明资料
三、技术/商务响应与偏离表
一、磋商响应声明
致 （采购人、采购代理机构）：
根据贵方为 （项目名称）的磋商邀请（采购代理编号： ），签字代表 （姓名、职务）经正式授权并代表供应商 （供应商名称）提交响应文件。
供应商名称（盖单位公章）：
日期： 年 月 日
二、供应商的资格证明资料
供应商应提供营业执照、信用查询等证明材料。
第五章 评分办法
1 技术方案完整、合理的得10分。
"""

CHAPTERS = [
    {"section_type": "bid_commitment", "title": "磋商响应声明", "attachments": []},
    {
        "section_type": "qualification",
        "title": "供应商的资格证明资料",
        "attachments": ["附件2-1 营业执照副本", "附件2-2 信用查询"],
    },
    {"section_type": "deviation", "title": "技术/商务响应与偏离表", "attachments": []},
]

FACTS = {
    "purchaser": "岳阳市粮食购销有限公司",
    "project_name": "仓顶面吊顶隔热降温改造",
    "agent_code": "HNZJ-TEST",
    "signatory": "胡扩军",
    "supplier_name": "湖南远发建筑工程有限公司",
    "date": "2026-06-16",
}

PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def _doc_text(data: bytes) -> str:
    document = WordDocument(BytesIO(data))
    parts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _assemble_bytes(**kwargs: Any) -> tuple[bytes, dict[str, Any]]:
    available_materials = kwargs.pop("available_materials", [])
    chapters = kwargs.pop("chapters", CHAPTERS)
    facts = kwargs.pop("facts", FACTS)
    data, diag = assemble_format_docx(
        text=SAMPLE_TEXT,
        chapters=chapters,
        facts=facts,
        available_materials=available_materials,
        **kwargs,
    )
    return data, diag


def _assemble(**kwargs: Any) -> tuple[str, dict[str, Any]]:
    data, diag = _assemble_bytes(**kwargs)
    return _doc_text(data), diag


def _docx_xml(data: bytes, name: str = "word/document.xml") -> str:
    with ZipFile(BytesIO(data)) as archive:
        return archive.read(name).decode("utf-8")


def test_review_export_includes_compliance_checklist_and_status_tags() -> None:
    text, diag = _assemble(export_mode="review")

    assert diag["export_mode"] == "review"
    assert diag["review_checklist_included"] is True
    assert "合规自检清单" in text
    assert "废标风险" in text
    assert "需上传材料" in text
    assert "扫描件待插入" in text or "正式投标前需替换" in text


def test_submission_export_omits_review_only_checklist_and_risk_words() -> None:
    text, diag = _assemble(export_mode="submission")

    assert diag["export_mode"] == "submission"
    assert diag["review_checklist_included"] is False
    assert diag["disqualifying_gaps"]
    assert "合规自检清单" not in text
    assert "废标风险" not in text
    assert "需上传材料" not in text
    assert "扫描件待插入" not in text
    assert "以下证明材料按招标文件要求随本章提交" in text


def test_unknown_export_mode_is_rejected() -> None:
    try:
        assemble_format_docx(
            text=SAMPLE_TEXT,
            chapters=CHAPTERS,
            facts=FACTS,
            export_mode="external",
        )
    except ValueError as exc:
        assert "export_mode" in str(exc)
    else:
        raise AssertionError("unknown export mode should be rejected")


def test_confirmed_material_image_is_embedded_into_attachment_chapter() -> None:
    data, diag = _assemble_bytes(
        export_mode="review",
        available_materials=[
            {
                "material_id": "license-1",
                "material_name": "营业执照副本",
                "material_type": "license",
                "verification_status": "confirmed",
                "data_level": "internal",
                "embedded_images": [
                    {"data": PNG_1X1, "content_type": "image/png", "caption": "营业执照第1页"}
                ],
            }
        ],
    )

    document = WordDocument(BytesIO(data))
    assert len(document.inline_shapes) == 1
    assert document.inline_shapes[0].width == Cm(14.0)
    text = _doc_text(data)
    assert "材料：营业执照副本" in text
    assert "营业执照第1页" in text
    assert diag["embedded_materials"] == [
        {
            "attachment": "附件2-1 营业执照副本",
            "material_id": "license-1",
            "material_name": "营业执照副本",
            "image_count": 1,
        }
    ]


def test_candidate_material_image_is_not_embedded() -> None:
    data, diag = _assemble_bytes(
        export_mode="review",
        available_materials=[
            {
                "material_id": "license-candidate",
                "material_name": "营业执照副本",
                "material_type": "license",
                "verification_status": "pending_confirm",
                "data_level": "internal",
                "embedded_images": [
                    {"data": PNG_1X1, "content_type": "image/png", "caption": "候选营业执照"}
                ],
            }
        ],
    )

    document = WordDocument(BytesIO(data))
    assert len(document.inline_shapes) == 0
    assert diag["embedded_materials"] == []


def test_scoring_items_generate_review_scoring_index_chapter() -> None:
    text, diag = _assemble(
        export_mode="review",
        compliance_items=[
            {
                "item_type": "scoring",
                "requirement_text": "类似业绩每个得2分，最高8分。",
                "status": "confirmed",
                "evidence_bindings": [
                    {"status": "active", "material_name": "近三年类似业绩合同"}
                ],
            }
        ],
    )

    assert "评分索引" in text
    assert "覆盖状态" in text
    assert "类似业绩每个得2分，最高8分。" in text
    assert "近三年类似业绩合同" in text
    assert diag["scoring_index_rows"][0]["status"] == "已绑定证据"
    assert any(item["title"] == "评分索引" for item in diag["rendered"])


def test_submission_scoring_index_omits_review_status_column() -> None:
    text, diag = _assemble(
        export_mode="submission",
        compliance_items=[
            {
                "item_type": "scoring",
                "requirement_text": "技术方案完整、合理得10分。",
                "status": "pending_confirm",
                "evidence_bindings": [],
            }
        ],
    )

    assert "评分索引" in text
    assert "技术方案完整、合理得10分。" in text
    assert "覆盖状态" not in text
    assert "待应答" not in text
    assert diag["scoring_index_rows"][0]["status"] == "待应答"


def test_scoring_index_page_uses_pageref_for_matched_chapter() -> None:
    chapters = [
        *CHAPTERS,
        {"section_type": "technical_solution", "title": "技术方案", "attachments": []},
    ]
    data, diag = _assemble_bytes(
        export_mode="submission",
        chapters=chapters,
        compliance_items=[
            {
                "item_type": "scoring",
                "requirement_text": "技术方案完整、合理得10分。",
                "status": "pending_confirm",
                "evidence_bindings": [],
            }
        ],
    )

    row = diag["scoring_index_rows"][0]
    assert row["target"] == "技术/商务响应与偏离表" or row["target"] == "技术方案"
    assert row["page_refs"][0]["kind"] == "chapter"
    document_xml = _docx_xml(data)
    settings_xml = _docx_xml(data, "word/settings.xml")
    assert f"PAGEREF {row['page_refs'][0]['bookmark']}" in document_xml
    assert f'w:name="{row["page_refs"][0]["bookmark"]}"' in document_xml
    assert "w:updateFields" in settings_xml


def test_scoring_index_page_uses_pageref_for_embedded_material() -> None:
    chapters = [
        CHAPTERS[0],
        {
            "section_type": "qualification",
            "title": "供应商的资格证明资料",
            "attachments": ["附件2-3 类似业绩证明材料"],
        },
        CHAPTERS[2],
    ]
    data, diag = _assemble_bytes(
        export_mode="review",
        chapters=chapters,
        available_materials=[
            {
                "material_id": "performance-1",
                "material_name": "近三年类似业绩合同",
                "material_type": "performance",
                "verification_status": "confirmed",
                "data_level": "internal",
                "embedded_images": [
                    {"data": PNG_1X1, "content_type": "image/png", "caption": "业绩合同第1页"}
                ],
            }
        ],
        compliance_items=[
            {
                "item_type": "scoring",
                "requirement_text": "类似业绩每个得2分，最高8分。",
                "status": "confirmed",
                "evidence_bindings": [
                    {"status": "active", "material_name": "近三年类似业绩合同"}
                ],
            }
        ],
    )

    row = diag["scoring_index_rows"][0]
    assert row["target"] == "近三年类似业绩合同"
    assert row["page_refs"][0]["kind"] == "material"
    document_xml = _docx_xml(data)
    assert f"PAGEREF {row['page_refs'][0]['bookmark']}" in document_xml
    assert f'w:name="{row["page_refs"][0]["bookmark"]}"' in document_xml
    assert "材料：近三年类似业绩合同" in _doc_text(data)


def test_scoring_index_does_not_pageref_material_without_embeddable_image() -> None:
    chapters = [
        CHAPTERS[0],
        {
            "section_type": "qualification",
            "title": "供应商的资格证明资料",
            "attachments": ["附件2-3 类似业绩证明材料"],
        },
        CHAPTERS[2],
    ]
    data, diag = _assemble_bytes(
        export_mode="review",
        chapters=chapters,
        available_materials=[
            {
                "material_id": "performance-empty",
                "material_name": "近三年类似业绩合同",
                "material_type": "performance",
                "verification_status": "confirmed",
                "data_level": "internal",
                "embedded_images": [],
            }
        ],
        compliance_items=[
            {
                "item_type": "scoring",
                "requirement_text": "类似业绩每个得2分，最高8分。",
                "status": "confirmed",
                "evidence_bindings": [
                    {"status": "active", "material_name": "近三年类似业绩合同"}
                ],
            }
        ],
    )

    row = diag["scoring_index_rows"][0]
    assert row["target"] == "近三年类似业绩合同"
    assert row["page_refs"] == []
    document_xml = _docx_xml(data)
    assert "PAGEREF" not in document_xml
    assert "材料：近三年类似业绩合同" not in _doc_text(data)
    assert diag["embedded_materials"] == []


def test_technical_response_items_render_deviation_rows_in_review() -> None:
    text, diag = _assemble(
        export_mode="review",
        compliance_items=[
            {
                "item_type": "technical_response",
                "requirement_text": "燃气管道设计压力应满足0.4MPa要求。",
                "status": "pending_confirm",
                "evidence_bindings": [],
            }
        ],
    )

    assert "燃气管道设计压力应满足0.4MPa要求。" in text
    assert "［待填:我方响应］" in text
    assert "［待确认］" in text
    assert diag["technical_response_rows"][0]["requirement_text"] == "燃气管道设计压力应满足0.4MPa要求。"


def test_submission_deviation_rows_keep_response_cells_blank() -> None:
    text, diag = _assemble(
        export_mode="submission",
        compliance_items=[
            {
                "item_type": "technical_response",
                "requirement_text": "管材应符合现行国家标准。",
                "status": "pending_confirm",
                "evidence_bindings": [],
            }
        ],
    )

    assert "管材应符合现行国家标准。" in text
    assert "［待填:我方响应］" not in text
    assert "［待确认］" not in text
    assert diag["technical_response_rows"][0]["response_text"] == ""


def test_submission_empty_deviation_chapter_omits_helper_text() -> None:
    text, _ = _assemble(export_mode="submission")

    assert "说明：『响应与偏离』" not in text
    assert "逐条对应招标技术/商务条款填写" not in text


def test_review_price_chapter_renders_boq_rows_and_validation_gaps() -> None:
    chapters = [*CHAPTERS, {"section_type": "bid_price_cover", "title": "报价一览表", "attachments": []}]
    text, diag = _assemble(
        export_mode="review",
        chapters=chapters,
        pricing_rows=[
            {"item_name": "DN100燃气管道安装", "unit": "米", "quantity": "120.5"},
            {"item_name": "阀门井砌筑", "unit": "座", "quantity": "2", "unit_price": "3000"},
        ],
        budget_amount="10000",
    )

    assert "DN100燃气管道安装" in text
    assert "［待填:综合单价］" in text
    assert "缺综合单价" in text
    assert "投标总价校验：待补齐综合单价/合价后计算" in text
    assert diag["pricing"]["row_count"] == 2
    assert diag["pricing"]["budget_status"] == "pending_prices"


def test_submission_price_chapter_omits_review_validation_words() -> None:
    chapters = [*CHAPTERS, {"section_type": "bid_price_cover", "title": "报价一览表", "attachments": []}]
    text, diag = _assemble(
        export_mode="submission",
        chapters=chapters,
        pricing_rows=[
            {"item_name": "DN100燃气管道安装", "unit": "米", "quantity": "10"},
        ],
        budget_amount="10000",
    )

    assert "DN100燃气管道安装" in text
    assert "［待填:综合单价］" not in text
    assert "缺综合单价" not in text
    assert "投标总价校验" not in text
    assert diag["pricing"]["issues"]


def test_l4_gas_narrative_uses_fact_boundary_without_fabricating_parameters() -> None:
    chapters = [
        *CHAPTERS,
        {
            "section_type": "gas_pipeline_construction_method",
            "title": "中压GB1燃气管道施工组织方案",
            "attachments": [],
        },
    ]
    facts = {
        **FACTS,
        "construction_period_days": "90日历天",
        "quality_standard": "合格",
    }
    text, diag = _assemble(
        export_mode="review",
        chapters=chapters,
        facts=facts,
        compliance_items=[
            {
                "item_type": "technical_response",
                "requirement_text": "应按招标文件要求完成燃气管道施工、试验及验收。",
                "status": "pending_confirm",
                "evidence_bindings": [],
            }
        ],
    )

    assert "中压GB1燃气管道施工组织方案" in text
    assert "一、已知事实与编制边界" in text
    assert "计划工期：90日历天" in text
    assert "应按招标文件要求完成燃气管道施工、试验及验收。" in text
    assert "管材、连接方式、防腐、压力等级和分段长度以施工图、设计说明和工程量清单为准" in text
    assert "DN100" not in text
    assert "0.4MPa" not in text
    assert any(item["title"] == "中压GB1燃气管道施工组织方案" and item["as"] == "l4_narrative" for item in diag["rendered"])


def test_l4_submission_narrative_omits_review_boundary_words() -> None:
    chapters = [
        *CHAPTERS,
        {"section_type": "gas_completion_handover_service", "title": "缺陷责任期服务和竣工资料移交", "attachments": []},
    ]
    text, diag = _assemble(
        export_mode="submission",
        chapters=chapters,
        facts={**FACTS, "warranty_period": "24个月"},
    )

    assert "缺陷责任期内建立响应、排查、处置和回访闭环" in text
    assert "缺陷责任/保修期：24个月" in text
    assert "审阅稿" not in text
    assert "不得由系统补写" not in text
    assert any(item["as"] == "l4_narrative" for item in diag["rendered"])
