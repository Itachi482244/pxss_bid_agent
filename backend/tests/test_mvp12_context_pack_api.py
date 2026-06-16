from __future__ import annotations

import base64
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from uuid import uuid4
from uuid import UUID

import pytest
from alembic import command
from alembic.config import Config
from docx import Document as WordDocument
from fastapi.testclient import TestClient
from sqlalchemy import delete, select

from app.db.session import SessionLocal
from app.main import app
from app.models import ComplianceItem
from app.models import (
    BidSection,
    Document,
    DocumentChunk,
    DocumentVersion,
    EnterpriseMaterial,
    Project,
    ProjectMember,
    Tenant,
    User,
)
from scripts.seed_dev_data import seed


@pytest.fixture(scope="module", autouse=True)
def prepare_database() -> None:
    backend_dir = Path(__file__).resolve().parents[1]
    config = Config(str(backend_dir / "alembic.ini"))
    config.set_main_option("script_location", str(backend_dir / "migrations"))
    command.upgrade(config, "head")
    seed()


def get_seed_project_and_section(client: TestClient) -> tuple[str, str]:
    projects = client.get("/api/v1/projects").json()
    project = next(item for item in projects if item["name"] == "智慧园区弱电工程投标")
    sections = client.get(f"/api/v1/projects/{project['id']}/sections").json()
    return project["id"], sections[0]["id"]


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def _ensure_exportable_tender_doc(project_id: str, section_id: str) -> None:
    sample_text = """第五章 响应文件组成
供应商的响应文件应包含以下部分：
一、磋商响应声明
二、供应商的资格证明资料
三、技术/商务响应与偏离表
四、报价一览表
一、磋商响应声明
致 （采购人、采购代理机构）：
根据贵方为 （项目名称）的磋商邀请（采购代理编号： ），签字代表 （姓名、职务）经正式授权并代表供应商 （供应商名称）提交响应文件。
供应商名称（盖单位公章）：
日期： 年 月 日
二、供应商的资格证明资料
供应商应提供营业执照、信用查询等证明材料。
第五章 评分办法
1 技术方案完整、合理的得10分。
第四章 工程量清单
序号 | 项目名称 | 单位 | 数量 | 综合单价 | 合价
1 | DN100燃气管道安装 | 米 | 120.5
"""
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id)) if project else None
        assert project is not None and user is not None
        suffix = uuid4().hex
        doc = Document(
            tenant_id=project.tenant_id,
            project_id=UUID(project_id),
            section_id=UUID(section_id),
            doc_type="tender",
            title="格式装配测试招标文件",
            source_type="upload",
            original_filename="format-export-test.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_ext="docx",
            file_size=1024,
            sha256=f"{suffix:0<64}"[:64],
            bucket="test",
            object_key=f"tests/{suffix}/format-export-test.docx",
            status="available",
            created_by=user.id,
            acquired_at=datetime.now(UTC),
        )
        db.add(doc)
        db.flush()
        version = DocumentVersion(
            tenant_id=project.tenant_id,
            document_id=doc.id,
            version_no=1,
            version_label="format-export-test",
            object_key=f"tests/{suffix}/parsed.json",
            sha256=f"{suffix:1<64}"[:64],
            parse_status="succeeded",
            parser_name="test",
            parser_version="1.0",
            frozen_at=datetime.now(UTC),
            created_by=user.id,
            change_reason="格式装配导出测试",
        )
        db.add(version)
        db.flush()
        doc.current_version_id = version.id
        db.add(
            DocumentChunk(
                tenant_id=project.tenant_id,
                document_id=doc.id,
                document_version_id=version.id,
                section_id=UUID(section_id),
                chunk_index=1,
                page_no=1,
                heading_path="响应文件组成",
                content_text=sample_text,
                content_hash=f"{1:064x}",
                bbox_json=None,
                table_json=None,
            )
        )
        db.commit()


def _ensure_embeddable_business_license_material(project_id: str) -> str:
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id)) if project else None
        assert project is not None and user is not None
        suffix = uuid4().hex
        object_key = f"tests/{suffix}/business-license.png"
        material = EnterpriseMaterial(
            tenant_id=project.tenant_id,
            material_type="license",
            name=f"格式装配测试营业执照 {suffix[:8]}",
            certificate_no=f"LIC-{suffix[:8]}",
            data_level="internal",
            verification_status="confirmed",
            evidence_text="营业执照副本，可覆盖供应商主体资格证明。",
            file_name="business-license.png",
            content_type="image/png",
            file_size=len(PNG_1X1),
            sha256=f"{suffix:2<64}"[:64],
            bucket="test-bucket",
            object_key=object_key,
            created_by=user.id,
            updated_by=user.id,
        )
        db.add(material)
        db.commit()
        return object_key


def _ensure_embeddable_credit_material(project_id: str) -> str:
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id)) if project else None
        assert project is not None and user is not None
        suffix = uuid4().hex
        object_key = f"tests/{suffix}/credit-query.png"
        material = EnterpriseMaterial(
            tenant_id=project.tenant_id,
            material_type="commitment",
            name=f"格式装配测试信用中国查询截图 {suffix[:8]}",
            certificate_no=None,
            data_level="internal",
            verification_status="confirmed",
            evidence_text="信用中国和中国政府采购网查询截图，可覆盖信用查询要求。",
            file_name="credit-query.png",
            content_type="image/png",
            file_size=len(PNG_1X1),
            sha256=f"{suffix:4<64}"[:64],
            bucket="test-bucket",
            object_key=object_key,
            created_by=user.id,
            updated_by=user.id,
        )
        db.add(material)
        db.commit()
        return object_key


def _ensure_irrelevant_product_catalog_material(project_id: str) -> str:
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        user = db.scalar(select(User))
        assert project is not None and user is not None
        suffix = uuid4().hex
        object_key = f"tests/{suffix}/product-catalog.png"
        material = EnterpriseMaterial(
            tenant_id=project.tenant_id,
            material_type="product_catalog",
            name=f"格式装配测试产品图册 {suffix[:8]}",
            certificate_no=None,
            data_level="internal",
            verification_status="confirmed",
            evidence_text="产品图册材料，与供应商资格证明资料无关。",
            file_name="product-catalog.png",
            content_type="image/png",
            file_size=len(PNG_1X1),
            sha256=f"{suffix:3<64}"[:64],
            bucket="test-bucket",
            object_key=object_key,
            created_by=user.id,
            updated_by=user.id,
        )
        db.add(material)
        db.commit()
        return object_key


def _add_isolated_quality_project() -> tuple[dict[str, str], str, str]:
    """Create a tenant with no enterprise materials so quality-gate blocking is deterministic."""
    now = datetime.now(UTC)
    suffix = uuid4().hex[:8]
    tenant_code = f"quality-{suffix}"
    user_external_id = f"quality-user-{suffix}"
    sample_text = """第五章 响应文件组成
供应商的响应文件应包含以下部分：
一、磋商响应声明
二、供应商的资格证明资料
三、技术/商务响应与偏离表
二、供应商的资格证明资料
附件2-1 施工总承包特级许可证书
第五章 评分办法
1 技术方案完整、合理的得10分。
"""
    with SessionLocal() as db:
        tenant = Tenant(code=tenant_code, name=f"质量体检租户 {suffix}", status="active")
        db.add(tenant)
        db.flush()
        user = User(
            tenant_id=tenant.id,
            external_id=user_external_id,
            name="质量体检测试用户",
            email=f"{suffix}@example.com",
            status="active",
        )
        db.add(user)
        db.flush()
        project = Project(
            tenant_id=tenant.id,
            name=f"质量体检阻断项目 {suffix}",
            purchaser="测试采购人",
            agency="测试代理",
            budget_amount=None,
            region_code="CN-TEST",
            industry_code="quality-test",
            status="pending_confirm",
            created_by=user.id,
        )
        db.add(project)
        db.flush()
        section = BidSection(
            tenant_id=tenant.id,
            project_id=project.id,
            code=f"quality-{suffix}",
            name="质量体检测试标段",
            budget_amount=None,
            status="pending_confirm",
            created_by=user.id,
        )
        db.add(section)
        db.flush()
        db.add(
            ProjectMember(
                tenant_id=tenant.id,
                project_id=project.id,
                user_id=user.id,
                role_code="owner",
                status="active",
                created_by=user.id,
            )
        )
        doc = Document(
            tenant_id=tenant.id,
            project_id=project.id,
            section_id=section.id,
            doc_type="tender",
            title="质量体检测试招标文件",
            source_type="upload",
            original_filename="quality-gate.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_ext="docx",
            file_size=1024,
            sha256=f"{suffix:5<64}"[:64],
            bucket="test",
            object_key=f"tests/{suffix}/quality-gate.docx",
            status="available",
            created_by=user.id,
            acquired_at=now,
        )
        db.add(doc)
        db.flush()
        version = DocumentVersion(
            tenant_id=tenant.id,
            document_id=doc.id,
            version_no=1,
            version_label="quality-gate",
            object_key=f"tests/{suffix}/quality-gate.json",
            sha256=f"{suffix:6<64}"[:64],
            parse_status="succeeded",
            parser_name="test",
            parser_version="1.0",
            frozen_at=now,
            created_by=user.id,
            change_reason="质量体检阻断测试",
        )
        db.add(version)
        db.flush()
        doc.current_version_id = version.id
        db.add(
            DocumentChunk(
                tenant_id=tenant.id,
                document_id=doc.id,
                document_version_id=version.id,
                section_id=section.id,
                chunk_index=1,
                page_no=1,
                heading_path="响应文件组成",
                content_text=sample_text,
                content_hash=f"{7:064x}",
                bbox_json=None,
                table_json=None,
            )
        )
        db.commit()
        return (
            {"X-Tenant-Code": tenant_code, "X-User-External-Id": user_external_id},
            str(project.id),
            str(section.id),
        )


def _add_isolated_format_export_project() -> tuple[dict[str, str], str, str]:
    """Create a tenant/project with clean draft state for format-export gate tests."""
    now = datetime.now(UTC)
    suffix = uuid4().hex[:8]
    tenant_code = f"format-{suffix}"
    user_external_id = f"format-user-{suffix}"
    sample_text = """第五章 响应文件组成
供应商的响应文件应包含以下部分：
一、磋商响应声明
二、供应商的资格证明资料
三、技术/商务响应与偏离表
四、报价一览表
一、磋商响应声明
致 （采购人、采购代理机构）：
根据贵方为 （项目名称）的磋商邀请（采购代理编号： ），签字代表 （姓名、职务）经正式授权并代表供应商 （供应商名称）提交响应文件。
供应商名称（盖单位公章）：
日期： 年 月 日
二、供应商的资格证明资料
供应商应提供营业执照、信用查询等证明材料。
第五章 评分办法
1 技术方案完整、合理的得10分。
第四章 工程量清单
序号 | 项目名称 | 单位 | 数量 | 综合单价 | 合价
1 | DN100燃气管道安装 | 米 | 120.5
"""
    with SessionLocal() as db:
        tenant = Tenant(code=tenant_code, name=f"格式装配租户 {suffix}", status="active")
        db.add(tenant)
        db.flush()
        user = User(
            tenant_id=tenant.id,
            external_id=user_external_id,
            name="格式装配测试用户",
            email=f"format-{suffix}@example.com",
            status="active",
        )
        db.add(user)
        db.flush()
        project = Project(
            tenant_id=tenant.id,
            name=f"格式装配项目 {suffix}",
            purchaser="测试采购人",
            agency="测试代理",
            budget_amount="1000000",
            region_code="CN-TEST",
            industry_code="format-test",
            status="pending_confirm",
            created_by=user.id,
        )
        db.add(project)
        db.flush()
        section = BidSection(
            tenant_id=tenant.id,
            project_id=project.id,
            code=f"format-{suffix}",
            name="格式装配测试标段",
            budget_amount="1000000",
            status="pending_confirm",
            created_by=user.id,
        )
        db.add(section)
        db.flush()
        db.add(
            ProjectMember(
                tenant_id=tenant.id,
                project_id=project.id,
                user_id=user.id,
                role_code="owner",
                status="active",
                created_by=user.id,
            )
        )
        doc = Document(
            tenant_id=tenant.id,
            project_id=project.id,
            section_id=section.id,
            doc_type="tender",
            title="格式装配测试招标文件",
            source_type="upload",
            original_filename="format-export-test.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_ext="docx",
            file_size=1024,
            sha256=f"{suffix:8<64}"[:64],
            bucket="test",
            object_key=f"tests/{suffix}/format-export-test.docx",
            status="available",
            created_by=user.id,
            acquired_at=now,
        )
        db.add(doc)
        db.flush()
        version = DocumentVersion(
            tenant_id=tenant.id,
            document_id=doc.id,
            version_no=1,
            version_label="format-export-test",
            object_key=f"tests/{suffix}/parsed.json",
            sha256=f"{suffix:9<64}"[:64],
            parse_status="succeeded",
            parser_name="test",
            parser_version="1.0",
            frozen_at=now,
            created_by=user.id,
            change_reason="格式装配导出测试",
        )
        db.add(version)
        db.flush()
        doc.current_version_id = version.id
        chunk = DocumentChunk(
            tenant_id=tenant.id,
            document_id=doc.id,
            document_version_id=version.id,
            section_id=section.id,
            chunk_index=1,
            page_no=1,
            heading_path="响应文件组成",
            content_text=sample_text,
            content_hash=f"{8:064x}",
            bbox_json=None,
            table_json=None,
        )
        db.add(chunk)
        db.flush()
        db.add(
            ComplianceItem(
                tenant_id=tenant.id,
                project_id=project.id,
                section_id=section.id,
                source_document_id=doc.id,
                source_version_id=version.id,
                source_chunk_id=chunk.id,
                source_page_no=chunk.page_no,
                item_type="scoring",
                requirement_text="技术方案完整、合理的得10分。",
                normalized_requirement=f"format_scoring:{suffix}",
                response_suggestion="在评分索引中定位技术方案响应。",
                evidence_text="第五章 评分办法：技术方案完整、合理的得10分。",
                status="pending_confirm",
                risk_level="medium",
                is_mandatory=False,
                is_batch_confirm_allowed=True,
                created_by=user.id,
            )
        )
        db.commit()
        return (
            {"X-Tenant-Code": tenant_code, "X-User-External-Id": user_external_id},
            str(project.id),
            str(section.id),
        )


def test_mvp12_context_pack_preview_confirm_and_supersede_without_draft_blocks() -> None:
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    base = f"/api/v1/projects/{project_id}/sections/{section_id}/business-draft"
    payload = {
        "profile_id": "engineering_construction_business_v1",
        "section_types": [
            "bid_letter",
            "bid_commitment",
            "qualification_performance_summary",
            "priced_boq",
        ],
    }

    decision_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/qualification-decision/generate"
    )
    assert decision_response.status_code == 200
    decision = decision_response.json()
    confirm_decision_response = client.post(
        (
            f"/api/v1/projects/{project_id}/sections/{section_id}/"
            f"qualification-decision/{decision['id']}/confirm"
        ),
        json={"reason": "测试确认参标建议，允许进入 ContextPack"},
    )
    assert confirm_decision_response.status_code == 200
    assert confirm_decision_response.json()["status"] == "confirmed"

    with SessionLocal() as db:
        items = list(
            db.scalars(
                select(ComplianceItem).where(
                    ComplianceItem.project_id == UUID(project_id),
                    ComplianceItem.section_id == UUID(section_id),
                    ComplianceItem.deleted_at.is_(None),
                )
            ).all()
        )
        for item in items:
            item.status = "confirmed"
            explanation = dict(item.explanation_json or {})
            explanation["enterprise_evidence_not_required"] = True
            explanation["enterprise_evidence_not_required_reason"] = "MVP1.2 ContextPack API 测试已人工核验"
            item.explanation_json = explanation
        db.commit()

    blocks_before_response = client.get(f"{base}/blocks")
    assert blocks_before_response.status_code == 200
    block_ids_before = {item["id"] for item in blocks_before_response.json()}

    preview_response = client.post(f"{base}/context-pack/preview", json=payload)
    assert preview_response.status_code == 200
    preview = preview_response.json()
    assert preview["profile_id"] == "engineering_construction_business_v1"
    assert preview["schema_version"] == "1.0"
    assert preview["readiness_status"] in {"pass", "warn", "block"}
    assert preview["context_json"]["matrix_summary"]["total"] > 0
    assert preview["context_json"]["qualification_decision"]["status"] == "confirmed"
    assert isinstance(preview["context_json"]["bound_evidence"], list)
    assert isinstance(preview["context_json"]["missing_facts"], list)
    policy = preview["context_json"]["content_quality_policy"]
    assert policy["mode"] == "evidence_first_bid_writing"
    assert policy["domain"] in {
        "engineering_general",
        "engineering_epc",
        "municipal_gas_pipeline",
        "municipal_gas_pipeline_epc",
    }
    constraints = preview["context_json"]["generation_constraints"]
    assert "source_discipline" in constraints
    assert any("工程量清单" in claim for claim in constraints["forbidden_claims"])
    readiness_checks = preview["readiness_json"]["checks"]
    assert all(check.get("summary") and check.get("action") for check in readiness_checks)
    assert not [
        check
        for check in readiness_checks
        if check["status"] == "block" and check["code"] != "qualification.no_go_confirmed"
    ]
    section_types = [item["section_type"] for item in preview["outline_plan_json"]["sections"]]
    assert set(section_types) == set(payload["section_types"])

    create_response = client.post(f"{base}/context-pack", json=payload)
    assert create_response.status_code == 200
    context_pack = create_response.json()
    assert context_pack["status"] == "confirmed"
    assert context_pack["readiness_status"] == preview["readiness_status"]
    assert len(context_pack["section_context_packs"]) == len(payload["section_types"])
    assert context_pack["outline_plan_json"]["sections"][0]["section_type"] == "bid_letter"
    assert all(section_pack["context_json"]["section"] for section_pack in context_pack["section_context_packs"])
    assert all(
        section_pack["context_json"]["content_quality_policy"]["mode"] == "evidence_first_bid_writing"
        for section_pack in context_pack["section_context_packs"]
    )

    blocks_after_first_context_pack = client.get(f"{base}/blocks")
    assert blocks_after_first_context_pack.status_code == 200
    assert {item["id"] for item in blocks_after_first_context_pack.json()} == block_ids_before

    second_payload = {
        "profile_id": "engineering_construction_business_v1",
        "section_types": ["bid_letter", "bid_commitment"],
    }
    second_create_response = client.post(f"{base}/context-pack", json=second_payload)
    assert second_create_response.status_code == 200
    second_context_pack = second_create_response.json()
    assert second_context_pack["id"] != context_pack["id"]
    assert len(second_context_pack["section_context_packs"]) == len(second_payload["section_types"])

    list_response = client.get(f"{base}/context-pack")
    assert list_response.status_code == 200
    active_context_packs = list_response.json()
    assert active_context_packs
    assert active_context_packs[0]["id"] == second_context_pack["id"]
    assert context_pack["id"] not in {item["id"] for item in active_context_packs}

    blocks_after_second_context_pack = client.get(f"{base}/blocks")
    assert blocks_after_second_context_pack.status_code == 200
    assert {item["id"] for item in blocks_after_second_context_pack.json()} == block_ids_before


def test_tender_format_docx_export_supports_review_and_submission_modes(monkeypatch) -> None:
    written: list[dict] = []
    fetched: list[dict] = []

    def fake_put_object_bytes(**kwargs):
        written.append(kwargs)

    def fake_get_object_bytes(**kwargs):
        fetched.append(kwargs)
        return PNG_1X1

    monkeypatch.setattr("app.services.tender_format_export.put_object_bytes", fake_put_object_bytes)
    monkeypatch.setattr("app.services.tender_format_export.get_object_bytes", fake_get_object_bytes)
    client = TestClient(app)
    headers, project_id, section_id = _add_isolated_format_export_project()
    license_object_key = _ensure_embeddable_business_license_material(project_id)
    credit_object_key = _ensure_embeddable_credit_material(project_id)
    irrelevant_object_key = _ensure_irrelevant_product_catalog_material(project_id)
    base = f"/api/v1/projects/{project_id}/sections/{section_id}/business-draft/format-docx/export"

    summary_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/section-quality-summary",
        headers=headers,
    )
    assert summary_response.status_code == 200, summary_response.text
    summary = summary_response.json()
    assert summary["status"] == "warn"
    assert summary["export_preview"]["submission_allowed"] is True
    assert summary["export_preview"]["scoring_index_count"] >= 1
    assert summary["material_summary"]["embeddable_count"] >= 2

    review = client.post(base, json={"export_mode": "review"}, headers=headers)
    assert review.status_code == 200, review.text
    review_body = review.json()
    assert review_body["export_type"] == "tender_format_docx"
    assert review_body["source_snapshot_json"]["export_mode"] == "review"
    assert review_body["source_snapshot_json"]["assembler_diag"]["review_checklist_included"] is True
    pricing_diag = review_body["source_snapshot_json"]["assembler_diag"]["pricing"]
    assert pricing_diag["row_count"] >= 1
    assert pricing_diag["budget_amount"] is not None
    assert pricing_diag["budget_status"] == "pending_prices"
    assert any(item["object_key"] == license_object_key for item in fetched)
    assert any(item["object_key"] == credit_object_key for item in fetched)
    assert not any(item["object_key"] == irrelevant_object_key for item in fetched)
    review_doc = WordDocument(BytesIO(written[0]["data"]))
    assert len(review_doc.inline_shapes) >= 1
    assert review_body["source_snapshot_json"]["assembler_diag"]["embedded_materials"]
    assert review_body["source_snapshot_json"]["material_image_diag"]["material_image_candidate_count"] >= 1

    submission = client.post(base, json={"export_mode": "submission"}, headers=headers)
    assert submission.status_code == 200, submission.text
    submission_body = submission.json()
    assert submission_body["source_snapshot_json"]["export_mode"] == "submission"
    assert submission_body["source_snapshot_json"]["assembler_diag"]["review_checklist_included"] is False

    assert len(written) == 2
    assert all(item["data"].startswith(b"PK") for item in written)
    assert not any(item["object_key"] == irrelevant_object_key for item in fetched)


def test_section_quality_summary_blocks_submission_when_core_material_missing() -> None:
    client = TestClient(app)
    headers, project_id, section_id = _add_isolated_quality_project()
    base = f"/api/v1/projects/{project_id}/sections/{section_id}"

    summary_response = client.get(f"{base}/section-quality-summary", headers=headers)
    assert summary_response.status_code == 200, summary_response.text
    summary = summary_response.json()
    assert summary["status"] == "block"
    assert summary["status_label"] == "阻断"
    assert summary["export_preview"]["submission_allowed"] is False
    assert any(check["code"] == "coverage.disqualifying_gaps" for check in summary["checks"])

    submission = client.post(
        f"{base}/business-draft/format-docx/export",
        json={"export_mode": "submission"},
        headers=headers,
    )
    assert submission.status_code == 409
    assert "正式版导出已拦截" in submission.text


def test_section_quality_summary_marks_submission_unavailable_without_tender_text() -> None:
    client = TestClient(app)
    headers, project_id, section_id = _add_isolated_quality_project()
    with SessionLocal() as db:
        db.execute(delete(DocumentChunk).where(DocumentChunk.section_id == UUID(section_id)))
        db.commit()

    response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/section-quality-summary",
        headers=headers,
    )
    assert response.status_code == 200, response.text
    summary = response.json()
    assert summary["status"] == "block"
    assert summary["export_preview"]["review_allowed"] is True
    assert summary["export_preview"]["submission_allowed"] is False
    assert "缺少可体检的招标正文" in summary["export_preview"]["submission_blocked_reason"]
