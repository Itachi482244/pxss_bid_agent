from __future__ import annotations

import hashlib
import json
from io import BytesIO
from decimal import Decimal
from pathlib import Path
from types import SimpleNamespace
from datetime import UTC, datetime
from uuid import UUID, uuid4

import pytest
import fitz
from alembic import command
from alembic.config import Config
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from fastapi.testclient import TestClient
from openpyxl import load_workbook
from sqlalchemy import func, select
from sqlalchemy.exc import IntegrityError

from app.core.config import settings
from app.db.session import SessionLocal
from app.main import app
from app.models import (
    AgentReviewItem,
    AsyncTask,
    AuditLog,
    BidSection,
    ComplianceEvidenceBinding,
    ComplianceItem,
    Document,
    DocumentChunk,
    DocumentVersion,
    EnterpriseMaterial,
    ExportFile,
    Project,
    ProjectMember,
    QualificationDecision,
    QualificationEvaluation,
    SectionConfirmation,
    User,
)
from app.services import agent_assist as agent_assist_service
from app.services.document_conversion import LegacyDocConversionError
from app.services.file_acquisition import DownloadedFile
from scripts.seed_dev_data import seed


@pytest.fixture(scope="module", autouse=True)
def prepare_database() -> None:
    backend_dir = Path(__file__).resolve().parents[1]
    config = Config(str(backend_dir / "alembic.ini"))
    config.set_main_option("script_location", str(backend_dir / "migrations"))
    command.upgrade(config, "head")
    seed()


def get_seed_project_and_section(client: TestClient) -> tuple[str, str]:
    projects = client.get("/api/v1/projects", params={"limit": 200}).json()
    project = next(item for item in projects if item["name"] == "智慧园区弱电工程投标")
    sections = client.get(f"/api/v1/projects/{project['id']}/sections").json()
    return project["id"], sections[0]["id"]


def get_cleanroom_project_and_section(client: TestClient) -> tuple[str, str]:
    projects = client.get("/api/v1/projects", params={"limit": 200}).json()
    project = next(item for item in projects if item["name"] == "洁净车间净化设备采购与安装项目")
    sections = client.get(f"/api/v1/projects/{project['id']}/sections").json()
    return project["id"], sections[0]["id"]


def create_agent_assist_fixture(
    project_id: str,
    section_id: str,
    *,
    token: str,
    status: str = "needs_material",
    risk_level: str = "medium",
    is_mandatory: bool = True,
) -> tuple[str, str]:
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        document = db.scalar(
            select(Document).where(
                Document.project_id == project.id,
                Document.section_id == section.id,
                Document.current_version_id.is_not(None),
            )
        )
        assert document is not None
        version = db.get(DocumentVersion, document.current_version_id)
        assert version is not None
        chunk = db.scalar(
            select(DocumentChunk)
            .where(DocumentChunk.document_version_id == version.id)
            .order_by(DocumentChunk.chunk_index.asc())
        )
        assert chunk is not None
        material = EnterpriseMaterial(
            tenant_id=project.tenant_id,
            material_type="license",
            name=f"Agent测试营业执照-{token}",
            data_level="internal",
            verification_status="confirmed",
            evidence_text=f"Agent测试营业执照 {token} 可证明主体资格，供 agent 证据建议测试使用。",
            file_name=f"agent-license-{token}.pdf",
            created_by=user.id,
            updated_by=user.id,
        )
        db.add(material)
        db.flush()
        item = ComplianceItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            source_document_id=document.id,
            source_version_id=version.id,
            source_chunk_id=chunk.id,
            source_page_no=chunk.page_no,
            item_type="qualification",
            requirement_text=f"投标人须提供 Agent测试营业执照 {token}，并保证真实有效。",
            normalized_requirement=f"agent_assist_license_{token}",
            response_suggestion=f"绑定 Agent测试营业执照-{token} 作为资格证明。",
            evidence_text=chunk.content_text,
            status=status,
            risk_level=risk_level,
            is_mandatory=is_mandatory,
            is_batch_confirm_allowed=False,
            confidence_score=Decimal("0.9300"),
            created_by=user.id,
        )
        db.add(item)
        db.commit()
        return str(item.id), str(material.id)


def create_confirmable_section_fixture(*, assist_stage: str = "awaiting_confirm") -> tuple[str, str]:
    with SessionLocal() as db:
        user = db.scalar(select(User).where(User.status == "active"))
        assert user is not None
        now = datetime.now(UTC)
        project = Project(
            tenant_id=user.tenant_id,
            name=f"Agent单次确认测试项目-{uuid4().hex[:8]}",
            purchaser="测试采购人",
            agency="测试代理",
            status="draft",
            bid_deadline_at=now,
            created_by=user.id,
        )
        db.add(project)
        db.flush()
        section = BidSection(
            tenant_id=user.tenant_id,
            project_id=project.id,
            code="A",
            name="测试标段",
            status="draft",
            assist_stage=assist_stage,
            bid_deadline_at=now,
            created_by=user.id,
        )
        db.add(section)
        db.flush()
        member = ProjectMember(
            tenant_id=user.tenant_id,
            project_id=project.id,
            user_id=user.id,
            role_code="owner",
            status="active",
            created_by=user.id,
        )
        db.add(member)
        decision = QualificationDecision(
            tenant_id=user.tenant_id,
            project_id=project.id,
            section_id=section.id,
            recommendation="go",
            status="confirmed",
            summary="测试参标建议已确认。",
            satisfied_count=0,
            blocking_count=0,
            missing_count=0,
            pending_count=0,
            reasons_json={},
            created_by=user.id,
            confirmed_by=user.id,
            confirmed_at=now,
            confirm_reason="测试确认",
        )
        db.add(decision)
        db.commit()
        return str(project.id), str(section.id)


def test_project_read_api_returns_seeded_workspace() -> None:
    client = TestClient(app)

    projects_response = client.get("/api/v1/projects", params={"limit": 200})
    assert projects_response.status_code == 200
    projects = projects_response.json()
    assert len(projects) >= 1
    project = next(item for item in projects if item["name"] == "智慧园区弱电工程投标")
    assert project["section_count"] == 1
    assert project["compliance_item_count"] >= 1
    assert project["pending_confirm_count"] >= 1

    project_id = project["id"]
    detail_response = client.get(f"/api/v1/projects/{project_id}")
    assert detail_response.status_code == 200
    detail = detail_response.json()
    assert detail["notice_url"] == "https://example.com/tender/demo"

    sections_response = client.get(f"/api/v1/projects/{project_id}/sections")
    assert sections_response.status_code == 200
    sections = sections_response.json()
    assert len(sections) == 1
    section = sections[0]
    assert section["name"] == "一标段：综合布线与安防"
    assert section["document_count"] >= 1

    items_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section['id']}/compliance-items",
        params={"status": "pending_confirm", "risk_level": "medium"},
    )
    assert items_response.status_code == 200
    items = items_response.json()
    seed_item = next(item for item in items if item["source_page_no"] == 12)
    assert seed_item["source_version_label"] == "v0.1"
    assert seed_item["owner_name"] == "演示管理员"
    assert seed_item["priority_rank"] in {0, 1, 2, 3}
    assert seed_item["priority_label"].startswith("P")

    audit_response = client.get(
        f"/api/v1/projects/{project_id}/audit-logs",
        params={"action": "seed.dev_data_created"},
    )
    assert audit_response.status_code == 200
    audit_logs = audit_response.json()
    assert len(audit_logs) >= 1
    assert any(item["action"] == "seed.dev_data_created" for item in audit_logs)


def test_cleanroom_demo_sample_is_seeded_for_mvp1_hardening() -> None:
    client = TestClient(app)

    projects_response = client.get("/api/v1/projects", params={"limit": 200})
    assert projects_response.status_code == 200
    project = next(
        item for item in projects_response.json() if item["name"] == "洁净车间净化设备采购与安装项目"
    )
    sections_response = client.get(f"/api/v1/projects/{project['id']}/sections")
    assert sections_response.status_code == 200
    section = sections_response.json()[0]

    items_response = client.get(f"/api/v1/projects/{project['id']}/sections/{section['id']}/compliance-items")
    assert items_response.status_code == 200
    item_types = {item["item_type"] for item in items_response.json()}
    assert {"qualification", "mandatory_response", "technical_response", "scoring", "deadline"}.issubset(item_types)

    preflight_response = client.get(f"/api/v1/projects/{project['id']}/sections/{section['id']}/preflight-check")
    assert preflight_response.status_code == 200
    preflight = preflight_response.json()
    assert preflight["status"] == "block"
    assert preflight["mandatory_missing_evidence_count"] >= 1
    assert preflight["technical_pending_count"] >= 1

    search_response = client.get(
        "/api/v1/enterprise/materials/search",
        params={"query": "高效过滤器 检测报告 过滤效率 洁净等级", "limit": 50},
    )
    assert search_response.status_code == 200
    results = search_response.json()
    assert any(item["material_type"] == "test_report" and item["recommend_reason"] for item in results)


def test_agent_assist_creates_exception_review_items() -> None:
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_id, section_id = get_cleanroom_project_and_section(client)
    with SessionLocal() as db:
        active_decisions = db.scalars(
            select(QualificationDecision).where(
                QualificationDecision.project_id == UUID(project_id),
                QualificationDecision.section_id == UUID(section_id),
                QualificationDecision.status != "superseded",
            )
        ).all()
        for decision in active_decisions:
            decision.status = "superseded"
        db.commit()

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={"async_processing": False},
    )

    assert response.status_code == 202
    task = response.json()
    assert task["task_type"] == "agent_assist"
    assert task["status"] == "succeeded"
    assert task["output_json"]["open_count"] >= 1
    assert task["output_json"]["matrix_review_count"] >= 1
    assert task["output_json"]["qualification_technical_count"] >= 1

    items_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items",
        params={"status": "open", "run_key": task["output_json"]["run_key"]},
    )
    assert items_response.status_code == 200
    review_items = items_response.json()
    actions = {item["action"] for item in review_items}
    assert "confirm_matrix_item" in actions
    assert "confirm_qualification_decision" in actions
    assert "review_technical_response" in actions
    assert any(action in actions for action in {"accept_evidence_binding", "missing_evidence"})
    assert all(item["requires_human"] is True for item in review_items)

    summary_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/summary",
        params={"run_key": task["output_json"]["run_key"]},
    )
    assert summary_response.status_code == 200
    summary = summary_response.json()
    assert summary["open_count"] == task["output_json"]["open_count"]
    assert summary["suggested_actions"]

    with SessionLocal() as db:
        logs = db.scalars(
            select(AuditLog).where(
                AuditLog.project_id == UUID(project_id),
                AuditLog.section_id == UUID(section_id),
                AuditLog.action == "agent.assist_finished",
            )
        ).all()
        assert logs
        assert all(log.actor_type == "agent" for log in logs)


def test_agent_assist_preserves_confirmed_qualification_decision_on_rerun() -> None:
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_id, section_id = get_cleanroom_project_and_section(client)

    with SessionLocal() as db:
        active_decisions = db.scalars(
            select(QualificationDecision).where(
                QualificationDecision.project_id == UUID(project_id),
                QualificationDecision.section_id == UUID(section_id),
                QualificationDecision.status != "superseded",
            )
        ).all()
        for decision in active_decisions:
            decision.status = "superseded"
        db.commit()

    first_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={"async_processing": False},
    )
    assert first_response.status_code == 202
    first_run_key = first_response.json()["output_json"]["run_key"]
    review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items",
        params={"status": "open", "step": "qualification_technical", "run_key": first_run_key, "limit": 500},
    )
    assert review_response.status_code == 200
    decision_item = next(
        item for item in review_response.json() if item["action"] == "confirm_qualification_decision"
    )
    decision_id = decision_item["qualification_decision_id"]

    accept_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{decision_item['id']}/accept",
        json={"reason": "测试人工确认 Go/No-Go 结论"},
    )
    assert accept_response.status_code == 200

    second_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={"async_processing": False},
    )
    assert second_response.status_code == 202
    second_run_key = second_response.json()["output_json"]["run_key"]

    with SessionLocal() as db:
        decision = db.get(QualificationDecision, UUID(decision_id))
        assert decision is not None
        assert decision.status == "confirmed"
        assert decision.confirm_reason == "测试人工确认 Go/No-Go 结论"
        open_decision_item = db.scalar(
            select(AgentReviewItem).where(
                AgentReviewItem.run_key == second_run_key,
                AgentReviewItem.action == "confirm_qualification_decision",
                AgentReviewItem.status == "open",
            )
        )
        assert open_decision_item is None
        preserved_item = db.scalar(
            select(AgentReviewItem).where(
                AgentReviewItem.run_key == second_run_key,
                AgentReviewItem.action == "qualification_decision_preserved",
                AgentReviewItem.status == "auto_passed",
                AgentReviewItem.qualification_decision_id == UUID(decision_id),
            )
        )
        assert preserved_item is not None


def test_agent_assist_preserves_confirmed_qualification_evaluation_on_rerun() -> None:
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_id, section_id = get_cleanroom_project_and_section(client)

    run_response = client.post(f"/api/v1/projects/{project_id}/sections/{section_id}/qualification-evaluations/run")
    assert run_response.status_code == 200
    evaluation = run_response.json()[0]

    confirm_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/qualification-evaluations/{evaluation['id']}/confirm",
        json={"reason": "测试人工确认资格评估项"},
    )
    assert confirm_response.status_code == 200
    assert confirm_response.json()["confirmed_by"] is not None

    assist_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={"async_processing": False},
    )
    assert assist_response.status_code == 202

    evaluations_response = client.get(f"/api/v1/projects/{project_id}/sections/{section_id}/qualification-evaluations")
    assert evaluations_response.status_code == 200
    preserved = next(item for item in evaluations_response.json() if item["id"] == evaluation["id"])
    assert preserved["confirmed_by"] is not None
    assert preserved["confirmed_at"] is not None
    assert preserved["confirm_reason"] == "测试人工确认资格评估项"
    run_key = assist_response.json()["output_json"]["run_key"]
    with SessionLocal() as db:
        open_review = db.scalar(
            select(AgentReviewItem).where(
                AgentReviewItem.run_key == run_key,
                AgentReviewItem.action == "review_qualification_evaluation",
                AgentReviewItem.qualification_evaluation_id == UUID(evaluation["id"]),
                AgentReviewItem.status == "open",
            )
        )
        assert open_review is None
        preserved_review = db.scalar(
            select(AgentReviewItem).where(
                AgentReviewItem.run_key == run_key,
                AgentReviewItem.action == "qualification_evaluation_preserved",
                AgentReviewItem.qualification_evaluation_id == UUID(evaluation["id"]),
                AgentReviewItem.status == "auto_passed",
            )
        )
        assert preserved_review is not None


def test_agent_assist_marks_changed_confirmed_qualification_evaluation() -> None:
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_id, section_id = get_cleanroom_project_and_section(client)

    run_response = client.post(f"/api/v1/projects/{project_id}/sections/{section_id}/qualification-evaluations/run")
    assert run_response.status_code == 200
    evaluation = run_response.json()[0]
    confirm_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/qualification-evaluations/{evaluation['id']}/confirm",
        json={"reason": "测试确认后结果变化保留确认"},
    )
    assert confirm_response.status_code == 200
    confirmed_by = confirm_response.json()["confirmed_by"]
    confirmed_at = confirm_response.json()["confirmed_at"]

    with SessionLocal() as db:
        legacy_evaluation = db.get(QualificationEvaluation, UUID(evaluation["id"]))
        assert legacy_evaluation is not None
        assert legacy_evaluation.confirmed_snapshot_json is not None
        legacy_evaluation.confirmed_snapshot_json = None
        item = db.get(ComplianceItem, UUID(evaluation["compliance_item_id"]))
        assert item is not None
        item.requirement_text = f"{item.requirement_text} 补充测试变更 {uuid4().hex[:6]}"
        db.commit()

    assist_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={"async_processing": False},
    )
    assert assist_response.status_code == 202
    run_key = assist_response.json()["output_json"]["run_key"]

    with SessionLocal() as db:
        updated = db.get(QualificationEvaluation, UUID(evaluation["id"]))
        assert updated is not None
        assert str(updated.confirmed_by) == confirmed_by
        assert updated.confirmed_at is not None
        assert updated.confirmed_at.isoformat().replace("+00:00", "Z") == confirmed_at
        assert updated.confirm_reason == "测试确认后结果变化保留确认"
        changed_review = db.scalar(
            select(AgentReviewItem).where(
                AgentReviewItem.run_key == run_key,
                AgentReviewItem.action == "review_qualification_evaluation",
                AgentReviewItem.qualification_evaluation_id == UUID(evaluation["id"]),
                AgentReviewItem.status == "open",
                AgentReviewItem.conclusion_changed.is_(True),
            )
        )
        assert changed_review is not None
        assert changed_review.tier == "blocking"

    second_assist_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={"async_processing": False},
    )
    assert second_assist_response.status_code == 202
    second_run_key = second_assist_response.json()["output_json"]["run_key"]

    with SessionLocal() as db:
        repeated_changed_review = db.scalar(
            select(AgentReviewItem).where(
                AgentReviewItem.run_key == second_run_key,
                AgentReviewItem.action == "review_qualification_evaluation",
                AgentReviewItem.qualification_evaluation_id == UUID(evaluation["id"]),
                AgentReviewItem.status == "open",
                AgentReviewItem.conclusion_changed.is_(True),
            )
        )
        assert repeated_changed_review is not None
        updated = db.get(QualificationEvaluation, UUID(evaluation["id"]))
        assert updated is not None
        assert updated.confirmed_snapshot_json is not None
        assert updated.confirmed_snapshot_json["requirement_text"] != updated.requirement_text
        preserved_review = db.scalar(
            select(AgentReviewItem).where(
                AgentReviewItem.run_key == second_run_key,
                AgentReviewItem.action == "qualification_evaluation_preserved",
                AgentReviewItem.qualification_evaluation_id == UUID(evaluation["id"]),
            )
        )
        assert preserved_review is None


def test_agent_assist_reuses_inflight_task_for_same_section() -> None:
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)

    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        assert project is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        active_tasks = db.scalars(
            select(AsyncTask).where(
                AsyncTask.tenant_id == project.tenant_id,
                AsyncTask.project_id == UUID(project_id),
                AsyncTask.section_id == UUID(section_id),
                AsyncTask.task_type == "agent_assist",
                AsyncTask.status.in_(["pending", "running", "retrying"]),
            )
        ).all()
        for task in active_tasks:
            task.status = "canceled"
        task = AsyncTask(
            tenant_id=project.tenant_id,
            project_id=UUID(project_id),
            section_id=UUID(section_id),
            task_type="agent_assist",
            status="running",
            idempotency_key=f"agent-assist-inflight-test:{uuid4().hex}",
            progress=35,
            input_json={"run_key": f"agent-assist-inflight-test:{uuid4().hex[:8]}"},
            retry_count=0,
            max_retries=0,
            created_by=user.id,
        )
        db.add(task)
        db.commit()
        task_id = str(task.id)

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={"async_processing": False},
    )

    assert response.status_code == 202
    payload = response.json()
    assert payload["id"] == task_id
    assert payload["status"] == "running"
    with SessionLocal() as db:
        active_count = db.scalar(
            select(func.count(AsyncTask.id)).where(
                AsyncTask.project_id == UUID(project_id),
                AsyncTask.section_id == UUID(section_id),
                AsyncTask.task_type == "agent_assist",
                AsyncTask.status.in_(["pending", "running", "retrying"]),
            )
        )
        assert active_count == 1
        active_task = db.get(AsyncTask, UUID(task_id))
        assert active_task is not None
        active_task.status = "canceled"
        db.commit()


def test_agent_assist_active_task_unique_index_blocks_duplicate_inflight() -> None:
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)

    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        assert project is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        active_tasks = db.scalars(
            select(AsyncTask).where(
                AsyncTask.tenant_id == project.tenant_id,
                AsyncTask.project_id == UUID(project_id),
                AsyncTask.section_id == UUID(section_id),
                AsyncTask.task_type == "agent_assist",
                AsyncTask.status.in_(["pending", "running", "retrying"]),
            )
        ).all()
        for task in active_tasks:
            task.status = "canceled"
        first = AsyncTask(
            tenant_id=project.tenant_id,
            project_id=UUID(project_id),
            section_id=UUID(section_id),
            task_type="agent_assist",
            status="pending",
            idempotency_key=f"agent-assist-unique-test:{uuid4().hex}",
            progress=0,
            input_json={"run_key": f"agent-assist-unique-test:{uuid4().hex[:8]}"},
            retry_count=0,
            max_retries=0,
            created_by=user.id,
        )
        db.add(first)
        db.commit()
        first_id = first.id

    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        assert project is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        duplicate = AsyncTask(
            tenant_id=project.tenant_id,
            project_id=UUID(project_id),
            section_id=UUID(section_id),
            task_type="agent_assist",
            status="running",
            idempotency_key=f"agent-assist-unique-test:{uuid4().hex}",
            progress=10,
            input_json={"run_key": f"agent-assist-unique-test:{uuid4().hex[:8]}"},
            retry_count=0,
            max_retries=0,
            created_by=user.id,
        )
        db.add(duplicate)
        with pytest.raises(IntegrityError):
            db.commit()
        db.rollback()

    with SessionLocal() as db:
        first = db.get(AsyncTask, first_id)
        assert first is not None
        first.status = "canceled"
        db.commit()


def test_agent_assist_failure_supersedes_partial_run_items(monkeypatch: pytest.MonkeyPatch) -> None:
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    token = uuid4().hex[:8]
    item_id, _material_id = create_agent_assist_fixture(project_id, section_id, token=token)
    run_key = f"agent-assist-failure-test:{uuid4().hex[:8]}"

    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        assert project is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        previous_item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=UUID(project_id),
            section_id=UUID(section_id),
            run_key=f"agent-assist-previous-test:{uuid4().hex[:8]}",
            step="matrix_review",
            action="confirm_matrix_item",
            status="open",
            severity="high",
            title="失败回滚前的待办",
            detail="用于验证失败不清空旧待办",
            object_type="compliance_item",
            object_id=UUID(item_id),
            compliance_item_id=UUID(item_id),
            confidence_score=Decimal("0.5000"),
            requires_human=True,
            escalation_reasons=["测试旧待办"],
            triggered_by=user.id,
        )
        db.add(previous_item)
        task = AsyncTask(
            tenant_id=project.tenant_id,
            project_id=UUID(project_id),
            section_id=UUID(section_id),
            task_type="agent_assist",
            status="pending",
            idempotency_key=f"agent-assist-failure-test:{uuid4().hex}",
            progress=0,
            input_json={
                "run_key": run_key,
                "force": True,
                "scope": "section",
                "steps": ["matrix_review", "evidence_binding", "qualification_technical"],
            },
            retry_count=0,
            max_retries=0,
            created_by=user.id,
        )
        db.add(task)
        db.commit()
        task_id = task.id
        previous_item_id = previous_item.id

    def fail_material_search(*_args: object, **_kwargs: object) -> None:
        raise RuntimeError("material search exploded")

    monkeypatch.setattr(agent_assist_service, "search_material_hits", fail_material_search)
    with SessionLocal() as db:
        result = agent_assist_service.execute_agent_assist_task(db, task_id)

    assert result["status"] == "failed"
    assert result["error_code"] == "AGENT_ASSIST_FAILED"
    with SessionLocal() as db:
        task = db.get(AsyncTask, task_id)
        assert task is not None
        assert task.status == "failed"
        previous_item = db.get(AgentReviewItem, previous_item_id)
        assert previous_item is not None
        assert previous_item.status == "open"
        current_run_items = db.scalars(
            select(AgentReviewItem).where(
                AgentReviewItem.async_task_id == task_id,
                AgentReviewItem.run_key == run_key,
            )
        ).all()
        assert current_run_items == []


def test_agent_review_accept_evidence_suggestion_binds_material() -> None:
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    token = uuid4().hex[:8]
    item_id, material_id = create_agent_assist_fixture(project_id, section_id, token=token)
    with SessionLocal() as db:
        active_decisions = db.scalars(
            select(QualificationDecision).where(
                QualificationDecision.project_id == UUID(project_id),
                QualificationDecision.section_id == UUID(section_id),
                QualificationDecision.status != "superseded",
            )
        ).all()
        for decision in active_decisions:
            decision.status = "superseded"
        db.commit()

    assist_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={"async_processing": False},
    )
    assert assist_response.status_code == 202
    run_key = assist_response.json()["output_json"]["run_key"]
    with SessionLocal() as db:
        active_decision = db.scalar(
            select(QualificationDecision).where(
                QualificationDecision.project_id == UUID(project_id),
                QualificationDecision.section_id == UUID(section_id),
                QualificationDecision.status != "superseded",
            )
        )
        assert active_decision is not None
        active_decision_id = active_decision.id

    review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items",
        params={"status": "open", "step": "evidence_binding", "run_key": run_key, "limit": 500},
    )
    assert review_response.status_code == 200
    suggestion = next(
        item
        for item in review_response.json()
        if item["action"] == "accept_evidence_binding"
        and item["compliance_item_id"] == item_id
        and item["enterprise_material_id"] == material_id
    )

    accept_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{suggestion['id']}/accept",
        json={"reason": "测试采纳 agent 证据建议"},
    )

    assert accept_response.status_code == 200
    accepted = accept_response.json()
    assert accepted["status"] == "accepted"
    assert accepted["decided_by"]
    with SessionLocal() as db:
        binding = db.scalar(
            select(ComplianceEvidenceBinding).where(
                ComplianceEvidenceBinding.compliance_item_id == UUID(item_id),
                ComplianceEvidenceBinding.enterprise_material_id == UUID(material_id),
                ComplianceEvidenceBinding.status == "active",
            )
        )
        assert binding is not None
        assert "测试采纳 agent 证据建议" in binding.bind_reason
        item = db.get(ComplianceItem, UUID(item_id))
        assert item is not None
        assert item.status == "pending_confirm"
        decision = db.get(QualificationDecision, active_decision_id)
        assert decision is not None
        assert decision.status == "superseded"
        evaluation = db.scalar(
            select(QualificationEvaluation).where(
                QualificationEvaluation.compliance_item_id == UUID(item_id),
                QualificationEvaluation.project_id == UUID(project_id),
                QualificationEvaluation.section_id == UUID(section_id),
            )
        )
        assert evaluation is not None
        assert evaluation.evaluation_status != "needs_material"
        assert evaluation.matched_material_id == UUID(material_id)
        audit = db.scalar(
            select(AuditLog).where(
                AuditLog.object_id == binding.id,
                AuditLog.action == "agent.evidence_suggestion_accepted",
            )
        )
        assert audit is not None
        assert audit.actor_type == "user"
        assert audit.after_json is not None
        assert audit.after_json["qualification_refresh"]["invalidated_decision_count"] >= 1


def test_agent_review_accept_evidence_suggestion_closes_when_material_already_bound() -> None:
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    token = uuid4().hex[:8]
    item_id, material_id = create_agent_assist_fixture(project_id, section_id, token=token)

    assist_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={"async_processing": False},
    )
    assert assist_response.status_code == 202
    run_key = assist_response.json()["output_json"]["run_key"]

    review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items",
        params={"status": "open", "step": "evidence_binding", "run_key": run_key, "limit": 500},
    )
    assert review_response.status_code == 200
    suggestion = next(
        item
        for item in review_response.json()
        if item["action"] == "accept_evidence_binding"
        and item["compliance_item_id"] == item_id
        and item["enterprise_material_id"] == material_id
    )

    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        material = db.get(EnterpriseMaterial, UUID(material_id))
        assert project is not None
        assert material is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        binding = ComplianceEvidenceBinding(
            tenant_id=project.tenant_id,
            project_id=UUID(project_id),
            section_id=UUID(section_id),
            compliance_item_id=UUID(item_id),
            enterprise_material_id=UUID(material_id),
            evidence_text=material.evidence_text or material.name,
            material_snapshot=agent_assist_service.enterprise_material_snapshot(material),
            confidence_score=Decimal("0.9300"),
            bind_reason="测试人工已提前绑定等价资料",
            status="active",
            created_by=user.id,
        )
        db.add(binding)
        db.commit()
        binding_id = binding.id

    accept_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{suggestion['id']}/accept",
        json={"reason": "测试采纳已满足的证据建议"},
    )

    assert accept_response.status_code == 200, accept_response.text
    accepted = accept_response.json()
    assert accepted["status"] == "accepted"
    with SessionLocal() as db:
        binding_count = db.scalar(
            select(func.count(ComplianceEvidenceBinding.id)).where(
                ComplianceEvidenceBinding.compliance_item_id == UUID(item_id),
                ComplianceEvidenceBinding.status == "active",
            )
        )
        review_item = db.get(AgentReviewItem, UUID(suggestion["id"]))
        audit = db.scalar(
            select(AuditLog).where(
                AuditLog.object_id == UUID(suggestion["id"]),
                AuditLog.action == "agent.evidence_suggestion_already_bound",
            )
        )
        assert binding_count == 1
        assert db.get(ComplianceEvidenceBinding, binding_id) is not None
        assert review_item is not None
        assert review_item.status == "accepted"
        assert audit is not None


def test_agent_assist_summary_suggested_actions_only_count_open_items() -> None:
    project_id = uuid4()
    section_id = uuid4()
    run_key = f"agent-summary-test:{uuid4().hex[:8]}"
    item = AgentReviewItem(
        tenant_id=uuid4(),
        project_id=project_id,
        section_id=section_id,
        run_key=run_key,
        step="matrix_review",
        action="confirm_matrix_item",
        status="accepted",
        severity="high",
        title="已处理条款",
        detail="已处理",
        object_type="compliance_item",
        object_id=uuid4(),
        confidence_score=Decimal("0.5000"),
        requires_human=True,
        triggered_by=uuid4(),
    )
    summary = agent_assist_service.agent_assist_summary_from_items(
        project_id=project_id,
        section_id=section_id,
        run_key=run_key,
        task_id=None,
        items=[item],
    )

    assert summary["open_count"] == 0
    assert summary["suggested_actions"] == ["当前没有需要人工拍板的 Agent 例外项"]


def test_agent_assist_summary_counts_llm_readonly_advice() -> None:
    project_id = uuid4()
    section_id = uuid4()
    run_key = f"agent-summary-test:{uuid4().hex[:8]}"
    item = AgentReviewItem(
        tenant_id=uuid4(),
        project_id=project_id,
        section_id=section_id,
        run_key=run_key,
        step="qualification_technical",
        action="ack_llm_technical_advice",
        status="open",
        severity="medium",
        title="LLM 只读建议",
        detail="建议人工查看",
        object_type="compliance_item",
        object_id=uuid4(),
        confidence_score=Decimal("0.7000"),
        requires_human=True,
        triggered_by=uuid4(),
    )

    summary = agent_assist_service.agent_assist_summary_from_items(
        project_id=project_id,
        section_id=section_id,
        run_key=run_key,
        task_id=None,
        items=[item],
    )

    assert summary["open_count"] == 1
    assert summary["llm_advice_count"] == 1
    assert summary["technical_review_count"] == 0
    assert summary["suggested_actions"] == ["查看 1 条 LLM 只读建议"]


def test_agent_review_accept_matrix_item_requires_source_verification() -> None:
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    token = uuid4().hex[:8]
    item_id, _material_id = create_agent_assist_fixture(
        project_id,
        section_id,
        token=token,
        status="pending_confirm",
        risk_level="high",
        is_mandatory=True,
    )

    assist_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={"async_processing": False},
    )
    assert assist_response.status_code == 202
    run_key = assist_response.json()["output_json"]["run_key"]

    review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items",
        params={"status": "open", "step": "matrix_review", "run_key": run_key, "limit": 500},
    )
    assert review_response.status_code == 200
    matrix_item = next(
        item
        for item in review_response.json()
        if item["action"] == "confirm_matrix_item" and item["compliance_item_id"] == item_id
    )

    blocked_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{matrix_item['id']}/accept",
        json={"reason": "测试不核验来源时阻断"},
    )
    assert blocked_response.status_code == 409
    assert "必须核验来源" in blocked_response.json()["detail"]

    accepted_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{matrix_item['id']}/accept",
        json={"reason": "测试人工核验来源后确认", "source_verified": True},
    )
    assert accepted_response.status_code == 200
    assert accepted_response.json()["status"] == "accepted"
    with SessionLocal() as db:
        item = db.get(ComplianceItem, UUID(item_id))
        assert item is not None
        assert item.status == "confirmed"
        review_item = db.get(AgentReviewItem, UUID(matrix_item["id"]))
        assert review_item is not None
        assert review_item.status == "accepted"


def test_section_final_review_and_confirm_lock_success() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()

    final_review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/final-review"
    )
    assert final_review_response.status_code == 200
    final_review = final_review_response.json()
    assert final_review["assist_stage"] == "awaiting_confirm"
    assert final_review["red"]["open_count"] == 0
    assert final_review["can_confirm"] is True

    confirm_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/confirm-lock",
        json={"reason": "测试最终确认锁定"},
    )
    assert confirm_response.status_code == 200, confirm_response.text
    confirmed = confirm_response.json()
    assert confirmed["assist_stage"] == "confirmed"
    assert confirmed["can_generate"] is True

    with SessionLocal() as db:
        section = db.get(BidSection, UUID(section_id))
        assert section is not None
        assert section.assist_stage == "confirmed"
        confirmation = db.scalar(
            select(SectionConfirmation).where(
                SectionConfirmation.project_id == UUID(project_id),
                SectionConfirmation.section_id == UUID(section_id),
                SectionConfirmation.status == "active",
            )
        )
        assert confirmation is not None
        assert confirmation.snapshot_json["assist_stage"] == "confirmed"


def test_tender_format_export_requires_confirmed_assist_stage() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture(assist_stage="awaiting_confirm")

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/business-draft/format-docx/export",
        json={"export_mode": "review"},
    )

    assert response.status_code == 409
    assert "尚不可生成标书" in response.text
    with SessionLocal() as db:
        section = db.get(BidSection, UUID(section_id))
        assert section is not None
        assert section.assist_stage == "awaiting_confirm"


def test_agent_assist_requires_unlock_for_confirmed_or_generated_section() -> None:
    client = TestClient(app)
    for assist_stage in ("confirmed", "generated"):
        project_id, section_id = create_confirmable_section_fixture(assist_stage=assist_stage)

        response = client.post(
            f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
            json={"async_processing": False, "force": True},
        )

        assert response.status_code == 409
        assert "撤回确认" in response.text
        with SessionLocal() as db:
            section = db.get(BidSection, UUID(section_id))
            assert section is not None
            assert section.assist_stage == assist_stage


@pytest.mark.parametrize("assist_stage", ["confirmed", "generated"])
def test_locked_section_blocks_legacy_edit_endpoints(assist_stage: str) -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture(assist_stage=assist_stage)
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        compliance_item = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标人应提供锁定后不可静默修改的服务承诺。",
        )
        material = EnterpriseMaterial(
            tenant_id=project.tenant_id,
            material_type="qualification",
            name="锁定围栏测试资料",
            data_level="internal",
            verification_status="confirmed",
            evidence_text="该资料用于验证锁定后的证据绑定被拒绝。",
            created_by=user.id,
            updated_by=user.id,
        )
        review_item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-lock-guard:{uuid4().hex[:8]}",
            step="matrix_review",
            action="confirm_matrix_item",
            status="open",
            severity="medium",
            tier="blocking",
            title="锁定围栏待办",
            detail="锁定后不应被旧入口处理。",
            object_type="compliance_item",
            object_id=compliance_item.id,
            compliance_item_id=compliance_item.id,
            confidence_score=Decimal("0.5000"),
            requires_human=True,
            triggered_by=user.id,
        )
        db.add_all([material, review_item])
        db.commit()
        item_id = compliance_item.id
        material_id = material.id
        review_item_id = review_item.id

    section_response = client.patch(
        f"/api/v1/projects/{project_id}/sections/{section_id}",
        json={"name": "锁定后非法修改", "reason": "测试锁定围栏"},
    )
    item_response = client.patch(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/{item_id}",
        json={"requirement_text": "锁定后非法修改条款", "reason": "测试锁定围栏"},
    )
    evidence_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/{item_id}/evidence-bindings",
        json={"enterprise_material_id": str(material_id), "reason": "测试锁定围栏"},
    )
    qualification_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/qualification-evaluations/run"
    )
    review_item_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{review_item_id}/accept",
        json={"reason": "测试锁定围栏", "source_verified": True},
    )

    responses = [
        section_response,
        item_response,
        evidence_response,
        qualification_response,
        review_item_response,
    ]
    for response in responses:
        assert response.status_code == 409, response.text
        assert "撤回确认" in response.text

    with SessionLocal() as db:
        section = db.get(BidSection, UUID(section_id))
        compliance_item = db.get(ComplianceItem, item_id)
        review_item = db.get(AgentReviewItem, review_item_id)
        binding_count = db.scalar(
            select(func.count(ComplianceEvidenceBinding.id)).where(
                ComplianceEvidenceBinding.compliance_item_id == item_id,
                ComplianceEvidenceBinding.status == "active",
            )
        )
        assert section is not None
        assert compliance_item is not None
        assert review_item is not None
        assert section.assist_stage == assist_stage
        assert section.name == "测试标段"
        assert compliance_item.requirement_text == "投标人应提供锁定后不可静默修改的服务承诺。"
        assert binding_count == 0
        assert review_item.status == "open"


def test_agent_assist_request_steps_are_exposed_and_ordered() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture(assist_stage="not_started")

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/assist",
        json={
            "async_processing": False,
            "force": True,
            "steps": ["matrix_review", "evidence_binding"],
        },
    )

    assert response.status_code == 202, response.text
    task = response.json()
    assert task["input_json"]["steps"] == ["evidence_binding", "matrix_review"]
    with SessionLocal() as db:
        section = db.get(BidSection, UUID(section_id))
        assert section is not None
        assert section.assist_stage == "awaiting_confirm"


def test_confirm_lock_blocks_when_red_zone_has_open_item() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        assert project is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=UUID(project_id),
            section_id=UUID(section_id),
            run_key=f"agent-final-review-test:{uuid4().hex[:8]}",
            step="matrix_review",
            action="confirm_matrix_item",
            status="open",
            severity="high",
            tier="blocking",
            is_disqualifying=True,
            title="红牌测试项",
            detail="必须先处理",
            object_type="compliance_item",
            object_id=None,
            confidence_score=Decimal("0.5000"),
            requires_human=True,
            escalation_reasons=["测试红牌"],
            triggered_by=user.id,
        )
        db.add(item)
        db.commit()

    final_review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/final-review"
    )
    assert final_review_response.status_code == 200
    assert final_review_response.json()["red"]["open_count"] == 1

    confirm_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/confirm-lock",
        json={"reason": "测试红牌拦截"},
    )
    assert confirm_response.status_code == 409
    assert "阻塞" in confirm_response.text


def test_final_review_counts_all_effective_items_when_item_list_is_limited() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        assert project is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        db.add_all(
            AgentReviewItem(
                tenant_id=project.tenant_id,
                project_id=UUID(project_id),
                section_id=UUID(section_id),
                run_key=f"agent-final-review-limit:{uuid4().hex[:8]}",
                step="matrix_review",
                action="agent_matrix_low_risk_pass",
                status="auto_passed",
                severity="low",
                tier="silent",
                is_disqualifying=False,
                auto_applied=True,
                title=f"白区自动通过测试项 {index}",
                detail="用于验证 final review 计数不受明细列表上限影响",
                object_type="compliance_item",
                object_id=None,
                confidence_score=Decimal("0.9000"),
                requires_human=False,
                escalation_reasons=[],
                triggered_by=user.id,
            )
            for index in range(501)
        )
        db.commit()

    final_review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/final-review"
    )

    assert final_review_response.status_code == 200
    final_review = final_review_response.json()
    assert final_review["white"]["total_count"] == 501
    assert final_review["white"]["auto_passed_count"] == 501
    assert len(final_review["white"]["items"]) == 500


def test_project_sections_overview_returns_assist_stage_and_readiness_counts() -> None:
    client = TestClient(app)
    project_id, first_section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        assert project is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        second_section = BidSection(
            tenant_id=project.tenant_id,
            project_id=project.id,
            code="B",
            name="第二标段",
            status="draft",
            assist_stage="awaiting_confirm",
            bid_deadline_at=project.bid_deadline_at,
            created_by=user.id,
        )
        db.add(second_section)
        db.flush()
        db.add(
            QualificationDecision(
                tenant_id=project.tenant_id,
                project_id=project.id,
                section_id=second_section.id,
                recommendation="go",
                status="confirmed",
                summary="第二标段参标建议已确认。",
                satisfied_count=0,
                blocking_count=0,
                missing_count=0,
                pending_count=0,
                reasons_json={},
                created_by=user.id,
                confirmed_by=user.id,
                confirmed_at=datetime.now(UTC),
                confirm_reason="测试确认",
            )
        )
        db.add(
            AgentReviewItem(
                tenant_id=project.tenant_id,
                project_id=project.id,
                section_id=second_section.id,
                run_key=f"agent-overview-red:{uuid4().hex[:8]}",
                step="matrix_review",
                action="confirm_matrix_item",
                status="open",
                severity="high",
                tier="blocking",
                title="第二标段红牌",
                detail="用于项目级概览统计。",
                object_type="compliance_item",
                object_id=None,
                confidence_score=Decimal("0.5000"),
                requires_human=True,
                triggered_by=user.id,
            )
        )
        db.commit()
        second_section_id = second_section.id

    response = client.get(f"/api/v1/projects/{project_id}/sections-overview")
    assert response.status_code == 200, response.text
    overview = response.json()
    assert overview["project_id"] == project_id
    assert overview["total_count"] >= 2
    assert overview["awaiting_confirm_count"] >= 2
    assert overview["red_open_count"] >= 1
    first_item = next(item for item in overview["sections"] if item["id"] == first_section_id)
    second_item = next(item for item in overview["sections"] if item["id"] == str(second_section_id))
    assert first_item["assist_stage"] == "awaiting_confirm"
    assert first_item["can_confirm"] is True
    assert second_item["red_open_count"] == 1
    assert second_item["suggested_action"] == "去确认(1)"


def test_final_review_keeps_unrerun_step_items_after_partial_rerun() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        assert project is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        old_run_key = f"agent-effective-old:{uuid4().hex[:8]}"
        new_run_key = f"agent-effective-new:{uuid4().hex[:8]}"
        db.add_all(
            [
                AgentReviewItem(
                    tenant_id=project.tenant_id,
                    project_id=UUID(project_id),
                    section_id=UUID(section_id),
                    run_key=old_run_key,
                    step="evidence_binding",
                    action="missing_evidence",
                    status="open",
                    severity="high",
                    tier="blocking",
                    title="旧证据红牌仍有效",
                    detail="局部重跑矩阵时证据步骤未重跑，仍应阻塞确认。",
                    object_type="compliance_item",
                    object_id=None,
                    confidence_score=Decimal("0.0000"),
                    requires_human=True,
                    escalation_reasons=["缺少企业资料证据"],
                    triggered_by=user.id,
                ),
                AgentReviewItem(
                    tenant_id=project.tenant_id,
                    project_id=UUID(project_id),
                    section_id=UUID(section_id),
                    run_key=new_run_key,
                    step="matrix_review",
                    action="agent_matrix_low_risk_pass",
                    status="auto_passed",
                    severity="low",
                    tier="silent",
                    title="新矩阵自动完成",
                    detail="局部重跑矩阵产生的新结果。",
                    object_type="compliance_item",
                    object_id=None,
                    confidence_score=Decimal("0.9500"),
                    requires_human=False,
                    triggered_by=user.id,
                    auto_applied=True,
                ),
            ]
        )
        db.commit()

    final_review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/final-review"
    )
    assert final_review_response.status_code == 200
    final_review = final_review_response.json()
    assert final_review["red"]["open_count"] == 1
    assert final_review["white"]["auto_passed_count"] == 1
    assert final_review["can_confirm"] is False

    summary_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/summary"
    )
    assert summary_response.status_code == 200
    summary = summary_response.json()
    assert summary["open_count"] == 1
    assert summary["auto_passed_count"] == 1

    confirm_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/confirm-lock",
        json={"reason": "测试局部重跑后旧步骤仍阻塞"},
    )
    assert confirm_response.status_code == 409
    assert "阻塞" in confirm_response.text


def _add_low_risk_source_item(
    db,
    *,
    project: Project,
    section: BidSection,
    user: User,
    requirement_text: str,
    status: str = "pending_confirm",
    item_type: str = "other",
    risk_level: str = "low",
    is_mandatory: bool = False,
) -> ComplianceItem:
    token = uuid4().hex
    document = Document(
        tenant_id=project.tenant_id,
        project_id=project.id,
        section_id=section.id,
        doc_type="tender",
        title="测试招标文件",
        source_type="upload",
        original_filename="test.pdf",
        file_size=1,
        sha256=hashlib.sha256(token.encode()).hexdigest(),
        bucket="test",
        object_key=f"test/{token}.pdf",
        status="available",
        created_by=user.id,
        acquired_at=datetime.now(UTC),
    )
    db.add(document)
    db.flush()
    version = DocumentVersion(
        tenant_id=project.tenant_id,
        document_id=document.id,
        version_no=1,
        version_label="v1",
        object_key=document.object_key,
        sha256=document.sha256,
        parse_status="succeeded",
        created_by=user.id,
    )
    db.add(version)
    db.flush()
    document.current_version_id = version.id
    chunk = DocumentChunk(
        tenant_id=project.tenant_id,
        document_id=document.id,
        document_version_id=version.id,
        section_id=section.id,
        chunk_index=1,
        page_no=1,
        content_text=requirement_text,
        content_hash=hashlib.sha256(f"chunk-{token}".encode()).hexdigest(),
    )
    db.add(chunk)
    db.flush()
    compliance_item = ComplianceItem(
        tenant_id=project.tenant_id,
        project_id=project.id,
        section_id=section.id,
        source_document_id=document.id,
        source_version_id=version.id,
        source_chunk_id=chunk.id,
        source_page_no=1,
        item_type=item_type,
        requirement_text=requirement_text,
        status=status,
        risk_level=risk_level,
        is_mandatory=is_mandatory,
        is_batch_confirm_allowed=risk_level != "high" and not is_mandatory and status != "needs_material",
        confidence_score=Decimal("0.9500"),
        created_by=user.id,
    )
    db.add(compliance_item)
    db.flush()
    return compliance_item


def test_final_review_keeps_other_items_after_single_item_rerun() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        item_one = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标文件应包含服务承诺。",
        )
        item_two = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标文件应包含售后承诺。",
        )
        old_run_key = f"agent-item-rerun-old:{uuid4().hex[:8]}"
        new_run_key = f"agent-item-rerun-new:{uuid4().hex[:8]}"
        db.add_all(
            [
                AgentReviewItem(
                    tenant_id=project.tenant_id,
                    project_id=project.id,
                    section_id=section.id,
                    run_key=old_run_key,
                    step="matrix_review",
                    action="confirm_matrix_item",
                    status="open",
                    severity="medium",
                    tier="blocking",
                    title="旧红牌 item one",
                    detail="该条稍后被单条重评替换。",
                    object_type="compliance_item",
                    object_id=item_one.id,
                    compliance_item_id=item_one.id,
                    confidence_score=Decimal("0.5000"),
                    requires_human=True,
                    triggered_by=user.id,
                ),
                AgentReviewItem(
                    tenant_id=project.tenant_id,
                    project_id=project.id,
                    section_id=section.id,
                    run_key=old_run_key,
                    step="matrix_review",
                    action="confirm_matrix_item",
                    status="open",
                    severity="medium",
                    tier="blocking",
                    title="旧红牌 item two",
                    detail="该条未重评，仍应保留。",
                    object_type="compliance_item",
                    object_id=item_two.id,
                    compliance_item_id=item_two.id,
                    confidence_score=Decimal("0.5000"),
                    requires_human=True,
                    triggered_by=user.id,
                ),
            ]
        )
        db.commit()
        db.add(
            AgentReviewItem(
                tenant_id=project.tenant_id,
                project_id=project.id,
                section_id=section.id,
                run_key=new_run_key,
                step="matrix_review",
                action="agent_matrix_low_risk_pass",
                status="auto_passed",
                severity="low",
                tier="silent",
                title="单条重评自动完成 item one",
                detail="只替换 item one 的旧红牌。",
                object_type="compliance_item",
                object_id=item_one.id,
                compliance_item_id=item_one.id,
                confidence_score=Decimal("0.9500"),
                requires_human=False,
                auto_applied=True,
                triggered_by=user.id,
            )
        )
        db.commit()

    final_review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/final-review"
    )
    assert final_review_response.status_code == 200
    final_review = final_review_response.json()
    assert final_review["red"]["open_count"] == 1
    open_red_titles = [
        item["title"] for item in final_review["red"]["items"] if item["status"] == "open"
    ]
    assert open_red_titles == ["旧红牌 item two"]
    assert final_review["white"]["auto_passed_count"] == 1


def test_agent_review_resolve_evidence_not_required_triggers_single_item_rerun() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        item_one = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标人应提供服务能力证书。",
            status="needs_material",
        )
        item_two = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标人应提供售后能力证书。",
            status="needs_material",
        )
        current_item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-resolve-old:{uuid4().hex[:8]}",
            step="evidence_binding",
            action="missing_evidence",
            status="open",
            severity="medium",
            tier="blocking",
            title="缺证据 item one",
            detail="可通过标记无需证据补救。",
            object_type="compliance_item",
            object_id=item_one.id,
            compliance_item_id=item_one.id,
            confidence_score=Decimal("0.0000"),
            requires_human=True,
            triggered_by=user.id,
        )
        other_item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-resolve-old:{uuid4().hex[:8]}",
            step="evidence_binding",
            action="missing_evidence",
            status="open",
            severity="medium",
            tier="blocking",
            title="缺证据 item two",
            detail="未处理，仍应保留。",
            object_type="compliance_item",
            object_id=item_two.id,
            compliance_item_id=item_two.id,
            confidence_score=Decimal("0.0000"),
            requires_human=True,
            triggered_by=user.id,
        )
        db.add_all([current_item, other_item])
        db.commit()
        current_item_id = current_item.id
        item_one_id = item_one.id

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{current_item_id}/resolve",
        json={
            "resolution": "evidence_not_required",
            "reason": "测试人工判定无需企业资料证据",
        },
    )
    assert response.status_code == 200, response.text
    assert response.json()["status"] == "accepted"

    with SessionLocal() as db:
        item_one = db.get(ComplianceItem, item_one_id)
        current_item = db.get(AgentReviewItem, current_item_id)
        assert item_one is not None
        assert current_item is not None
        assert item_one.status == "pending_confirm"
        assert item_one.explanation_json is not None
        assert item_one.explanation_json["enterprise_evidence_not_required"] is True
        assert current_item.status == "accepted"
        rerun_item = db.scalar(
            select(AgentReviewItem).where(
                AgentReviewItem.compliance_item_id == item_one_id,
                AgentReviewItem.action == "agent_matrix_low_risk_pass",
                AgentReviewItem.status == "auto_passed",
            )
        )
        assert rerun_item is not None

    final_review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/final-review"
    )
    assert final_review_response.status_code == 200
    final_review = final_review_response.json()
    assert final_review["red"]["open_count"] == 1
    open_red_titles = [
        item["title"] for item in final_review["red"]["items"] if item["status"] == "open"
    ]
    assert open_red_titles == ["缺证据 item two"]
    assert final_review["white"]["auto_passed_count"] >= 1


def test_agent_review_resolve_keeps_unrerun_technical_item_open() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        technical_item = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标人应提供设备技术响应方案。",
            status="needs_material",
            item_type="technical_response",
        )
        missing_evidence = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-tech-rerun-old:{uuid4().hex[:8]}",
            step="evidence_binding",
            action="missing_evidence",
            status="open",
            severity="medium",
            tier="blocking",
            title="技术项缺证据",
            detail="补救后只应替换证据/矩阵结果。",
            object_type="compliance_item",
            object_id=technical_item.id,
            compliance_item_id=technical_item.id,
            confidence_score=Decimal("0.0000"),
            requires_human=True,
            triggered_by=user.id,
        )
        technical_review = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-tech-rerun-old:{uuid4().hex[:8]}",
            step="qualification_technical",
            action="review_technical_response",
            status="open",
            severity="medium",
            tier="blocking",
            title="技术响应仍需人工确认",
            detail="该待办不在本轮证据/矩阵定向重评范围内。",
            object_type="compliance_item",
            object_id=technical_item.id,
            compliance_item_id=technical_item.id,
            confidence_score=Decimal("0.5000"),
            requires_human=True,
            triggered_by=user.id,
        )
        db.add_all([missing_evidence, technical_review])
        db.commit()
        missing_evidence_id = missing_evidence.id
        technical_review_id = technical_review.id

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{missing_evidence_id}/resolve",
        json={
            "resolution": "evidence_not_required",
            "reason": "测试技术项无需企业资料证据",
        },
    )
    assert response.status_code == 200, response.text

    with SessionLocal() as db:
        technical_review = db.get(AgentReviewItem, technical_review_id)
        assert technical_review is not None
        assert technical_review.status == "open"

    final_review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/final-review"
    )
    assert final_review_response.status_code == 200
    final_review = final_review_response.json()
    open_red_titles = [
        item["title"] for item in final_review["red"]["items"] if item["status"] == "open"
    ]
    assert "技术响应仍需人工确认" in open_red_titles


def test_qualification_resolution_rebuilds_decision_review_item() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        qualification_item = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标人须具备有效营业执照。",
            status="needs_material",
            item_type="qualification",
            risk_level="medium",
            is_mandatory=True,
        )
        old_decision = db.scalar(
            select(QualificationDecision).where(
                QualificationDecision.project_id == project.id,
                QualificationDecision.section_id == section.id,
                QualificationDecision.status == "confirmed",
            )
        )
        assert old_decision is not None
        old_decision_review = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-qualification-old:{uuid4().hex[:8]}",
            step="qualification_technical",
            action="confirm_qualification_decision",
            status="open",
            severity="high",
            tier="blocking",
            title="旧参标建议待确认",
            detail="资格证据补救后应失效。",
            object_type="qualification_decision",
            object_id=old_decision.id,
            qualification_decision_id=old_decision.id,
            confidence_score=Decimal("0.7000"),
            requires_human=True,
            triggered_by=user.id,
        )
        missing_evidence = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-qualification-old:{uuid4().hex[:8]}",
            step="evidence_binding",
            action="missing_evidence",
            status="open",
            severity="high",
            tier="blocking",
            title="资格项缺证据",
            detail="补救后应重建资格汇总待办。",
            object_type="compliance_item",
            object_id=qualification_item.id,
            compliance_item_id=qualification_item.id,
            confidence_score=Decimal("0.0000"),
            requires_human=True,
            triggered_by=user.id,
        )
        db.add_all([old_decision_review, missing_evidence])
        db.commit()
        old_decision_id = old_decision.id
        old_decision_review_id = old_decision_review.id
        missing_evidence_id = missing_evidence.id

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{missing_evidence_id}/resolve",
        json={
            "resolution": "evidence_not_required",
            "reason": "测试资格项无需企业资料证据",
        },
    )
    assert response.status_code == 200, response.text

    with SessionLocal() as db:
        old_decision = db.get(QualificationDecision, old_decision_id)
        old_decision_review = db.get(AgentReviewItem, old_decision_review_id)
        assert old_decision is not None
        assert old_decision_review is not None
        assert old_decision.status == "superseded"
        assert old_decision_review.status == "superseded"
        new_decision_review = db.scalar(
            select(AgentReviewItem).where(
                AgentReviewItem.project_id == UUID(project_id),
                AgentReviewItem.section_id == UUID(section_id),
                AgentReviewItem.action == "confirm_qualification_decision",
                AgentReviewItem.status == "open",
                AgentReviewItem.qualification_decision_id != old_decision_id,
            )
        )
        assert new_decision_review is not None


def test_agent_review_closes_superseded_qualification_decision_item() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        decision = QualificationDecision(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            recommendation="conditional_go",
            status="superseded",
            summary="已失效的参标建议。",
            satisfied_count=0,
            blocking_count=0,
            missing_count=1,
            pending_count=0,
            reasons_json={},
            created_by=user.id,
        )
        db.add(decision)
        db.flush()
        review_item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-stale-decision:{uuid4().hex[:8]}",
            step="qualification_technical",
            action="confirm_qualification_decision",
            status="open",
            severity="high",
            tier="blocking",
            title="已失效参标建议待确认",
            detail="服务层应拒绝确认已 superseded 的 decision。",
            object_type="qualification_decision",
            object_id=decision.id,
            qualification_decision_id=decision.id,
            confidence_score=Decimal("0.7000"),
            requires_human=True,
            triggered_by=user.id,
        )
        db.add(review_item)
        db.commit()
        review_item_id = review_item.id

    list_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items",
        params={"status": "open"},
    )
    assert list_response.status_code == 200
    assert str(review_item_id) not in {item["id"] for item in list_response.json()}

    final_review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/final-review"
    )
    assert final_review_response.status_code == 200
    red_titles = [
        item["title"] for item in final_review_response.json()["red"]["items"] if item["status"] == "open"
    ]
    assert "已失效参标建议待确认" not in red_titles

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{review_item_id}/accept",
        json={"reason": "测试关闭已失效参标建议待办"},
    )
    assert response.status_code == 200
    assert response.json()["status"] == "superseded"
    with SessionLocal() as db:
        review_item = db.get(AgentReviewItem, review_item_id)
        audit = db.scalar(
            select(AuditLog).where(
                AuditLog.object_id == review_item_id,
                AuditLog.action == "agent.qualification_decision_stale_closed",
            )
        )
        assert review_item is not None
        assert review_item.status == "superseded"
        assert audit is not None


def test_legacy_evidence_bind_refreshes_agent_missing_evidence_item() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        item = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标人应提供服务能力证明。",
            status="needs_material",
        )
        material = EnterpriseMaterial(
            tenant_id=project.tenant_id,
            material_type="qualification",
            name="服务能力证明",
            data_level="internal",
            verification_status="confirmed",
            evidence_text="已确认的服务能力证明。",
            file_name="service-proof.pdf",
            created_by=user.id,
            updated_by=user.id,
        )
        review_item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-legacy-bind:{uuid4().hex[:8]}",
            step="evidence_binding",
            action="missing_evidence",
            status="open",
            severity="medium",
            tier="blocking",
            title="旧入口待补证据",
            detail="旧证据 Tab 绑定后应关闭该红牌。",
            object_type="compliance_item",
            object_id=item.id,
            compliance_item_id=item.id,
            confidence_score=Decimal("0.0000"),
            requires_human=True,
            triggered_by=user.id,
        )
        db.add_all([material, review_item])
        db.commit()
        item_id = item.id
        material_id = material.id
        review_item_id = review_item.id

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/{item_id}/evidence-bindings",
        json={
            "enterprise_material_id": str(material_id),
            "evidence_text": "已确认的服务能力证明。",
            "confidence_score": "0.9300",
            "reason": "测试旧入口绑定证据",
        },
    )
    assert response.status_code == 201, response.text

    with SessionLocal() as db:
        review_item = db.get(AgentReviewItem, review_item_id)
        assert review_item is not None
        assert review_item.status == "superseded"
        refreshed_item = db.scalar(
            select(AgentReviewItem).where(
                AgentReviewItem.compliance_item_id == item_id,
                AgentReviewItem.run_key.like("agent-evidence-bind:%"),
                AgentReviewItem.status.in_(["open", "auto_passed"]),
            )
        )
        assert refreshed_item is not None

    final_review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/final-review"
    )
    assert final_review_response.status_code == 200
    open_red_titles = [
        item["title"] for item in final_review_response.json()["red"]["items"] if item["status"] == "open"
    ]
    assert "旧入口待补证据" not in open_red_titles


def test_missing_evidence_cannot_be_accepted_or_dismissed_without_resolution() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        compliance_item = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标人应提供服务能力证书。",
            status="needs_material",
        )
        review_item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-missing-evidence-guard:{uuid4().hex[:8]}",
            step="evidence_binding",
            action="missing_evidence",
            status="open",
            severity="medium",
            tier="blocking",
            title="缺少证据测试",
            detail="必须通过补救动作关闭。",
            object_type="compliance_item",
            object_id=compliance_item.id,
            compliance_item_id=compliance_item.id,
            confidence_score=Decimal("0.0000"),
            requires_human=True,
            triggered_by=user.id,
        )
        db.add(review_item)
        db.commit()
        review_item_id = review_item.id

    accept_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{review_item_id}/accept",
        json={"reason": "测试不能直接采纳"},
    )
    dismiss_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/agent-review-items/{review_item_id}/dismiss",
        json={"reason": "测试不能直接忽略"},
    )

    assert accept_response.status_code == 409
    assert dismiss_response.status_code == 409
    assert "补救" in accept_response.text
    with SessionLocal() as db:
        review_item = db.get(AgentReviewItem, review_item_id)
        assert review_item is not None
        assert review_item.status == "open"


def test_confirm_lock_applies_preaccepted_matrix_item() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        token = uuid4().hex
        document = Document(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            doc_type="tender",
            title="测试招标文件",
            source_type="upload",
            original_filename="test.pdf",
            file_size=1,
            sha256=hashlib.sha256(token.encode()).hexdigest(),
            bucket="test",
            object_key=f"test/{token}.pdf",
            status="available",
            created_by=user.id,
            acquired_at=datetime.now(UTC),
        )
        db.add(document)
        db.flush()
        version = DocumentVersion(
            tenant_id=project.tenant_id,
            document_id=document.id,
            version_no=1,
            version_label="v1",
            object_key=document.object_key,
            sha256=document.sha256,
            parse_status="succeeded",
            created_by=user.id,
        )
        db.add(version)
        db.flush()
        document.current_version_id = version.id
        chunk = DocumentChunk(
            tenant_id=project.tenant_id,
            document_id=document.id,
            document_version_id=version.id,
            section_id=section.id,
            chunk_index=1,
            page_no=1,
            content_text="投标文件应包含服务承诺。",
            content_hash=hashlib.sha256(f"chunk-{token}".encode()).hexdigest(),
        )
        db.add(chunk)
        db.flush()
        compliance_item = ComplianceItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            source_document_id=document.id,
            source_version_id=version.id,
            source_chunk_id=chunk.id,
            source_page_no=1,
            item_type="other",
            requirement_text="投标文件应包含服务承诺。",
            status="pending_confirm",
            risk_level="low",
            is_mandatory=False,
            confidence_score=Decimal("0.9500"),
            created_by=user.id,
        )
        db.add(compliance_item)
        db.flush()
        review_item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-preaccept-test:{uuid4().hex[:8]}",
            step="matrix_review",
            action="pre_accept_matrix_item",
            status="open",
            severity="low",
            tier="pre_accepted",
            title="预采纳条款测试",
            detail="确认锁定时生效",
            object_type="compliance_item",
            object_id=compliance_item.id,
            compliance_item_id=compliance_item.id,
            confidence_score=Decimal("0.9500"),
            requires_human=False,
            triggered_by=user.id,
        )
        db.add(review_item)
        db.commit()
        item_id = compliance_item.id
        review_item_id = review_item.id

    confirm_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/confirm-lock",
        json={"reason": "测试预采纳批量生效"},
    )
    assert confirm_response.status_code == 200, confirm_response.text
    with SessionLocal() as db:
        compliance_item = db.get(ComplianceItem, item_id)
        review_item = db.get(AgentReviewItem, review_item_id)
        assert compliance_item is not None
        assert review_item is not None
        assert compliance_item.status == "confirmed"
        assert review_item.status == "accepted"
        assert review_item.auto_applied is True


def test_preaccepted_evidence_for_needs_material_item_stays_yellow_and_confirms(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        compliance_item = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标人应提供服务能力证书。",
            status="needs_material",
        )
        best_material = EnterpriseMaterial(
            tenant_id=project.tenant_id,
            material_type="qualification",
            name="服务能力证书 A",
            data_level="internal",
            verification_status="confirmed",
            evidence_text="服务能力证书 A 已确认。",
            created_by=user.id,
            updated_by=user.id,
        )
        second_material = EnterpriseMaterial(
            tenant_id=project.tenant_id,
            material_type="qualification",
            name="服务能力证书 B",
            data_level="internal",
            verification_status="confirmed",
            evidence_text="服务能力证书 B 已确认。",
            created_by=user.id,
            updated_by=user.id,
        )
        db.add_all([best_material, second_material])
        db.flush()

        def fake_search_material_hits(*args, **kwargs):
            return [
                SimpleNamespace(
                    material=best_material,
                    chunk=None,
                    snippet="服务能力证书 A 已确认。",
                    confidence_score=0.94,
                    base_score=0.94,
                    rerank_score=None,
                    recommend_reason="测试明显最优候选",
                ),
                SimpleNamespace(
                    material=second_material,
                    chunk=None,
                    snippet="服务能力证书 B 已确认。",
                    confidence_score=0.72,
                    base_score=0.72,
                    rerank_score=None,
                    recommend_reason="测试次优候选",
                ),
            ]

        monkeypatch.setattr(agent_assist_service, "search_material_hits", fake_search_material_hits)
        preaccepted_evidence_item_ids: set[UUID] = set()
        evidence_counts: dict[UUID, int] = {}
        evidence_items = agent_assist_service._add_evidence_review_items(
            db,
            project=project,
            section=section,
            items=[compliance_item],
            evidence_counts=evidence_counts,
            run_key=f"agent-preaccepted-evidence:{uuid4().hex[:8]}",
            async_task_id=None,
            actor_user_id=user.id,
            preaccepted_evidence_item_ids=preaccepted_evidence_item_ids,
        )
        matrix_items = agent_assist_service._add_matrix_review_items(
            db,
            project=project,
            section=section,
            items=[compliance_item],
            evidence_counts=evidence_counts,
            run_key=f"agent-preaccepted-matrix:{uuid4().hex[:8]}",
            async_task_id=None,
            actor_user_id=user.id,
            preaccepted_evidence_item_ids=preaccepted_evidence_item_ids,
        )
        evidence_actions = [item.action for item in evidence_items]
        matrix_actions = [item.action for item in matrix_items]
        db.commit()
        item_id = compliance_item.id
        evidence_review_item_id = evidence_items[0].id
        matrix_review_item_id = matrix_items[0].id

    assert evidence_actions == ["pre_accept_evidence_binding"]
    assert matrix_actions == ["pre_accept_matrix_item"]

    final_review_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/final-review"
    )
    assert final_review_response.status_code == 200
    final_review = final_review_response.json()
    assert final_review["red"]["open_count"] == 0
    assert final_review["yellow"]["open_count"] == 2
    assert final_review["can_confirm"] is True

    confirm_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/confirm-lock",
        json={"reason": "测试预采纳证据和条款一起生效"},
    )
    assert confirm_response.status_code == 200, confirm_response.text
    with SessionLocal() as db:
        compliance_item = db.get(ComplianceItem, item_id)
        evidence_review_item = db.get(AgentReviewItem, evidence_review_item_id)
        matrix_review_item = db.get(AgentReviewItem, matrix_review_item_id)
        binding_count = db.scalar(
            select(func.count(ComplianceEvidenceBinding.id)).where(
                ComplianceEvidenceBinding.compliance_item_id == item_id,
                ComplianceEvidenceBinding.status == "active",
            )
        )
        assert compliance_item is not None
        assert evidence_review_item is not None
        assert matrix_review_item is not None
        assert compliance_item.status == "confirmed"
        assert evidence_review_item.status == "accepted"
        assert matrix_review_item.status == "accepted"
        assert binding_count == 1


def test_confirm_lock_closes_preaccepted_evidence_when_binding_already_active() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        compliance_item = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标人应提供服务能力证书。",
        )
        material = EnterpriseMaterial(
            tenant_id=project.tenant_id,
            material_type="qualification",
            name="服务能力证书",
            data_level="internal",
            verification_status="confirmed",
            evidence_text="服务能力证书已确认。",
            created_by=user.id,
            updated_by=user.id,
        )
        db.add(material)
        db.flush()
        binding = ComplianceEvidenceBinding(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            compliance_item_id=compliance_item.id,
            enterprise_material_id=material.id,
            evidence_text="服务能力证书已确认。",
            material_snapshot={"id": str(material.id), "name": material.name},
            confidence_score=Decimal("0.9300"),
            bind_reason="测试中模拟人工已提前绑定",
            status="active",
            created_by=user.id,
        )
        review_item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-preaccept-evidence-satisfied:{uuid4().hex[:8]}",
            step="evidence_binding",
            action="pre_accept_evidence_binding",
            status="open",
            severity="low",
            tier="pre_accepted",
            title="预采纳证据已满足测试",
            detail="确认锁定时发现证据已存在，应关闭黄区项。",
            object_type="enterprise_material",
            object_id=material.id,
            compliance_item_id=compliance_item.id,
            enterprise_material_id=material.id,
            confidence_score=Decimal("0.9300"),
            requires_human=False,
            recommendation_json={"evidence_text": "服务能力证书已确认。"},
            triggered_by=user.id,
        )
        db.add_all([binding, review_item])
        db.commit()
        item_id = compliance_item.id
        review_item_id = review_item.id

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/confirm-lock",
        json={"reason": "测试预采纳证据已由人工满足"},
    )
    assert response.status_code == 200, response.text
    with SessionLocal() as db:
        section = db.get(BidSection, UUID(section_id))
        review_item = db.get(AgentReviewItem, review_item_id)
        active_binding_count = db.scalar(
            select(func.count(ComplianceEvidenceBinding.id)).where(
                ComplianceEvidenceBinding.compliance_item_id == item_id,
                ComplianceEvidenceBinding.status == "active",
            )
        )
        assert section is not None
        assert review_item is not None
        assert section.assist_stage == "confirmed"
        assert review_item.status == "accepted"
        assert review_item.auto_applied is True
        assert active_binding_count == 1


def test_confirm_lock_returns_conflict_when_preaccepted_evidence_is_invalid() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture()
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        compliance_item = _add_low_risk_source_item(
            db,
            project=project,
            section=section,
            user=user,
            requirement_text="投标人应提供服务能力证书。",
        )
        material = EnterpriseMaterial(
            tenant_id=project.tenant_id,
            material_type="qualification",
            name="失效服务能力证书",
            data_level="internal",
            verification_status="draft",
            evidence_text="该资料尚未确认。",
            created_by=user.id,
            updated_by=user.id,
        )
        db.add(material)
        db.flush()
        review_item = AgentReviewItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            run_key=f"agent-preaccept-evidence-invalid:{uuid4().hex[:8]}",
            step="evidence_binding",
            action="pre_accept_evidence_binding",
            status="open",
            severity="low",
            tier="pre_accepted",
            title="预采纳证据失效测试",
            detail="资料状态变化后确认锁定应返回业务冲突。",
            object_type="enterprise_material",
            object_id=material.id,
            compliance_item_id=compliance_item.id,
            enterprise_material_id=material.id,
            confidence_score=Decimal("0.9300"),
            requires_human=False,
            recommendation_json={"evidence_text": "该资料尚未确认。"},
            triggered_by=user.id,
        )
        db.add(review_item)
        db.commit()
        review_item_id = review_item.id

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/confirm-lock",
        json={"reason": "测试预采纳证据失效"},
    )
    assert response.status_code == 409
    assert "预采纳项" in response.text
    with SessionLocal() as db:
        section = db.get(BidSection, UUID(section_id))
        review_item = db.get(AgentReviewItem, review_item_id)
        assert section is not None
        assert review_item is not None
        assert section.assist_stage == "awaiting_confirm"
        assert review_item.status == "open"


def test_unlock_generated_section_invalidates_tender_format_exports() -> None:
    client = TestClient(app)
    project_id, section_id = create_confirmable_section_fixture(assist_stage="generated")
    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        section = db.get(BidSection, UUID(section_id))
        assert project is not None
        assert section is not None
        user = db.scalar(select(User).where(User.tenant_id == project.tenant_id, User.status == "active"))
        assert user is not None
        confirmation = SectionConfirmation(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            status="active",
            snapshot_json={"assist_stage": "generated"},
            confirmed_by=user.id,
            confirmed_at=datetime.now(UTC),
        )
        export_file = ExportFile(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section.id,
            export_type="tender_format_docx",
            file_name="generated.docx",
            bucket="test",
            object_key=f"test/generated-{uuid4().hex}.docx",
            sha256="a" * 64,
            status="available",
            source_snapshot_json={"export_mode": "submission"},
            created_by=user.id,
        )
        db.add_all([confirmation, export_file])
        db.commit()
        confirmation_id = confirmation.id
        export_file_id = export_file.id

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/unlock",
        json={"reason": "测试生成后撤回"},
    )
    assert response.status_code == 200, response.text
    assert response.json()["assist_stage"] == "awaiting_confirm"

    with SessionLocal() as db:
        section = db.get(BidSection, UUID(section_id))
        confirmation = db.get(SectionConfirmation, confirmation_id)
        export_file = db.get(ExportFile, export_file_id)
        assert section is not None
        assert confirmation is not None
        assert export_file is not None
        assert section.assist_stage == "awaiting_confirm"
        assert confirmation.status == "withdrawn"
        assert export_file.status == "deleted"
        assert export_file.source_snapshot_json is not None
        assert export_file.source_snapshot_json["invalidated_by_unlock_reason"] == "测试生成后撤回"


def test_preflight_check_detects_outdated_matrix_version() -> None:
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)

    initial_response = client.get(f"/api/v1/projects/{project_id}/sections/{section_id}/preflight-check")
    assert initial_response.status_code == 200

    with SessionLocal() as db:
        seed_matrix_item = db.scalar(
            select(ComplianceItem).where(
                ComplianceItem.project_id == UUID(project_id),
                ComplianceItem.section_id == UUID(section_id),
                ComplianceItem.normalized_requirement == "provide_valid_business_license",
                ComplianceItem.deleted_at.is_(None),
            )
        )
        assert seed_matrix_item is not None
        document = db.get(Document, seed_matrix_item.source_document_id)
        assert document is not None
        assert document.current_version_id is not None
        current = db.get(DocumentVersion, document.current_version_id)
        assert current is not None
        source_version = db.scalar(
            select(DocumentVersion)
            .join(DocumentChunk, DocumentChunk.document_version_id == DocumentVersion.id)
            .where(DocumentVersion.document_id == document.id)
            .order_by(DocumentVersion.version_no.asc())
        )
        assert source_version is not None
        max_version_no = db.scalar(
            select(func.max(DocumentVersion.version_no)).where(DocumentVersion.document_id == document.id)
        ) or current.version_no
        new_version = DocumentVersion(
            tenant_id=document.tenant_id,
            document_id=document.id,
            version_no=max_version_no + 1,
            version_label=f"v0.outdated-test-{max_version_no + 1}",
            object_key=current.object_key,
            sha256=current.sha256,
            parse_status="succeeded",
            parser_name="manual-editor",
            parser_version="1.0",
            created_by=document.created_by,
            change_reason="测试矩阵过期提醒",
        )
        db.add(new_version)
        db.flush()
        chunks = db.scalars(
            select(DocumentChunk).where(DocumentChunk.document_version_id == source_version.id)
        ).all()
        for chunk in chunks:
            db.add(
                DocumentChunk(
                    tenant_id=chunk.tenant_id,
                    document_id=chunk.document_id,
                    document_version_id=new_version.id,
                    section_id=chunk.section_id,
                    chunk_index=chunk.chunk_index,
                    page_no=chunk.page_no,
                    heading_path=chunk.heading_path,
                    content_text=chunk.content_text,
                    content_hash=chunk.content_hash,
                    bbox_json=chunk.bbox_json,
                    table_json=chunk.table_json,
                )
            )
        document.current_version_id = new_version.id
        db.commit()

    response = client.get(f"/api/v1/projects/{project_id}/sections/{section_id}/preflight-check")
    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "block"
    assert payload["matrix_outdated"] is True
    assert payload["outdated_item_count"] >= 1
    assert any(item["code"] == "matrix_version" and item["status"] == "block" for item in payload["checks"])


def test_project_read_api_enforces_tenant_context() -> None:
    client = TestClient(app)

    response = client.get("/api/v1/projects", headers={"X-Tenant-Code": "missing"})

    assert response.status_code == 404


def test_create_project_and_section_api() -> None:
    client = TestClient(app)

    create_response = client.post(
        "/api/v1/projects",
        json={
            "name": f"API 新建项目 {uuid4().hex}",
            "purchaser": "测试采购人",
            "agency": "测试代理机构",
            "budget_amount": "1000000.00",
            "region_code": "CN-4306",
            "industry_code": "municipal",
            "notice_url": "https://example.com/notice",
            "section_name": "一标段：测试标段",
        },
    )
    assert create_response.status_code == 201
    project = create_response.json()
    assert project["section_count"] == 1
    assert project["status"] == "pending_files"

    section_response = client.post(
        f"/api/v1/projects/{project['id']}/sections",
        json={"code": "section-002", "name": "二标段：补充测试"},
    )
    assert section_response.status_code == 201
    section = section_response.json()
    assert section["name"] == "二标段：补充测试"

    sections = client.get(f"/api/v1/projects/{project['id']}/sections").json()
    assert len(sections) == 2


def test_update_project_and_section_key_fields_writes_audit_log() -> None:
    client = TestClient(app)
    create_response = client.post(
        "/api/v1/projects",
        json={
            "name": f"关键字段更新项目 {uuid4().hex}",
            "section_name": "一标段：待更新",
        },
    )
    assert create_response.status_code == 201
    created_project = create_response.json()
    project_id = created_project["id"]
    sections = client.get(f"/api/v1/projects/{project_id}/sections").json()
    section_id = sections[0]["id"]

    project_response = client.patch(
        f"/api/v1/projects/{project_id}",
        json={
            "purchaser": "更新后的采购人",
            "budget_amount": "1880000.00",
            "bid_deadline_at": "2026-01-08T09:30:00+08:00",
            "reason": "测试更新项目关键信息",
        },
    )
    assert project_response.status_code == 200
    project = project_response.json()
    assert project["purchaser"] == "更新后的采购人"
    assert project["budget_amount"] == "1880000.00"

    section_response = client.patch(
        f"/api/v1/projects/{project_id}/sections/{section_id}",
        json={
            "code": "cleanroom-001",
            "name": "更新后的标段",
            "budget_amount": "988000.00",
            "reason": "测试更新标段关键信息",
        },
    )
    assert section_response.status_code == 200
    section = section_response.json()
    assert section["code"] == "cleanroom-001"
    assert section["name"] == "更新后的标段"

    audit_response = client.get(f"/api/v1/projects/{project_id}/audit-logs")
    assert audit_response.status_code == 200
    actions = [item["action"] for item in audit_response.json()]
    assert "project.updated" in actions
    assert "section.updated" in actions


def test_delete_project_archives_and_hides_from_default_list() -> None:
    client = TestClient(app)
    project_name = f"待删除项目 {uuid4().hex}"

    create_response = client.post(
        "/api/v1/projects",
        json={"name": project_name, "section_name": "一标段：删除测试"},
    )
    assert create_response.status_code == 201
    project_id = create_response.json()["id"]

    delete_response = client.delete(f"/api/v1/projects/{project_id}")
    assert delete_response.status_code == 204

    projects = client.get("/api/v1/projects").json()
    assert all(item["id"] != project_id for item in projects)

    archived_projects = client.get("/api/v1/projects", params={"status": "archived"}).json()
    archived = next(item for item in archived_projects if item["id"] == project_id)
    assert archived["status"] == "archived"

    with SessionLocal() as db:
        project = db.get(Project, UUID(project_id))
        assert project is not None
        assert project.archived_at is not None
        audit = db.scalar(
            select(AuditLog).where(
                AuditLog.project_id == UUID(project_id),
                AuditLog.action == "project.archived",
            )
        )
        assert audit is not None


def build_import_docx_bytes(project_name: str) -> bytes:
    doc = DocxDocument()
    doc.add_heading(f"{project_name}招标公告", level=1)
    doc.add_paragraph(f"项目名称：{project_name}")
    doc.add_paragraph("招标人：岳阳市君山区城市管理局")
    doc.add_paragraph("招标代理机构：中技建设咨询有限公司")
    doc.add_paragraph("最高投标限价：1349.09万元")
    doc.add_paragraph("投标截止时间：2025年12月11日09时30分")
    doc.add_paragraph("本项目分为一个标段")
    doc.add_paragraph("投标人须提供有效营业执照，并加盖公章。")
    doc.add_paragraph("保修要求：按国务院2000年279号令相关规定及合同约定进行施工质量保修。")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def test_project_import_from_uploaded_document_creates_project_document_and_matrix(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(settings, "llm_provider", "mock")
    monkeypatch.setattr(settings, "llm_api_key", "")
    settings.run_tasks_inline = True

    def fake_chat_completion(db, **kwargs):  # noqa: ANN001, ARG001
        prompt_version = kwargs["prompt_version"]
        if prompt_version == "document_section_plan@1.1.0":
            # Word 导入文本无真实分页，分块按 chunk 顺序映射为合成页码，
            # 章节范围需覆盖全部分块，否则后续按页抽取会漏掉正文条款。
            user_content = kwargs["messages"][-1]["content"]
            pages = json.loads(user_content.rsplit("pages:\n", 1)[1])
            end_page = max(page["page_no"] for page in pages)
            content = {
                "sections": [
                    {
                        "section_index": 1,
                        "title": "招标公告",
                        "section_type": "announcement",
                        "start_page": 1,
                        "end_page": end_page,
                        "confidence_score": 0.9,
                        "evidence": "导入文件招标公告。",
                    }
                ]
            }
        elif prompt_version == "compliance_extract_by_section@1.1.0":
            user_content = kwargs["messages"][-1]["content"]
            chunks = json.loads(user_content.rsplit("chunks:\n", 1)[1])

            # 大章节会被拆成多段抽取，每次只看到本段分块；按 needle 命中本段才产出条款，
            # 命中不到则返回 None 并过滤掉，避免对不含该条款的段落误报或抛异常。
            def chunk_for(needle: str) -> dict | None:
                return next((item for item in chunks if needle in item["text"]), None)

            def make_item(needle: str, requirement_text: str, item_type: str) -> dict | None:
                chunk = chunk_for(needle)
                if chunk is None:
                    return None
                return {
                    "source_chunk_index": chunk["chunk_index"],
                    "item_type": item_type,
                    "requirement_text": requirement_text,
                    "risk_level": "high",
                    "is_mandatory": True,
                    "source_quote": requirement_text,
                    "confidence_score": 0.9,
                }

            content = {
                "items": [
                    candidate
                    for candidate in (
                        make_item(
                            "营业执照",
                            "投标人须提供有效营业执照，并加盖公章。",
                            "qualification",
                        ),
                        make_item(
                            "最高投标限价",
                            "最高投标限价：1349.09万元。",
                            "mandatory_response",
                        ),
                    )
                    if candidate is not None
                ]
            }
        else:
            content = {"status": "passed", "issues": []}
        return SimpleNamespace(
            content=json.dumps(content, ensure_ascii=False),
            provider="fake",
            model_name="unit-test",
            log_id=None,
            usage={},
        )

    monkeypatch.setattr("app.services.compliance_generation.chat_completion", fake_chat_completion)
    client = TestClient(app)
    project_name = f"导入文件项目 {uuid4().hex}"

    draft_response = client.post(
        "/api/v1/projects/import-drafts/upload",
        files={
            "file": (
                f"{project_name}.docx",
                build_import_docx_bytes(project_name),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert draft_response.status_code == 201
    draft = draft_response.json()
    assert draft["source"]["source_type"] == "manual_import"
    assert draft["project"]["name"] == project_name
    assert draft["project"]["purchaser"] == "岳阳市君山区城市管理局"
    assert draft["project"]["budget_amount"] == "13490900.00"
    assert draft["sections"][0]["name"].startswith("一标段")

    confirm_response = client.post("/api/v1/projects/import-drafts/confirm", json=draft)
    assert confirm_response.status_code == 201
    result = confirm_response.json()
    assert result["project"]["name"] == project_name
    assert result["document_id"]
    assert result["parse_task_id"]
    assert result["matrix_task_id"]

    documents_response = client.get(
        f"/api/v1/projects/{result['project']['id']}/sections/{result['section_id']}/documents"
    )
    assert documents_response.status_code == 200
    documents = documents_response.json()
    assert len(documents) == 1
    assert documents[0]["source_type"] == "manual_import"
    assert documents[0]["current_version"]["parse_status"] == "succeeded"

    items_response = client.get(
        f"/api/v1/projects/{result['project']['id']}/sections/{result['section_id']}/compliance-items"
    )
    assert items_response.status_code == 200
    items = items_response.json()
    assert any("营业执照" in item["requirement_text"] for item in items)


def test_project_import_upload_legacy_doc_uses_conversion_for_draft(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    client = TestClient(app)
    project_name = f"旧版导入项目 {uuid4().hex}"
    conversion_calls: list[dict[str, object]] = []

    def fake_convert_legacy_doc_to_docx(data: bytes, *, filename: str) -> bytes:
        conversion_calls.append({"data": data, "filename": filename})
        return build_import_docx_bytes(project_name)

    monkeypatch.setattr(
        "app.services.project_import.convert_legacy_doc_to_docx",
        fake_convert_legacy_doc_to_docx,
    )

    response = client.post(
        "/api/v1/projects/import-drafts/upload",
        files={"file": (f"{project_name}.doc", b"legacy-word-binary", "application/msword")},
    )

    assert response.status_code == 201
    draft = response.json()
    assert draft["source"]["file_ext"] == "doc"
    assert conversion_calls == [{"data": b"legacy-word-binary", "filename": draft["source"]["original_filename"]}]
    assert draft["project"]["name"] == project_name
    assert draft["project"]["purchaser"] == "岳阳市君山区城市管理局"
    assert draft["project"]["budget_amount"] == "13490900.00"
    assert "营业执照" in draft["preview_text"]


def test_project_import_upload_legacy_doc_conversion_failure_returns_client_error(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    client = TestClient(app)

    def fake_convert_legacy_doc_to_docx(data: bytes, *, filename: str) -> bytes:  # noqa: ARG001
        raise LegacyDocConversionError(
            "旧版 .doc 自动转换依赖 LibreOffice/soffice，当前环境未安装转换器",
            code="LEGACY_DOC_CONVERTER_UNAVAILABLE",
        )

    monkeypatch.setattr(
        "app.services.project_import.convert_legacy_doc_to_docx",
        fake_convert_legacy_doc_to_docx,
    )

    response = client.post(
        "/api/v1/projects/import-drafts/upload",
        files={"file": ("转换失败.doc", b"legacy-word-binary", "application/msword")},
    )

    assert response.status_code == 422
    assert "旧版 .doc 自动转换依赖 LibreOffice/soffice" in response.json()["detail"]
    assert "请手工转换为 .docx" in response.json()["detail"]


def test_project_import_confirm_accepts_async_processing(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(settings, "llm_provider", "mock")
    monkeypatch.setattr(settings, "llm_api_key", "")
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_name = f"异步导入文件项目 {uuid4().hex}"

    draft_response = client.post(
        "/api/v1/projects/import-drafts/upload",
        files={
            "file": (
                f"{project_name}.docx",
                build_import_docx_bytes(project_name),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert draft_response.status_code == 201
    draft = draft_response.json()
    draft["async_processing"] = True

    confirm_response = client.post("/api/v1/projects/import-drafts/confirm", json=draft)
    assert confirm_response.status_code == 201
    result = confirm_response.json()
    assert result["project"]["name"] == project_name
    assert result["parse_task_id"]
    assert result["matrix_task_id"]


def test_project_import_upload_uses_tender_document_size_limit(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("app.api.v1.routes.projects.TENDER_DOCUMENT_FILE_MAX_BYTES", 1024 * 1024)
    client = TestClient(app)

    response = client.post(
        "/api/v1/projects/import-drafts/upload",
        files={"file": ("large.pdf", b"x" * (1024 * 1024 + 1), "application/pdf")},
    )

    assert response.status_code == 413
    assert response.json()["detail"] == "File is too large; max 1 MiB"


def test_project_import_upload_returns_client_error_for_image_only_pdf() -> None:
    document = fitz.open()
    document.new_page(width=595, height=842)
    pdf_bytes = document.tobytes()
    document.close()
    client = TestClient(app)

    response = client.post(
        "/api/v1/projects/import-drafts/upload",
        files={"file": ("scan.pdf", pdf_bytes, "application/pdf")},
    )

    assert response.status_code == 422
    assert "PDF 未提取到可用文本" in response.json()["detail"]
    assert "当前只支持可复制文本" in response.json()["detail"]


def test_project_import_from_public_url_html_creates_parsed_project(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(settings, "llm_provider", "mock")
    monkeypatch.setattr(settings, "llm_api_key", "")
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_name = f"网页导入项目 {uuid4().hex}"
    html = f"""
    <html><body>
    <h1>{project_name}招标公告</h1>
    <p>项目名称：{project_name}</p>
    <p>招标人：岳阳市君山区城市管理局</p>
    <p>招标代理机构：中技建设咨询有限公司</p>
    <p>最高投标限价：1349.09万元</p>
    <p>开标时间：2025年12月11日09时30分</p>
    <p>投标人须提供有效营业执照，并加盖公章。</p>
    </body></html>
    """.encode()

    def fake_fetch_public_file(url: str) -> DownloadedFile:
        return DownloadedFile(
            final_url=url,
            filename="notice.html",
            content_type="text/html; charset=utf-8",
            content_length=len(html),
            http_status=200,
            redirect_chain=[],
            data=html,
        )

    monkeypatch.setattr("app.services.project_import.fetch_public_file", fake_fetch_public_file)

    draft_response = client.post(
        "/api/v1/projects/import-drafts/public-url",
        json={
            "source_url": "https://example.com/tender/notice.html",
            "source_site": "示例公共资源交易中心",
        },
    )
    assert draft_response.status_code == 201
    draft = draft_response.json()
    assert draft["source"]["source_type"] == "public_url"
    assert draft["project"]["name"] == project_name

    confirm_response = client.post("/api/v1/projects/import-drafts/confirm", json=draft)
    assert confirm_response.status_code == 201
    result = confirm_response.json()
    assert result["project"]["notice_url"] == "https://example.com/tender/notice.html"
    assert result["parse_task_id"] is None
    assert result["matrix_task_id"]

    with SessionLocal() as db:
        document = db.get(Document, UUID(result["document_id"]))
        assert document is not None
        assert document.file_ext == "html"
        assert document.current_version_id is not None
        version = db.get(DocumentVersion, document.current_version_id)
        assert version is not None
        assert version.parse_status == "succeeded"
        chunk_ids = db.scalars(
            select(DocumentChunk.id).where(DocumentChunk.document_version_id == version.id)
        ).all()
        assert chunk_ids


def create_api_test_compliance_item(
    project_id: str,
    section_id: str,
    *,
    risk_level: str = "low",
    is_mandatory: bool = False,
    owner_user_id: UUID | None = None,
) -> str:
    with SessionLocal() as db:
        document = db.scalar(
            select(Document).where(
                Document.project_id == UUID(project_id),
                Document.section_id == UUID(section_id),
                Document.original_filename == "招标文件.pdf",
            )
        )
        assert document is not None
        chunk = db.scalar(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == document.id)
            .order_by(DocumentChunk.created_at.asc())
        )
        assert chunk is not None
        version = db.get(DocumentVersion, chunk.document_version_id)
        assert version is not None
        user = db.scalar(select(User).where(User.external_id == "demo-admin"))
        assert user is not None
        item = ComplianceItem(
            tenant_id=document.tenant_id,
            project_id=document.project_id,
            section_id=UUID(section_id),
            source_document_id=document.id,
            source_version_id=version.id,
            source_chunk_id=chunk.id,
            source_page_no=chunk.page_no,
            item_type="mandatory_response",
            requirement_text=f"API 测试响应项 {uuid4().hex}",
            normalized_requirement=f"api_test:{uuid4().hex}",
            response_suggestion="请人工确认响应内容。",
            evidence_text=chunk.content_text,
            status="pending_confirm",
            risk_level=risk_level,
            is_mandatory=is_mandatory,
            is_batch_confirm_allowed=risk_level != "high" and not is_mandatory,
            owner_user_id=owner_user_id,
            confidence_score=Decimal("0.8000"),
            created_by=user.id,
        )
        db.add(item)
        db.commit()
        return str(item.id)


def create_review_test_chunks(
    project_id: str,
    section_id: str,
    contents: list[str],
) -> list[dict[str, object]]:
    with SessionLocal() as db:
        document = db.scalar(
            select(Document).where(
                Document.project_id == UUID(project_id),
                Document.section_id == UUID(section_id),
                Document.original_filename == "招标文件.pdf",
            )
        )
        assert document is not None
        assert document.current_version_id is not None
        max_index = db.scalar(
            select(func.max(DocumentChunk.chunk_index)).where(
                DocumentChunk.document_version_id == document.current_version_id
            )
        ) or 0
        chunks: list[DocumentChunk] = []
        for index, content in enumerate(contents, start=1):
            chunk = DocumentChunk(
                tenant_id=document.tenant_id,
                document_id=document.id,
                document_version_id=document.current_version_id,
                section_id=UUID(section_id),
                chunk_index=max_index + index,
                page_no=900 + index,
                heading_path=f"MVP1.1 P1 回归/第 {index} 处",
                content_text=content,
                content_hash=hashlib.sha256(content.encode("utf-8")).hexdigest(),
                bbox_json=None,
                table_json=None,
            )
            db.add(chunk)
            chunks.append(chunk)
        db.commit()
        return [
            {
                "id": str(chunk.id),
                "chunk_index": chunk.chunk_index,
                "page_no": chunk.page_no,
                "content_text": chunk.content_text,
            }
            for chunk in chunks
        ]


def create_review_test_item(
    project_id: str,
    section_id: str,
    chunk_id: str,
    *,
    requirement_text: str,
    dedup_key: str,
    risk_level: str = "low",
    is_mandatory: bool = False,
) -> str:
    with SessionLocal() as db:
        chunk = db.get(DocumentChunk, UUID(chunk_id))
        assert chunk is not None
        document = db.get(Document, chunk.document_id)
        assert document is not None
        user = db.scalar(select(User).where(User.external_id == "demo-admin"))
        assert user is not None
        item = ComplianceItem(
            tenant_id=document.tenant_id,
            project_id=UUID(project_id),
            section_id=UUID(section_id),
            source_document_id=document.id,
            source_version_id=chunk.document_version_id,
            source_chunk_id=chunk.id,
            source_page_no=chunk.page_no,
            item_type="mandatory_response",
            requirement_text=requirement_text,
            normalized_requirement=f"p1_test:{dedup_key}:{uuid4().hex}",
            dedup_key=dedup_key,
            response_suggestion="P1 回归测试项，请人工确认。",
            evidence_text=chunk.content_text,
            explanation_json={
                "rule_code": "TEST-MVP11-P1",
                "rule_reason": "MVP1.1 P1 关联组回归测试注入项。",
            },
            status="pending_confirm",
            risk_level=risk_level,
            is_mandatory=is_mandatory,
            is_batch_confirm_allowed=risk_level != "high" and not is_mandatory,
            owner_user_id=user.id,
            confidence_score=Decimal("0.9000"),
            created_by=user.id,
        )
        db.add(item)
        db.commit()
        return str(item.id)


def test_compliance_item_update_confirm_assign_and_audit() -> None:
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    item_id = create_api_test_compliance_item(project_id, section_id)

    missing_reason_response = client.patch(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/{item_id}",
        json={"risk_level": "medium"},
    )
    assert missing_reason_response.status_code == 422

    update_response = client.patch(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/{item_id}",
        json={
            "risk_level": "medium",
            "is_mandatory": True,
            "status": "needs_material",
            "response_suggestion": "需补充承诺函并复核。",
            "reason": "测试人工修改矩阵项",
        },
    )
    assert update_response.status_code == 200
    updated = update_response.json()
    assert updated["risk_level"] == "medium"
    assert updated["is_mandatory"] is True
    assert updated["status"] == "needs_material"
    assert updated["is_batch_confirm_allowed"] is False
    assert updated["modify_reason"] == "测试人工修改矩阵项"

    assign_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/{item_id}/assign",
        json={"reason": "测试指派给当前用户"},
    )
    assert assign_response.status_code == 200
    assigned = assign_response.json()
    assert assigned["owner_name"] == "演示管理员"

    blocked_confirm_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/{item_id}/confirm",
        json={"reason": "测试单条确认"},
    )
    assert blocked_confirm_response.status_code == 409
    assert "source verification" in blocked_confirm_response.json()["detail"]

    confirm_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/{item_id}/confirm",
        json={"reason": "测试单条确认", "source_verified": True},
    )
    assert confirm_response.status_code == 200
    confirmed = confirm_response.json()
    assert confirmed["status"] == "confirmed"
    assert confirmed["confirmed_by"] is not None

    audit_response = client.get(
        f"/api/v1/projects/{project_id}/audit-logs",
        params={"object_type": "compliance_item", "limit": 50},
    )
    assert audit_response.status_code == 200
    actions = {
        item["action"]
        for item in audit_response.json()
        if item["object_id"] == item_id
    }
    assert {"matrix.item_updated", "matrix.item_assigned", "matrix.item_confirmed"} <= actions


def test_matrix_review_p1_manual_source_similar_candidates_and_apply() -> None:
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    unique = uuid4().hex[:8]
    selected_text = f"投标人须提供洁净设备安装调试方案 {unique}。"
    fuzzy_text = f"投标人须提供洁净设备调试安装方案 {unique}。"
    chunks = create_review_test_chunks(
        project_id,
        section_id,
        [
            f"补充条款：{selected_text} 本条用于人工补漏回归。",
            f"重复条款：{selected_text} 请在响应文件中承诺。",
            f"近似条款：{fuzzy_text}",
            f"采购人：测试单位 联系电话：13800000000 {selected_text}",
        ],
    )
    base_chunk = chunks[0]

    review_before_response = client.get(f"/api/v1/projects/{project_id}/sections/{section_id}/matrix-review")
    assert review_before_response.status_code == 200
    review_before = review_before_response.json()
    assert any(item["chunk"]["id"] == base_chunk["id"] for item in review_before["uncovered_chunks"])

    start = str(base_chunk["content_text"]).index(selected_text)
    create_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/from-source",
        json={
            "source_chunk_id": base_chunk["id"],
            "selected_text": selected_text,
            "selection_start_offset": start,
            "selection_end_offset": start + len(selected_text),
            "item_type": "mandatory_response",
            "risk_level": "medium",
            "is_mandatory": True,
            "response_suggestion": "请提供安装调试方案并在响应文件中承诺。",
            "reason": "P1 回归：人工从原文补漏新增",
        },
    )
    assert create_response.status_code == 201
    created_result = create_response.json()
    created_item = created_result["item"]
    assert created_item["source_create_method"] == "manual_selection"
    assert created_item["selected_text"] == selected_text
    assert created_item["selection_start_offset"] == start
    assert created_item["selection_end_offset"] == start + len(selected_text)
    assert created_item["dedup_key"]

    candidates = created_result["similar_candidates"]
    candidate_chunk_ids = {candidate["source_chunk_id"] for candidate in candidates}
    assert chunks[1]["id"] in candidate_chunk_ids
    assert chunks[2]["id"] in candidate_chunk_ids
    assert chunks[3]["id"] not in candidate_chunk_ids
    fuzzy_candidate = next(candidate for candidate in candidates if candidate["source_chunk_id"] == chunks[2]["id"])
    assert fuzzy_candidate["match_type"] == "fuzzy"
    assert any(segment["operation"] in {"insert", "replace"} for segment in fuzzy_candidate["diff_segments"])

    exact_candidate = next(candidate for candidate in candidates if candidate["source_chunk_id"] == chunks[1]["id"])
    apply_response = client.post(
        (
            f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/"
            f"{created_item['id']}/similar-candidates/apply"
        ),
        json={
            "reason": "P1 回归：逐个确认相似补票",
            "candidates": [
                {
                    "candidate_key": exact_candidate["candidate_key"],
                    "source_chunk_id": exact_candidate["source_chunk_id"],
                    "selected_text": exact_candidate["selected_text"],
                    "selection_start_offset": exact_candidate["selection_start_offset"],
                    "selection_end_offset": exact_candidate["selection_end_offset"],
                    "action": "join_group",
                },
                {
                    "candidate_key": fuzzy_candidate["candidate_key"],
                    "source_chunk_id": fuzzy_candidate["source_chunk_id"],
                    "selected_text": fuzzy_candidate["selected_text"],
                    "selection_start_offset": fuzzy_candidate["selection_start_offset"],
                    "selection_end_offset": fuzzy_candidate["selection_end_offset"],
                    "action": "create_independent",
                },
            ],
        },
    )
    assert apply_response.status_code == 200
    apply_result = apply_response.json()
    assert apply_result["affected_item_count"] == 3
    joined_items = [
        item for item in apply_result["items"]
        if item["source_chunk_id"] in {created_item["source_chunk_id"], chunks[1]["id"]}
    ]
    independent_items = [item for item in apply_result["items"] if item["source_chunk_id"] == chunks[2]["id"]]
    assert len(joined_items) == 2
    assert {item["duplicate_group_id"] for item in joined_items} == {apply_result["duplicate_group_id"]}
    assert independent_items[0]["duplicate_group_id"] is None

    audit_response = client.get(
        f"/api/v1/projects/{project_id}/audit-logs",
        params={"limit": 100},
    )
    assert audit_response.status_code == 200
    actions = {item["action"] for item in audit_response.json()}
    assert {"matrix.item_created_from_source", "matrix.similar_candidate_applied"} <= actions


def test_matrix_review_p1_duplicate_group_unlink_split_and_cascade_confirmation() -> None:
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    unique = uuid4().hex[:8]
    requirement = f"投标人须提供项目经理驻场服务承诺 {unique}。"
    dedup_key = f"p1_duplicate_{unique}"
    chunks = create_review_test_chunks(
        project_id,
        section_id,
        [
            f"第一处：{requirement}",
            f"第二处：{requirement}",
            f"第三处：{requirement}",
        ],
    )
    item_ids = [
        create_review_test_item(
            project_id,
            section_id,
            str(chunk["id"]),
            requirement_text=requirement,
            dedup_key=dedup_key,
        )
        for chunk in chunks
    ]

    review_candidate_response = client.get(f"/api/v1/projects/{project_id}/sections/{section_id}/matrix-review")
    assert review_candidate_response.status_code == 200
    candidate_groups = [
        group for group in review_candidate_response.json()["duplicate_groups"]
        if group["group_key"] == dedup_key and group["group_type"] == "candidate"
    ]
    assert candidate_groups and candidate_groups[0]["item_count"] == 3

    confirm_group_response = client.post(
        (
            f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/"
            f"{item_ids[0]}/duplicate-group/confirm"
        ),
        json={"reason": "P1 回归：人工确认重复关联组"},
    )
    assert confirm_group_response.status_code == 200
    group_result = confirm_group_response.json()
    assert group_result["affected_item_count"] == 3
    assert group_result["duplicate_group_id"]
    assert {item["duplicate_group_status"] for item in group_result["items"]} == {"confirmed"}

    unlink_response = client.post(
        (
            f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/"
            f"{item_ids[2]}/duplicate-group/unlink"
        ),
        json={"reason": "P1 回归：第三处存在语义差异，解除联动"},
    )
    assert unlink_response.status_code == 200
    unlinked_item = unlink_response.json()["items"][0]
    assert unlinked_item["duplicate_group_id"] is None
    assert unlinked_item["duplicate_group_status"] == "unlinked"

    confirm_item_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/{item_ids[0]}/confirm",
        json={"reason": "P1 回归：确认同组条目并验证级联", "cascade": True},
    )
    assert confirm_item_response.status_code == 200
    confirmed_item = confirm_item_response.json()
    assert confirmed_item["status"] == "confirmed"
    assert confirmed_item["cascade_affected_count"] == 1
    assert confirmed_item["cascade_affected_items"][0]["id"] == item_ids[1]

    with SessionLocal() as db:
        first = db.get(ComplianceItem, UUID(item_ids[0]))
        second = db.get(ComplianceItem, UUID(item_ids[1]))
        third = db.get(ComplianceItem, UUID(item_ids[2]))
        assert first is not None and second is not None and third is not None
        assert first.status == "confirmed"
        assert second.status == "confirmed"
        assert third.status == "pending_confirm"
        assert third.duplicate_group_status == "unlinked"

    split_response = client.post(
        (
            f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/"
            f"{item_ids[2]}/duplicate-group/split"
        ),
        json={"reason": "P1 回归：拆分为独立关联组"},
    )
    assert split_response.status_code == 200
    split_item = split_response.json()["items"][0]
    assert split_item["duplicate_group_id"]
    assert split_item["duplicate_group_id"] != group_result["duplicate_group_id"]
    assert split_item["duplicate_group_status"] == "confirmed"

    audit_response = client.get(
        f"/api/v1/projects/{project_id}/audit-logs",
        params={"limit": 120},
    )
    assert audit_response.status_code == 200
    actions = {item["action"] for item in audit_response.json()}
    assert {
        "matrix.duplicate_group_confirmed",
        "matrix.duplicate_group_unlinked",
        "matrix.duplicate_group_split",
        "matrix.cascade_confirmed",
    } <= actions


def test_compliance_item_bulk_guards() -> None:
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    low_item_id = create_api_test_compliance_item(project_id, section_id)
    high_item_id = create_api_test_compliance_item(
        project_id,
        section_id,
        risk_level="high",
        is_mandatory=True,
    )

    blocked_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/bulk-confirm",
        json={"item_ids": [high_item_id], "reason": "测试禁止批量确认高风险项"},
    )
    assert blocked_response.status_code == 409

    assign_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/bulk-assign",
        json={"item_ids": [low_item_id, high_item_id], "reason": "测试批量指派"},
    )
    assert assign_response.status_code == 200
    assert {item["owner_name"] for item in assign_response.json()} == {"演示管理员"}

    confirm_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/bulk-confirm",
        json={"item_ids": [low_item_id], "reason": "测试批量确认低风险项"},
    )
    assert confirm_response.status_code == 200
    assert confirm_response.json()[0]["status"] == "confirmed"

    with SessionLocal() as db:
        batch_logs = db.scalars(
            select(AuditLog).where(
                AuditLog.project_id == UUID(project_id),
                AuditLog.action.in_(["matrix.items_assigned", "matrix.items_batch_confirmed"]),
            )
        ).all()
        assert len(batch_logs) >= 2


def test_compliance_matrix_excel_export_api() -> None:
    settings.run_tasks_inline = True
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)

    response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/export-excel",
        json={"risk_level": "low"},
    )
    assert response.status_code == 202
    task = response.json()
    assert task["status"] == "succeeded"
    assert task["output_json"]["export_file_id"]
    assert task["output_json"]["row_count"] >= 1

    export_id = task["output_json"]["export_file_id"]
    export_response = client.get(f"/api/v1/projects/{project_id}/export-files/{export_id}")
    assert export_response.status_code == 200
    export_file = export_response.json()
    assert export_file["status"] == "available"
    assert export_file["sha256"] == task["output_json"]["sha256"]
    assert export_file["file_name"].startswith("合规矩阵快照-")
    assert export_file["source_snapshot_json"]["snapshot_note"] == "只读快照；最新状态以平台为准"

    download_response = client.get(
        f"/api/v1/projects/{project_id}/export-files/{export_id}/download"
    )
    assert download_response.status_code == 200
    assert download_response.content.startswith(b"PK")
    workbook = load_workbook(BytesIO(download_response.content))
    sheet = workbook["合规矩阵快照"]
    assert sheet["A1"].value == "合规矩阵快照"
    assert "最新状态以投标 Agent 平台为准" in sheet["A2"].value
    assert sheet["A7"].value == "序号"


def test_matrix_review_returns_word_xml_document_and_highlights(monkeypatch: pytest.MonkeyPatch) -> None:
    client = TestClient(app)

    docx = DocxDocument()
    section = docx.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header.paragraphs[0].text = "测试页眉"
    section.footer.paragraphs[0].text = "测试页脚"
    docx.add_heading("第一章 投标人须知", level=1)
    paragraph = docx.add_paragraph()
    paragraph.alignment = 0
    run = paragraph.add_run("投标人须提供有效营业执照，并加盖公章。")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x22, 0x33, 0x44)
    table = docx.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "材料"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "检测报告"
    table.cell(1, 1).text = "须提供净化设备检测报告"
    buffer = BytesIO()
    docx.save(buffer)
    payload = buffer.getvalue()

    with SessionLocal() as db:
        user = db.scalar(select(User))
        assert user is not None
        project = Project(
            tenant_id=user.tenant_id,
            name=f"Word XML 审阅隔离测试-{uuid4()}",
            purchaser="测试采购人",
            agency="测试代理",
            status="pending_confirm",
            created_by=user.id,
        )
        db.add(project)
        db.flush()
        db.add(
            ProjectMember(
                tenant_id=project.tenant_id,
                project_id=project.id,
                user_id=user.id,
                role_code="owner",
                status="active",
                created_by=user.id,
            )
        )
        section_model = BidSection(
            tenant_id=project.tenant_id,
            project_id=project.id,
            code="word-review",
            name="Word 原文审阅测试标段",
            status="pending_confirm",
            created_by=user.id,
        )
        db.add(section_model)
        db.flush()
        project_id = str(project.id)
        section_id = str(section_model.id)
        document = Document(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section_model.id,
            doc_type="tender",
            title="Word XML 审阅测试招标文件",
            source_type="upload",
            original_filename="word-review-test.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_ext="docx",
            file_size=len(payload),
            sha256=hashlib.sha256(payload).hexdigest(),
            bucket=settings.minio_bucket,
            object_key="tests/word-review-test.docx",
            status="available",
            created_by=user.id,
            acquired_at=datetime.now(UTC),
        )
        db.add(document)
        db.flush()
        version = DocumentVersion(
            tenant_id=project.tenant_id,
            document_id=document.id,
            version_no=1,
            version_label="v0.word-review",
            object_key=document.object_key,
            sha256=document.sha256,
            parse_status="succeeded",
            parser_name="word",
            parser_version="1.0",
            created_by=user.id,
        )
        db.add(version)
        db.flush()
        document.current_version_id = version.id
        paragraph_text = "投标人须提供有效营业执照，并加盖公章。"
        table_text = "材料 | 要求\n检测报告 | 须提供净化设备检测报告"
        paragraph_chunk = DocumentChunk(
            tenant_id=project.tenant_id,
            document_id=document.id,
            document_version_id=version.id,
            section_id=section_model.id,
            chunk_index=2,
            page_no=None,
            heading_path="第一章 投标人须知",
            content_text=paragraph_text,
            content_hash=hashlib.sha256(paragraph_text.encode("utf-8")).hexdigest(),
        )
        table_chunk = DocumentChunk(
            tenant_id=project.tenant_id,
            document_id=document.id,
            document_version_id=version.id,
            section_id=section_model.id,
            chunk_index=3,
            page_no=None,
            heading_path="第一章 投标人须知",
            content_text=table_text,
            content_hash=hashlib.sha256(table_text.encode("utf-8")).hexdigest(),
            table_json={"rows": [["材料", "要求"], ["检测报告", "须提供净化设备检测报告"]]},
        )
        db.add_all([paragraph_chunk, table_chunk])
        db.flush()
        item = ComplianceItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section_model.id,
            source_document_id=document.id,
            source_version_id=version.id,
            source_chunk_id=paragraph_chunk.id,
            item_type="qualification",
            requirement_text="有效营业执照",
            normalized_requirement="word_review_business_license",
            response_suggestion="请提供营业执照。",
            evidence_text=paragraph_text,
            explanation_json={"source_quote": "有效营业执照"},
            status="pending_confirm",
            risk_level="high",
            is_mandatory=True,
            is_batch_confirm_allowed=False,
            confidence_score=Decimal("0.9500"),
            created_by=user.id,
        )
        db.add(item)
        db.commit()

    monkeypatch.setattr("app.api.v1.routes.projects.get_object_bytes", lambda **_: payload)

    response = client.get(f"/api/v1/projects/{project_id}/sections/{section_id}/matrix-review")

    assert response.status_code == 200
    review = response.json()
    assert review["review_document"]["mode"] == "word_xml"
    assert review["review_document"]["headers"] == ["测试页眉"]
    assert review["review_document"]["footers"] == ["测试页脚"]
    assert review["review_document"]["page_margins"]["top"] == 57.6
    assert any(block["type"] == "table" for block in review["review_document"]["blocks"])
    highlight = next(item for item in review["highlights"] if item["text"] == "有效营业执照")
    assert highlight["risk_level"] == "high"
    assert highlight["match_source"] == "source_quote"


def test_matrix_review_returns_pdf_layout_document_and_highlights() -> None:
    client = TestClient(app)
    payload = b"%PDF-test-layout-placeholder"

    with SessionLocal() as db:
        user = db.scalar(select(User))
        assert user is not None
        project = Project(
            tenant_id=user.tenant_id,
            name=f"PDF 原文审阅隔离测试-{uuid4()}",
            purchaser="测试采购人",
            agency="测试代理",
            status="pending_confirm",
            created_by=user.id,
        )
        db.add(project)
        db.flush()
        db.add(
            ProjectMember(
                tenant_id=project.tenant_id,
                project_id=project.id,
                user_id=user.id,
                role_code="owner",
                status="active",
                created_by=user.id,
            )
        )
        section_model = BidSection(
            tenant_id=project.tenant_id,
            project_id=project.id,
            code="pdf-review",
            name="PDF 原文审阅测试标段",
            status="pending_confirm",
            created_by=user.id,
        )
        db.add(section_model)
        db.flush()
        project_id = str(project.id)
        section_id = str(section_model.id)
        document = Document(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section_model.id,
            doc_type="tender",
            title="PDF 原文审阅测试招标文件",
            source_type="upload",
            original_filename="pdf-review-test.pdf",
            content_type="application/pdf",
            file_ext="pdf",
            file_size=len(payload),
            sha256=hashlib.sha256(payload).hexdigest(),
            bucket=settings.minio_bucket,
            object_key="tests/pdf-review-test.pdf",
            status="available",
            created_by=user.id,
            acquired_at=datetime.now(UTC),
        )
        db.add(document)
        db.flush()
        version = DocumentVersion(
            tenant_id=project.tenant_id,
            document_id=document.id,
            version_no=1,
            version_label="v0.pdf-review",
            object_key=document.object_key,
            sha256=document.sha256,
            parse_status="succeeded",
            parser_name="pdf-parser",
            parser_version="pymupdf-layout-v1",
            created_by=user.id,
        )
        db.add(version)
        db.flush()
        document.current_version_id = version.id
        paragraph_text = "Bidder must provide a valid business license."
        table_text = "Material | Requirement\nReport | Valid test report"
        paragraph_chunk = DocumentChunk(
            tenant_id=project.tenant_id,
            document_id=document.id,
            document_version_id=version.id,
            section_id=section_model.id,
            chunk_index=1,
            page_no=1,
            heading_path="PDF 第 1 页",
            content_text=paragraph_text,
            content_hash=hashlib.sha256(paragraph_text.encode("utf-8")).hexdigest(),
            bbox_json={
                "page_no": 1,
                "page_width": 595.0,
                "page_height": 842.0,
                "x0": 72.0,
                "y0": 72.0,
                "x1": 360.0,
                "y1": 96.0,
                "block_type": "text",
                "parser_version": "pymupdf-layout-v1",
            },
        )
        table_chunk = DocumentChunk(
            tenant_id=project.tenant_id,
            document_id=document.id,
            document_version_id=version.id,
            section_id=section_model.id,
            chunk_index=2,
            page_no=2,
            heading_path="PDF 第 2 页",
            content_text=table_text,
            content_hash=hashlib.sha256(table_text.encode("utf-8")).hexdigest(),
            bbox_json={
                "page_no": 2,
                "page_width": 595.0,
                "page_height": 842.0,
                "x0": 72.0,
                "y0": 120.0,
                "x1": 420.0,
                "y1": 180.0,
                "block_type": "table",
                "parser_version": "pymupdf-layout-v1",
            },
            table_json={"rows": [["Material", "Requirement"], ["Report", "Valid test report"]]},
        )
        db.add_all([paragraph_chunk, table_chunk])
        db.flush()
        item = ComplianceItem(
            tenant_id=project.tenant_id,
            project_id=project.id,
            section_id=section_model.id,
            source_document_id=document.id,
            source_version_id=version.id,
            source_chunk_id=paragraph_chunk.id,
            item_type="qualification",
            requirement_text="valid business license",
            normalized_requirement="pdf_review_business_license",
            response_suggestion="Please provide the business license.",
            evidence_text=paragraph_text,
            explanation_json={"source_quote": "valid business license"},
            status="pending_confirm",
            risk_level="high",
            is_mandatory=True,
            is_batch_confirm_allowed=False,
            confidence_score=Decimal("0.9500"),
            created_by=user.id,
        )
        db.add(item)
        db.commit()

    response = client.get(f"/api/v1/projects/{project_id}/sections/{section_id}/matrix-review")

    assert response.status_code == 200
    review = response.json()
    assert review["review_document"]["mode"] == "pdf_layout"
    assert review["review_document"]["pages"][0]["page_no"] == 1
    assert review["review_document"]["pages"][1]["page_no"] == 2
    assert review["review_document"]["blocks"][0]["page_no"] == 1
    assert review["review_document"]["blocks"][0]["bbox_json"]["page_width"] == 595.0
    assert any(block["type"] == "table" for block in review["review_document"]["blocks"])
    highlight = next(item for item in review["highlights"] if item["text"] == "valid business license")
    assert highlight["risk_level"] == "high"
    assert highlight["match_source"] == "source_quote"
