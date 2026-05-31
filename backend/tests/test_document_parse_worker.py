from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from uuid import UUID, uuid4

import pytest
from alembic import command
from alembic.config import Config
from docx import Document as WordDocument
from fastapi.testclient import TestClient
from sqlalchemy import select

import fitz
from app.core.config import settings
from app.db.session import SessionLocal
from app.main import app
from app.models import AsyncTask, AuditLog, Document, DocumentChunk, DocumentVersion, ParseTask
from app.parsers.pdf import PdfTextEmptyError, parse_pdf_bytes
from app.services.document_parse import execute_document_parse_task
from scripts.seed_dev_data import seed


@pytest.fixture(scope="module", autouse=True)
def prepare_database() -> None:
    backend_dir = Path(__file__).resolve().parents[1]
    config = Config(str(backend_dir / "alembic.ini"))
    config.set_main_option("script_location", str(backend_dir / "migrations"))
    command.upgrade(config, "head")
    seed()


def get_seed_project_and_section(client: TestClient) -> tuple[str, str]:
    projects = client.get("/api/v1/projects").json()
    project = next(item for item in projects if item["name"] == "智慧园区弱电工程投标")
    sections = client.get(f"/api/v1/projects/{project['id']}/sections").json()
    return project["id"], sections[0]["id"]


def build_docx_bytes() -> bytes:
    document = WordDocument()
    document.add_heading("资格要求", level=1)
    document.add_paragraph("投标人须提供有效营业执照，并加盖公章。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "人员"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "项目经理"
    table.cell(1, 1).text = "同类项目经验不少于 3 年"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Bidder must provide a valid business license.")
    page.insert_text((72, 96), "Bid deadline is 2025-12-11 09:30.")
    page2 = document.new_page()
    page2.insert_text((72, 72), "The bidder shall provide a safety certificate.")
    return document.tobytes()


def build_pdf_table_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page(width=420, height=300)
    page.insert_text((72, 50), "Tender qualification table")
    page.draw_rect(fitz.Rect(72, 90, 320, 150))
    page.draw_line((72, 120), (320, 120))
    page.draw_line((180, 90), (180, 150))
    page.insert_text((82, 110), "Material")
    page.insert_text((192, 110), "Requirement")
    page.insert_text((82, 140), "Report")
    page.insert_text((192, 140), "Valid test report")
    return document.tobytes()


def build_blank_pdf_bytes() -> bytes:
    document = fitz.open()
    document.new_page()
    return document.tobytes()


def test_pdf_parser_extracts_layout_blocks_and_tables() -> None:
    chunks = parse_pdf_bytes(build_pdf_table_bytes())

    assert chunks[0].bbox_json is not None
    assert chunks[0].bbox_json["page_no"] == 1
    assert chunks[0].bbox_json["page_width"] == 420
    table_chunk = next(chunk for chunk in chunks if chunk.table_json)
    assert table_chunk.bbox_json is not None
    assert table_chunk.bbox_json["block_type"] == "table"
    assert table_chunk.table_json is not None
    assert table_chunk.table_json["rows"][0] == ["Material", "Requirement"]
    assert table_chunk.table_json["rows"][1] == ["Report", "Valid test report"]


def test_pdf_parser_rejects_empty_text_pdf() -> None:
    with pytest.raises(PdfTextEmptyError) as exc_info:
        parse_pdf_bytes(build_blank_pdf_bytes())

    assert exc_info.value.code == "PDF_TEXT_EMPTY_OCR_REQUIRED"


def test_word_parse_worker_extracts_paragraphs_and_tables(monkeypatch) -> None:
    settings.run_tasks_inline = False
    from app.worker import run_document_parse_task

    monkeypatch.setattr(run_document_parse_task, "delay", lambda task_id: None)

    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    filename = f"测试招标文件-{uuid4().hex}.docx"

    upload_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/upload",
        data={"doc_type": "tender", "title": "Word 招标文件"},
        files={
            "file": (
                filename,
                build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    document_payload = upload_response.json()

    parse_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document_payload['id']}/parse-tasks",
        json={"parser_name": "word-parser", "parser_version": "0.1.0"},
    )
    assert parse_response.status_code == 202
    parse_payload = parse_response.json()
    assert parse_payload["parser_type"] == "word"

    task_id = UUID(parse_payload["task"]["id"])
    document_id = UUID(document_payload["id"])
    version_id = UUID(document_payload["current_version_id"])

    with SessionLocal() as db:
        result = execute_document_parse_task(db, task_id)
        assert result["status"] == "succeeded"
        assert result["chunk_count"] == 3

        task = db.get(AsyncTask, task_id)
        assert task is not None
        assert task.status == "succeeded"
        assert task.output_json is not None
        assert task.output_json["chunk_count"] == 3

        document = db.get(Document, document_id)
        version = db.get(DocumentVersion, version_id)
        assert document is not None
        assert version is not None
        assert document.status == "available"
        assert version.parse_status == "succeeded"
        assert version.parser_name == "word-parser"

        chunks = db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_version_id == version_id)
            .order_by(DocumentChunk.chunk_index)
        ).all()
        assert [chunk.chunk_index for chunk in chunks] == [1, 2, 3]
        assert chunks[0].heading_path == "资格要求"
        assert "营业执照" in chunks[1].content_text
        assert chunks[2].table_json is not None
        assert chunks[2].table_json["rows"][1] == ["项目经理", "同类项目经验不少于 3 年"]

        parse_task = db.scalar(select(ParseTask).where(ParseTask.task_id == task_id))
        assert parse_task is not None
        assert parse_task.result_summary_json is not None
        assert parse_task.result_summary_json["chunk_count"] == 3

        audit_actions = db.scalars(
            select(AuditLog.action).where(AuditLog.object_id == parse_task.id)
        ).all()
        assert "document.parse_succeeded" in audit_actions

        repeated = execute_document_parse_task(db, task_id)
        assert repeated["status"] == "already_succeeded"
        assert repeated["chunk_count"] == 3


def test_pdf_text_parse_worker_extracts_pages(monkeypatch) -> None:
    settings.run_tasks_inline = False
    from app.worker import run_document_parse_task

    monkeypatch.setattr(run_document_parse_task, "delay", lambda task_id: None)

    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    filename = f"测试招标文件-{uuid4().hex}.pdf"

    upload_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/upload",
        data={"doc_type": "tender", "title": "PDF 招标文件"},
        files={"file": (filename, build_pdf_bytes(), "application/pdf")},
    )
    assert upload_response.status_code == 201
    document_payload = upload_response.json()

    parse_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document_payload['id']}/parse-tasks",
        json={"parser_name": "pdf-parser", "parser_version": "0.1.0"},
    )
    assert parse_response.status_code == 202
    parse_payload = parse_response.json()
    assert parse_payload["parser_type"] == "pdf_text"

    with SessionLocal() as db:
        result = execute_document_parse_task(db, UUID(parse_payload["task"]["id"]))
        assert result["status"] == "succeeded"
        assert result["chunk_count"] >= 1

        chunks = db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_id == UUID(document_payload["id"]))
            .order_by(DocumentChunk.chunk_index)
        ).all()
        assert chunks[0].page_no == 1
        assert "business license" in chunks[0].content_text
        assert chunks[0].bbox_json is not None
        assert chunks[0].bbox_json["page_width"] > 0
        assert chunks[-1].page_no == 2


def test_pdf_text_parse_worker_marks_empty_pdf_as_ocr_required(monkeypatch) -> None:
    settings.run_tasks_inline = False
    from app.worker import run_document_parse_task

    monkeypatch.setattr(run_document_parse_task, "delay", lambda task_id: None)

    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    filename = f"空白招标文件-{uuid4().hex}.pdf"

    upload_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/upload",
        data={"doc_type": "tender", "title": "空白 PDF 招标文件"},
        files={"file": (filename, build_blank_pdf_bytes(), "application/pdf")},
    )
    assert upload_response.status_code == 201
    document_payload = upload_response.json()

    parse_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document_payload['id']}/parse-tasks",
        json={"parser_name": "pdf-parser", "parser_version": "0.1.0"},
    )
    assert parse_response.status_code == 202
    parse_payload = parse_response.json()

    with SessionLocal() as db:
        result = execute_document_parse_task(db, UUID(parse_payload["task"]["id"]))
        assert result["status"] == "failed"
        assert result["error_code"] == "PDF_TEXT_EMPTY_OCR_REQUIRED"


def test_manual_revision_creates_new_version_and_matrix_uses_revised_chunks(monkeypatch) -> None:
    settings.run_tasks_inline = True
    from app.worker import run_document_parse_task

    monkeypatch.setattr(run_document_parse_task, "delay", lambda task_id: None)

    def fake_chat_completion(db, **kwargs):  # noqa: ANN001, ARG001
        prompt_version = kwargs["prompt_version"]
        if prompt_version == "document_section_plan@1.1.0":
            content = {
                "sections": [
                    {
                        "section_index": 1,
                        "title": "资格要求",
                        "section_type": "announcement",
                        "start_page": 1,
                        "end_page": 1,
                        "confidence_score": 0.9,
                        "evidence": "人工修正后的资格要求段落。",
                    }
                ]
            }
        elif prompt_version == "compliance_extract_by_section@1.1.0":
            user_content = kwargs["messages"][-1]["content"]
            chunks = json.loads(user_content.rsplit("chunks:\n", 1)[1])
            chunk = next(item for item in chunks if "安全生产许可证" in item["text"])
            content = {
                "items": [
                    {
                        "source_chunk_index": chunk["chunk_index"],
                        "item_type": "qualification",
                        "requirement_text": "投标人须提供有效安全生产许可证。",
                        "risk_level": "high",
                        "is_mandatory": True,
                        "source_quote": "投标人须提供有效安全生产许可证。",
                        "confidence_score": 0.9,
                    }
                ]
            }
        else:
            content = {"status": "passed", "issues": []}
        return SimpleNamespace(
            content=json.dumps(content, ensure_ascii=False),
            provider="fake",
            model_name="unit-test",
            log_id=None,
            usage={},
        )

    monkeypatch.setattr("app.services.compliance_generation.chat_completion", fake_chat_completion)

    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    filename = f"人工修正测试-{uuid4().hex}.docx"

    upload_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/upload",
        data={"doc_type": "tender", "title": "人工修正测试"},
        files={
            "file": (
                filename,
                build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    document_payload = upload_response.json()
    document_id = document_payload["id"]
    source_version_id = document_payload["current_version_id"]

    parse_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document_id}/parse-tasks",
        json={"parser_name": "word-parser", "parser_version": "0.1.0"},
    )
    assert parse_response.status_code == 202

    chunks_response = client.get(
        (
            f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document_id}"
            f"/versions/{source_version_id}/chunks"
        )
    )
    assert chunks_response.status_code == 200
    chunks = chunks_response.json()
    assert chunks

    revised_chunks = [
        {
            "chunk_index": item["chunk_index"],
            "page_no": item["page_no"],
            "heading_path": item["heading_path"],
            "content_text": item["content_text"],
            "bbox_json": item["bbox_json"],
            "table_json": item["table_json"],
        }
        for item in chunks
    ]
    revised_chunks[0]["content_text"] = "资格要求"
    revised_chunks[1]["content_text"] = "投标人须提供有效安全生产许可证。"

    revision_response = client.post(
        (
            f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document_id}"
            f"/versions/{source_version_id}/manual-revisions"
        ),
        json={"reason": "人工修正解析分块", "chunks": revised_chunks},
    )
    assert revision_response.status_code == 201
    revision = revision_response.json()
    new_version = revision["new_version"]
    assert new_version["version_no"] == document_payload["current_version"]["version_no"] + 1
    assert new_version["parser_name"] == "manual-editor"
    assert revision["document"]["current_version_id"] == new_version["id"]
    assert "安全生产许可证" in revision["chunks"][1]["content_text"]

    matrix_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/generate",
        json={
            "document_id": document_id,
            "document_version_id": new_version["id"],
            "force": True,
        },
    )
    assert matrix_response.status_code == 202

    items_response = client.get(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items",
        params={"limit": 500},
    )
    assert items_response.status_code == 200
    items = items_response.json()
    revised_item = next(
        item
        for item in items
        if item["source_version_id"] == new_version["id"] and "安全生产许可证" in item["requirement_text"]
    )
    assert revised_item["source_version_id"] == new_version["id"]

    with SessionLocal() as db:
        logs = db.scalars(
            select(AuditLog).where(
                AuditLog.action == "document.manual_revision_published",
                AuditLog.object_id == UUID(new_version["id"]),
            )
        ).all()
        assert logs
