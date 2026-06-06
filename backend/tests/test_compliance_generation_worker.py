from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
import re
import threading
import time
from types import SimpleNamespace
from uuid import UUID, uuid4

import pytest
from alembic import command
from alembic.config import Config
from docx import Document as WordDocument
from fastapi.testclient import TestClient
from sqlalchemy import select

from app.core.config import settings
from app.db.session import SessionLocal
from app.main import app
from app.models import (
    AsyncTask,
    ComplianceItem,
    Document,
    DocumentChunk,
    DocumentExtractionQualityReport,
    DocumentSemanticSection,
    DocumentVersion,
)
from app.parsers.word import ParsedWordChunk, parse_docx_bytes
from app.services.compliance_generation import (
    LLMDocumentSection,
    LLMCoverageReview,
    _page_payload,
    _augment_section_candidates_from_table_guards,
    _augment_section_candidates_from_text_guards,
    _normalize_llm_compliance_payload,
    _normalize_coverage_review_payload,
    _is_boilerplate_semantic_section,
    _is_retriable_llm_response_error,
    _llm_retry_reason,
    _rule_extract,
    _source_quote_matches,
    _source_create_method,
    _split_large_section_plan,
    _source_quote_matches_adjacent_chunk_window,
    _source_quote_text_variants,
    _best_chunk_matching_quote,
    _best_chunk_matching_quote_window,
    _split_chunks_for_llm_retry,
    _unique_chunk_matching_quote,
    ensure_document_section_plan,
    execute_compliance_matrix_generation_task,
)
from app.services.document_parse import execute_document_parse_task
from scripts.seed_dev_data import seed


@pytest.fixture(scope="module", autouse=True)
def prepare_database() -> None:
    backend_dir = Path(__file__).resolve().parents[1]
    config = Config(str(backend_dir / "alembic.ini"))
    config.set_main_option("script_location", str(backend_dir / "migrations"))
    command.upgrade(config, "head")
    seed()


def get_seed_project_and_section(client: TestClient) -> tuple[str, str]:
    projects = client.get("/api/v1/projects").json()
    project = next(item for item in projects if item["name"] == "智慧园区弱电工程投标")
    sections = client.get(f"/api/v1/projects/{project['id']}/sections").json()
    return project["id"], sections[0]["id"]


def build_qualification_docx_bytes() -> bytes:
    document = WordDocument()
    document.add_heading("测试燃气项目（EPC）", level=1)
    document.add_paragraph("1.4 招标范围：建设规模范围内的设计、施工总承包，包括验收及缺陷责任期服务。")
    document.add_paragraph("1.6 保修要求：按国务院2000年279号令相关规定及合同约定进行施工质量保修。")
    document.add_paragraph("1.7 缺陷责任期：24个月。")
    document.add_paragraph("1.8 标段划分：本项目分为一个标段。")
    document.add_paragraph("1.9 技术响应要求：洁净设备应提供产品样本、检测报告和安装调试验收方案。")
    document.add_paragraph("2.资格要求")
    document.add_paragraph("2.1 具有独立法人资格并依法取得企业营业执照，营业执照处于有效期；")
    document.add_paragraph("2.2 须具备市政公用工程施工总承包贰级及以上资质，安全生产许可证处于有效期；")
    document.add_paragraph("2.3 拟任施工项目负责人具备市政公用工程专业二级及以上注册建造师资格，以联合体投标的，拟任工程总承包项目负责人须为联合体牵头人单位人员。")
    document.add_paragraph("2.4 本次招标接受联合体投标。")
    document.add_paragraph("2.8 类似工程业绩要求：")
    document.add_paragraph("？不要求")
    document.add_paragraph("3.资格审查")
    document.add_paragraph("采用资格后审方式")
    document.add_paragraph("4.评标办法")
    document.add_paragraph("本招标项目采用湘建监督[2024]34号文件规定的综合评估法。")
    document.add_paragraph("6.投标文件的递交")
    document.add_paragraph("6.1 电子投标文件递交的截止时间为2025年12月11日9时30分。")
    document.add_paragraph("8.行政监督")
    document.add_paragraph("本次招标项目招标投标监督机构为岳阳市君山区建设工程招投标管理办公室，电话0730-8178006。")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_cleanroom_public_notice_docx_bytes() -> bytes:
    document = WordDocument()
    document.add_heading("净化运维服务托管项目公开招标公告", level=1)
    document.add_paragraph(
        "项目概况\n"
        "净化运维服务托管招标项目的潜在投标人应在登录“邢台市公共资源交易网”"
        "选择“市级交易响应方登录”→“邢台市公共资源交易平台”自行下载招标文件。"
        "获取招标文件，并于2026年06月12日 09点00分（北京时间）前递交投标文件。"
    )
    document.add_paragraph("一、项目基本情况")
    document.add_paragraph("采购需求：序号 标的名称 预算金额（万元） 服务期限 简要技术需求或服务要求1 净化区域运维服务托管项目 180 三年 详见采购需求")
    document.add_paragraph("合同履行期限：三年，合同一年一签。")
    document.add_paragraph("本项目不接受联合体投标。")
    document.add_paragraph("二、申请人的资格要求：")
    document.add_paragraph("1.满足《中华人民共和国政府采购法》第二十二条规定；")
    document.add_paragraph(
        "2.落实政府采购政策需满足的资格要求：1.满足《中华人民共和国政府采购法》第二十二条规定；"
        "2.1本项目专门面向小微企业采购，投标人须为小微企业；"
        "2.2其他落实政府采购政策的资格要求（如有）： 无 ；"
        "2.3通过“信用中国”网站和中国政府采购网查询信用记录，被列入失信被执行人、"
        "重大税收违法案件当事人名单、政府采购严重违法失信行为记录名单的供应商，没有资格参加本项目的采购活动。"
    )
    document.add_paragraph(
        "3.本项目的特定资格要求：（1）供应商须具备建设行政主管部门颁发的"
        "建筑装修装饰工程专业承包贰级及以上资质、电子与智能化工程专业承包贰级及以上资质、"
        "建筑机电安装工程专业承包叁级及以上资质且具备有效的安全生产许可证； "
        "（2）供应商拟派项目经理须具备机电工程专业贰级及以上注册建造师执业资格，"
        "具备有效的安全生产考核(B 类)合格证书，且未担任其他在施建设工程项目的项目经理。"
    )
    document.add_paragraph("三、获取招标文件")
    document.add_paragraph("时间：2026年05月21日至2026年05月27日,每天上午00:00至12:00,下午12:00至23:59")
    document.add_paragraph("地点：登录“邢台市公共资源交易网”自行下载招标文件。")
    document.add_paragraph("四、提交投标文件截止时间、开标时间和地点")
    document.add_paragraph("2026年06月12日 09点00分（北京时间）")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def normalized_requirement_key(text: str) -> str:
    return re.sub(r"[\s，。；;：:（）()、]+", "", text)


def test_source_create_method_uses_short_storage_values() -> None:
    assert (
        _source_create_method(
            {"extraction_provider": "deepseek:deepseek-v4-pro:table_guard"},
            "ai_sectioned",
        )
        == "table_guard"
    )
    assert _source_create_method({"extraction_provider": "deepseek:deepseek-v4-pro"}, "rule") == "ai_sectioned"
    assert _source_create_method({"extraction_provider": "rules"}, "ai_sectioned") == "rule"


def test_source_quote_match_accepts_short_table_values() -> None:
    chunk = SimpleNamespace(
        content_text="是否允许联合体承包 | √不允许。\\n交易保证金 | 无",
        table_json={
            "rows": [
                ["是否允许联合体承包", "√不允许。", "□允许"],
                ["交易保证金", "无"],
            ]
        },
    )

    assert _source_quote_matches(chunk, "不允许")  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, "无")  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, "不允许联合体承包")  # type: ignore[arg-type]
    assert not _source_quote_matches(chunk, "接受联合体")  # type: ignore[arg-type]


def test_source_quote_match_treats_table_separator_as_colon() -> None:
    chunk = SimpleNamespace(
        content_text="交易方式-评审办\n法 | 公开竞标-综合评分法 | 最高报价限价\n(万元) | 388.7963",
        table_json={
            "rows": [
                ["交易方式-评审办法", "公开竞标-综合评分法", "最高报价限价\n(万元)", "388.7963"]
            ]
        },
    )

    assert _source_quote_matches(chunk, "最高报价限价 (万元)：388.7963")  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, "最高报价限价 (万元)\\n388.7963")  # type: ignore[arg-type]


def test_source_quote_match_ignores_checkbox_and_quote_variants() -> None:
    chunk = SimpleNamespace(
        content_text='☑ 下载安装最新版“上城投标工具"(适用综合评分法)',
        table_json=None,
    )

    assert _source_quote_matches(chunk, "下载安装最新版“上城投标工具”(适用综合评分法)")  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, '下载安装最新版“上城投标工具\\"(适用综合评分法)')  # type: ignore[arg-type]


def test_unique_chunk_matching_quote_corrects_adjacent_chunk_reference() -> None:
    wrong_chunk = SimpleNamespace(content_text="交易担保 | 本项目不收取交易保证金", table_json=None)
    right_chunk = SimpleNamespace(content_text="响应文件的质询 | ☑不质询。", table_json=None)

    assert _unique_chunk_matching_quote([wrong_chunk, right_chunk], "不质询。") is right_chunk  # type: ignore[list-item]


def test_best_chunk_matching_quote_prefers_nearest_duplicate() -> None:
    first = SimpleNamespace(chunk_index=12, content_text="本项目不收取交易保证金", table_json=None)
    second = SimpleNamespace(chunk_index=20, content_text="本项目不收取交易保证金", table_json=None)

    assert (
        _best_chunk_matching_quote([first, second], "本项目不收取交易保证金", preferred_chunk_index=10)  # type: ignore[list-item]
        is first
    )


def test_source_quote_match_allows_long_quotes_split_by_page_break() -> None:
    chunk = SimpleNamespace(
        content_text="13 | 评审小组 | 采用综合评分法的，由发包人组建的评审小组进行评审。评审小组由项目申请人、相关科室代表及相关专业人员等不少",
        table_json=None,
    )

    assert _source_quote_matches(
        chunk,
        "采用综合评分法的，由发包人组建的评审小组进行评审。评审小组由项目申请人、相关科室代表及相关专业人员等不少于5人的单数组成，相关专业人员不少于成员总数的三分之二。",
    )  # type: ignore[arg-type]


def test_source_quote_match_ignores_pdf_private_use_spacing() -> None:
    chunk = SimpleNamespace(
        content_text="2.合同价格形式：\ue5e5固定单价合同 \ue5e5。",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "合同价格形式：固定单价合同。")  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, "2.合同价格形式：\x97固定单价合同 \x97。")  # type: ignore[arg-type]


def test_source_quote_match_allows_omitted_middle_phrase() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "风险控制价 311.0370 万元，为防止恶意低价竞争，最高报价限价的 80 ％"
            "作为风险控制价，成交价低于风险控制价的，成交人还须提供差额担保。"
        ),
        table_json=None,
    )

    assert _source_quote_matches(
        chunk,
        "风险控制价 311.0370 万元，最高报价限价的 80 ％作为风险控制价",
    )  # type: ignore[arg-type]


def test_source_quote_match_allows_source_inserted_middle_phrase() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "施工企业的现场监控、现场临时宿舍取暖降温费用，"
            "未根据市政府、市建设行政主管部门颁发的有关文件对于现场监控和现场民工宿舍空调"
            "的设置要求或标准规定落实相应费用的报价的，作废标处理。"
        ),
        table_json=None,
    )

    assert _source_quote_matches(  # type: ignore[arg-type]
        chunk,
        "未根据市政府、市建设行政主管部门颁发的有关文件对于现场监控的设置要求或标准规定"
        "落实相应费用的报价的，作废标处理",
    )


def test_source_quote_match_allows_long_quote_with_exact_prefix_and_summarized_tail() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "10.1 对拟派项目负责人“有在建合同工程”的认定标准："
            "(1)拟派项目负责人在响应截止时间尚有在其他在建合同工程中担任项目负责人的情形为"
            "“有在建合同工程”。"
            "(2)其他工程项目，包括在中华人民共和国境内所有建设工程，不受地域、行业和投资性质的限制。"
            "(3)在建合同工程的时间界定：在建合同工程的开始时间为合同工程成交通知书发出日期。"
            "10.2 存在以下情形的视为“有在建合同工程”："
            "(1)合同协议书尚未签订的，成交通知书中载明的项目负责人。"
        ),
        table_json=None,
    )

    assert _source_quote_matches(  # type: ignore[arg-type]
        chunk,
        "对拟派项目负责人“有在建合同工程”的认定标准："
        "(1)拟派项目负责人在响应截止时间尚有在其他在建合同工程中担任项目负责人的情形为"
        "“有在建合同工程”。"
        "(2)其他工程项目，包括在中华人民共和国境内所有建设工程，不受地域、行业和投资性质的限制。"
        "(3)在建合同工程的时间界定：在建合同工程的开始时间为合同工程成交通知书发出日期。"
        "(4)存在以下情形的视为“有在建合同工程”："
        "(a)合同协议书尚未签订的，成交通知书中载明的项目负责人。",
    )


def test_source_quote_match_allows_short_parallel_clause_omission() -> None:
    chunk = SimpleNamespace(
        content_text="投标文件提出了不能满足竞标文件要求的工程验收、施工工期、保修期等要求的。",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "投标文件提出了不能满足竞标文件要求的施工工期")  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, "不能满足竞标文件要求的施工工期")  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, "不能满足竞标文件要求的保修期")  # type: ignore[arg-type]


def test_source_quote_match_allows_omitted_previous_list_item() -> None:
    chunk = SimpleNamespace(
        content_text="☑综合评分法项目 (1)响应函；(2)法定代表人身份证明；(3)授权委托书；",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "综合评分法项目 (2)法定代表人身份证明；")  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, "综合评分法项目 (3)授权委托书；")  # type: ignore[arg-type]


def test_source_quote_match_allows_selected_value_after_unselected_option() -> None:
    chunk = SimpleNamespace(
        content_text="是否专门面向\n中小企业 | □是。工程全部由中小企业承建。\n☑否。",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "是否专门面向中小企业：☑否。")  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, "是否专门面向中小企业 | ☑否。")  # type: ignore[arg-type]
    assert not _source_quote_matches(chunk, "是否专门面向中小企业：☑是。")  # type: ignore[arg-type]
    assert not _source_quote_matches(chunk, "是否专门面向中小企业 | ☑是。")  # type: ignore[arg-type]


def test_source_quote_match_allows_checked_option_after_ellipsis() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "是否专门面向\n中小企业 | □是。工程全部由符合政策要求的中小企业承建，"
            "潜在承包人应提供中小企业声明函。\n☑否。"
        ),
        table_json={
            "rows": [
                [
                    "是否专门面向\n中小企业",
                    "□是。工程全部由符合政策要求的中小企业承建，潜在承包人应提供中小企业声明函。\n☑否。",
                ]
            ]
        },
    )

    assert _source_quote_matches(chunk, "是否专门面向中小企业：... ☑否。")  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, "是否专门面向中小企业 | ☑否。")  # type: ignore[arg-type]
    assert not _source_quote_matches(chunk, "是否专门面向中小企业：... ☑是。")  # type: ignore[arg-type]


def test_source_quote_match_allows_selected_numbered_option_after_omitted_options() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "关于是否扣留质量保证金的约定：是。"
            "15.3.1 承包人提供质量保证金的方式 "
            "质量保证金采用以下第（3）种方式："
            "（1）质量保证金保函，保证金额为：；"
            "（2）1.5 %的工程款；"
            "（3）其他方式:缴纳形式同投标担保,保证金额为1.5%的工程款。"
        ),
        table_json=None,
    )

    assert _source_quote_matches(  # type: ignore[arg-type]
        chunk,
        "关于是否扣留质量保证金的约定：是。质量保证金采用以下第（3）种方式："
        "其他方式:缴纳形式同投标担保,保证金额为1.5%的工程款。",
    )


def test_source_quote_match_allows_trailing_context_from_list_intro() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "承包人有下列违约情形的，发包人有权单方解除合同："
            "(1)承包人原因超过开工日期【10】日以上未进场施工的；"
            "(2)承包人原因停工超【15】日以上的；"
        ),
        table_json=None,
    )

    assert _source_quote_matches(  # type: ignore[arg-type]
        chunk,
        "承包人原因停工超【15】日以上的，发包人有权单方解除合同",
    )


def test_source_quote_match_allows_omitted_parenthetical_example() -> None:
    chunk = SimpleNamespace(
        content_text="对不合格的材料(如水泥、砂、石、水质、钢筋等)严禁用于本工程。",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "对不合格的材料严禁用于本工程。")  # type: ignore[arg-type]


def test_source_quote_match_allows_omitted_nested_parenthetical_example() -> None:
    chunk = SimpleNamespace(
        content_text="按规定办理农民工工伤保险（按照杭政办函（2007）148 号规定的公式计算）、意外险、工程一切险等保险费用。",
        table_json=None,
    )

    assert _source_quote_matches(  # type: ignore[arg-type]
        chunk,
        "按规定办理农民工工伤保险、意外险、工程一切险等保险费用。",
    )


def test_source_quote_match_ignores_terminal_soft_punctuation_variants() -> None:
    chunk = SimpleNamespace(
        content_text="(7)发包人要求提交的其他资料。",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "(7)发包人要求提交的其他资料；")  # type: ignore[arg-type]


def test_source_quote_match_allows_ascii_terminal_period_for_selected_value() -> None:
    chunk = SimpleNamespace(
        content_text="是否允许分包 | √不允许。 □允许。",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "不允许.")  # type: ignore[arg-type]


def test_source_quote_match_allows_optional_chinese_particle() -> None:
    chunk = SimpleNamespace(
        content_text="发包人供应的材料设备保管费用由承包人承担，承包人已在投标报价中综合考虑。缺陷责任期为24个月。",
        table_json=None,
    )

    assert _source_quote_matches(
        chunk,
        "发包人供应的材料设备的保管费用由承包人承担，承包人已在投标报价中综合考虑",
    )  # type: ignore[arg-type]
    assert _source_quote_matches(chunk, "缺陷责任期24个月")  # type: ignore[arg-type]


def test_source_quote_match_allows_common_delegation_word_order_variant() -> None:
    chunk = SimpleNamespace(
        content_text="1.响应报价应由潜在承包人或受其委托具有相应能力的工程造价咨询人编制。",
        table_json=None,
    )

    assert _source_quote_matches(  # type: ignore[arg-type]
        chunk,
        "响应报价应由潜在承包人或其受委托具有相应能力的工程造价咨询人编制。",
    )


def test_source_quote_match_allows_carried_intro_for_short_list_item() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "信用信息的使用规则：经查询列入失信被执行人名单、重大税收违法案件当事人名单、"
            "严重失信黑名单的潜在承包人将取消成交人资格。"
        ),
        table_json=None,
    )

    assert _source_quote_matches(chunk, "列入严重失信黑名单")  # type: ignore[arg-type]
    assert not _source_quote_matches(chunk, "列入政府采购黑名单")  # type: ignore[arg-type]


def test_source_quote_match_allows_omitted_source_label_word() -> None:
    chunk = SimpleNamespace(
        content_text="工资性工程款比例约定：按相关文件执行。",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "工资性工程款比例按相关文件执行")  # type: ignore[arg-type]


def test_source_quote_match_allows_carried_subject_prefix() -> None:
    chunk = SimpleNamespace(
        content_text="承包人应严格按已确认的施工技术方案实施。无条件服从建设单位对现场总平面的协调。",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "承包人无条件服从建设单位对现场总平面的协调。")  # type: ignore[arg-type]


def test_source_quote_text_variants_align_common_clause_type_confusion() -> None:
    variants = _source_quote_text_variants("按专用条款第13.6.1的规定对施工现场进行清理")

    assert (
        "按通用条款第13.6.1的规定对施工现场进行清理",
        {"专用条款": "通用条款"},
    ) in variants


def test_source_quote_text_variants_align_measure_fee_list() -> None:
    variants = _source_quote_text_variants("文明施工、环境保护、安全施工、临时设施费四项费用的投标报价总额")

    assert (
        "文明施工费、环境保护费、安全施工费和临时设施费四项费用的投标报价总额",
        {
            "文明施工、环境保护、安全施工、临时设施费四项费用": (
                "文明施工费、环境保护费、安全施工费和临时设施费四项费用"
            )
        },
    ) in variants


def test_source_quote_text_variants_align_deduct_score_order() -> None:
    variants = _source_quote_text_variants("投标报价低于风险控制价的，每低于最佳报价1%的扣2分")

    assert (
        "投标报价低于风险控制价的，每低于最佳报价1%的扣分2分",
        {"扣2分": "扣分2分"},
    ) in variants


def test_source_quote_text_variants_align_risk_control_intro() -> None:
    variants = _source_quote_text_variants("风险控制价：最高投标限价的80％作为风险控制价")

    assert (
        "风险控制价：为防止投标人恶意低价竞标，最高投标限价的80％作为风险控制价",
        {"风险控制价：最高投标限价": "风险控制价：为防止投标人恶意低价竞标，最高投标限价"},
    ) in variants


def test_source_quote_text_variants_align_risk_control_intro_after_ellipsis() -> None:
    variants = _source_quote_text_variants("风险控制价：...最高投标限价的80％作为风险控制价")

    assert any("为防止投标人恶意低价竞标" in variant for variant, _ in variants)


def test_source_quote_text_variants_align_in_progress_contract_period() -> None:
    variants = _source_quote_text_variants(
        "在建合同工程的开始时间为合同工程中标通知书发出日期，或者合同签订日期，结束时间为验收合格或合同解除日期。"
    )

    assert any("不通过招标方式的则以合同签订日期为开始时间" in variant for variant, _ in variants)


def test_source_quote_text_variants_align_cross_page_truncated_last_character() -> None:
    variants = _source_quote_text_variants("补充、修改后重新传输递交")

    assert ("补充、修改后重新传输递", {}) in variants


def test_source_quote_text_variants_align_reordered_failed_import_quote() -> None:
    variants = _source_quote_text_variants("无法导入成功的响应文件")

    assert any("电子响应文件无法解密" in variant for variant, _ in variants)


def test_source_quote_text_variants_align_reordered_management_staff_deadline() -> None:
    variants = _source_quote_text_variants(
        "签订合同后 7天内，承包人提交项目管理机构及施工现场管理人员安排报告，并对相应人员的到位率作出承诺"
    )

    assert any("承包人提交项目管理机构及施工现场管理人员安排报告的期限" in variant for variant, _ in variants)


def test_source_quote_text_variants_align_omitted_qualification_list_prefix() -> None:
    variants = _source_quote_text_variants("不得低于投标时的资质")

    assert ("不得低于投标时的职称、资质", {"不得低于投标时的资质": "不得低于投标时的职称、资质"}) in variants


def test_source_quote_text_variants_align_omitted_civilized_site_basis() -> None:
    variants = _source_quote_text_variants("文明施工按杭州市人民政府令第278号文《杭州市建设工程文明施工管理规定》执行")

    assert any("杭建监总[2011]21号" in variant for variant, _ in variants)


def test_source_quote_match_allows_checked_party_alternative() -> None:
    chunk = SimpleNamespace(
        content_text="代理服务费用收取标准为：依据招标代理合同进行收费，由☑发包人／□承包人支付。",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "由☑发包人支付。")  # type: ignore[arg-type]


def test_source_quote_match_allows_full_quote_with_check_marker_and_table_separators() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "6.3 | 评标办法 | √ 技术标打分制的综合评估法："
            "技术标评分（30分）,资信标评\n分（10分）,商务标评分（60分）；"
        ),
        table_json=None,
    )

    assert _source_quote_matches(  # type: ignore[arg-type]
        chunk,
        "6.3 | 评标办法 | √ 技术标打分制的综合评估法：技术标评分（30分）,资信标评分（10分）,商务标评分（60分）；",
    )


def test_source_quote_match_allows_cost_rate_label_carried_from_sentence() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "规费以“人工费+机械费”为取费基数，不得低于标准费率的30%，"
            "即单独装饰工程不得低于8.38%；安装工程不得低于9.19%。"
        ),
        table_json=None,
    )

    assert _source_quote_matches(chunk, "单独装饰工程规费不得低于8.38%")  # type: ignore[arg-type]


def test_source_quote_match_allows_adjacent_chunk_page_split() -> None:
    previous_chunk = SimpleNamespace(
        chunk_index=105,
        content_text="承包人须为本工程办理建设\n- 86 -",
        table_json=None,
    )
    current_chunk = SimpleNamespace(
        chunk_index=106,
        content_text="工程一切险并支付保险费用，工程费用也已包含在总报价内。",
        table_json=None,
    )

    assert _source_quote_matches_adjacent_chunk_window(  # type: ignore[arg-type]
        [previous_chunk, current_chunk],
        current_chunk,
        "承包人须为本工程办理建设工程一切险并支付保险费用",
    )


def test_source_quote_match_allows_page_number_chunk_between_split() -> None:
    previous_chunk = SimpleNamespace(
        chunk_index=24,
        content_text="⑩工程量清单报价与工、料、机报价及对应的报价分析不相符的，",
        table_json=None,
    )
    page_number_chunk = SimpleNamespace(
        chunk_index=25,
        content_text="13",
        table_json=None,
    )
    current_chunk = SimpleNamespace(
        chunk_index=26,
        content_text="或与拟建工程的施工方案明显不匹配的；",
        table_json=None,
    )

    assert _source_quote_matches_adjacent_chunk_window(  # type: ignore[arg-type]
        [previous_chunk, page_number_chunk, current_chunk],
        previous_chunk,
        "工程量清单报价与工、料、机报价及对应的报价分析不相符的，或与拟建工程的施工方案明显不匹配的",
    )


def test_best_chunk_matching_quote_window_relocates_misindexed_cross_page_quote() -> None:
    unrelated_chunk = SimpleNamespace(
        chunk_index=14,
        content_text="5.2 | 开标程序 | 招标人代表应使用招标CA数字证书进行解密。",
        table_json=None,
    )
    start_chunk = SimpleNamespace(
        chunk_index=16,
        content_text="7.4 | 履约担保及工程款支付担保 | 中标价低于风险控制价的，中标人还须向招标人提供风险控制价",
        table_json=None,
    )
    page_number_chunk = SimpleNamespace(
        chunk_index=17,
        content_text="9",
        table_json=None,
    )
    continuation_chunk = SimpleNamespace(
        chunk_index=18,
        content_text="和中标价的差额担保，作为履约担保的一部分。",
        table_json=None,
    )

    corrected = _best_chunk_matching_quote_window(  # type: ignore[arg-type]
        [unrelated_chunk, start_chunk, page_number_chunk, continuation_chunk],
        "中标价低于风险控制价的，中标人还须向招标人提供风险控制价 和中标价的差额担保，作为履约担保的一部分。",
        preferred_chunk_index=14,
    )

    assert corrected is start_chunk


def test_source_quote_match_allows_parenthetical_and_punctuation_in_adjacent_split() -> None:
    previous_chunk = SimpleNamespace(
        chunk_index=24,
        content_text=(
            "10.3在建项目的项目负责人 (包括在工程总承包项目中担任施工负责人或总负责人) 办理"
            "更换后，递交响应文件时需提供的资料：(1)项目业主同意更换的证明。\n- 13 -"
        ),
        table_json=None,
    )
    current_chunk = SimpleNamespace(
        chunk_index=25,
        content_text=(
            "(2)原项目负责人在建项目信息有备案在建设主管部门的，应提供建设主管部门同意更换"
            "的证明或网上变更信息扫描件。(3)未提供有效证明材料的视作无变更，事后补充的变更证明材料均不予认可。"
        ),
        table_json=None,
    )

    assert _source_quote_matches_adjacent_chunk_window(  # type: ignore[arg-type]
        [previous_chunk, current_chunk],
        current_chunk,
        (
            "在建项目的项目负责人办理更换后，递交响应文件时需提供的资料："
            "(1)项目业主同意更换的证明；(2)原项目负责人在建项目信息有备案在建设主管部门的，"
            "应提供建设主管部门同意更换的证明或网上变更信息扫描件；"
            "(3)未提供有效证明材料的视作无变更，事后补充的变更证明材料均不予认可。"
        ),
    )


def test_source_quote_match_allows_ellipsis_omission() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "3.我公司将建立严格的质量保证体系，确保项目质量，加强施工质量验收制度，"
            "绝不违章施工，严格执行国家、省、市现行的有关施工验收规范和质量检验评定标准，"
            "确保项目质量验收合格或一次性验收合格。"
        ),
        table_json=None,
    )

    assert _source_quote_matches(  # type: ignore[arg-type]
        chunk,
        "3.我公司将建立严格的质量保证体系，确保项目质量，...确保项目质量验收合格或一次性验收合格。",
    )


def test_source_quote_match_allows_ellipsis_with_parenthetical_source_context() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "企业管理费包含安责险，安责险计取按照相关文件，"
            "最低费率不得低于相应弹性费率下限值（即单独装饰工程 11.37%；安装工程 16.29%）乘以20%。"
        ),
        table_json=None,
    )

    assert _source_quote_matches(  # type: ignore[arg-type]
        chunk,
        "企业管理费包含安责险，...最低费率不得低于相应弹性费率下限值乘以20%",
    )


def test_source_quote_match_allows_ellipsis_with_soft_punctuation() -> None:
    chunk = SimpleNamespace(
        content_text=(
            "5.3.2承包人提前通知监理人隐蔽工程检查的期限的约定：双方约定隐蔽工程和中间验收部位的程序："
            "承包人自检合格后，在隐蔽工程和中间验收48小时前通知监理工程师和发包人代表参加。"
        ),
        table_json=None,
    )

    assert _source_quote_matches(  # type: ignore[arg-type]
        chunk,
        "承包人提前通知监理人隐蔽工程检查的期限的约定：……在隐蔽工程和中间验收48小时前通知。",
    )


def test_source_quote_match_allows_ellipsis_with_short_tail_fragment() -> None:
    chunk = SimpleNamespace(
        content_text="投标文件提出了不能满足竞标文件要求的工程验收、施工工期、保修期等要求的。",
        table_json=None,
    )

    assert _source_quote_matches(chunk, "不能满足竞标文件要求的...保修期")  # type: ignore[arg-type]


def test_coverage_review_payload_accepts_description_alias() -> None:
    review = LLMCoverageReview.model_validate(
        _normalize_coverage_review_payload(
            {
                "status": "blocked",
                "issues": [
                    {
                        "severity": "high",
                        "description": "漏抽关键合规项：交易方式。",
                        "suggestion": "交易方式：公开竞标",
                    }
                ],
            }
        )
    )

    assert review.issues[0].message == "漏抽关键合规项：交易方式。"
    assert review.issues[0].suggested_requirement == "交易方式：公开竞标"


def test_coverage_review_payload_accepts_item_issue_shape() -> None:
    review = LLMCoverageReview.model_validate(
        _normalize_coverage_review_payload(
            {
                "status": "blocked",
                "issues": [
                    {
                        "type": "missing",
                        "item": "交易保证金：无",
                    }
                ],
            }
        )
    )

    assert review.issues[0].message == "missing: 交易保证金：无"
    assert review.issues[0].suggested_requirement == "交易保证金：无"


def test_compliance_payload_accepts_candidates_alias() -> None:
    payload = _normalize_llm_compliance_payload(
        {"candidates": [{"chunk_index": 1, "quote": "工期要求：180日历天"}]}
    )

    assert payload["items"][0]["source_chunk_index"] == 1
    assert payload["items"][0]["source_quote"] == "工期要求：180日历天"


def test_compliance_payload_accepts_matrix_alias() -> None:
    payload = _normalize_llm_compliance_payload(
        {"matrix": [{"chunk_index": 4, "requirement": "须具备建筑工程施工总承包三级资质。"}]}
    )

    assert payload["items"][0]["source_chunk_index"] == 4
    assert payload["items"][0]["requirement_text"] == "须具备建筑工程施工总承包三级资质。"


def test_compliance_payload_accepts_matrix_items_alias() -> None:
    payload = _normalize_llm_compliance_payload(
        {
            "matrix_items": [
                {
                    "source_chunk_index": 12,
                    "text": "响应文件有效期为90日历天。",
                }
            ]
        }
    )

    assert payload["items"][0]["source_chunk_index"] == 12
    assert payload["items"][0]["requirement_text"] == "响应文件有效期为90日历天。"


def test_compliance_payload_accepts_compliance_entries_alias() -> None:
    payload = _normalize_llm_compliance_payload(
        {"compliance_entries": [{"source_index": 7, "content": "交易保证金：无"}]}
    )

    assert payload["items"][0]["source_chunk_index"] == 7
    assert payload["items"][0]["requirement_text"] == "交易保证金：无"


def test_compliance_payload_accepts_compliance_matrix_alias() -> None:
    payload = _normalize_llm_compliance_payload(
        {"compliance_matrix": [{"chunk_id": 9, "text": "采用综合评分法。"}]}
    )

    assert payload["items"][0]["source_chunk_index"] == 9
    assert payload["items"][0]["requirement_text"] == "采用综合评分法。"


def test_compliance_payload_accepts_section_keyed_items() -> None:
    payload = _normalize_llm_compliance_payload(
        {"section9": [{"source_index": 99, "content": "质量保证金为1.5%的工程款。"}]}
    )

    assert payload["items"][0]["source_chunk_index"] == 99
    assert payload["items"][0]["requirement_text"] == "质量保证金为1.5%的工程款。"


def test_boilerplate_semantic_sections_are_identified() -> None:
    section = SimpleNamespace(
        section_type="other",
        title="封面及说明",
        evidence="第1页为封面，第2页为交易文件示范文本说明，第3页为目录。",
    )
    announcement = SimpleNamespace(
        section_type="announcement",
        title="第一章 交易公告",
        evidence="第4页出现交易公告。",
    )

    assert _is_boilerplate_semantic_section(section)  # type: ignore[arg-type]
    assert not _is_boilerplate_semantic_section(announcement)  # type: ignore[arg-type]


def test_table_guard_augments_multi_field_rows_and_deadline() -> None:
    chunk = SimpleNamespace(
        chunk_index=5,
        heading_path="PDF 第 4 页",
        table_json={
            "rows": [
                ["交易方式-评审办法", "公开竞标-综合评分法", "最高报价限价\n(万元)", "388.7963", ""],
                ["10", "响应文件有效期", "90 日历天 (从响应截止之日起算)。"],
                ["交易时间\n(响应文件递交截止时间)", "2024年12月30日14时0分"],
                ["是否专门面向\n中小企业", "□是。工程全部由符合政策要求的中小企业承建。\n☑否。"],
            ]
        },
    )

    candidates = _augment_section_candidates_from_table_guards(  # type: ignore[arg-type]
        [chunk],
        [],
        extraction_provider="deepseek:deepseek-v4-pro:table_guard",
    )
    texts = [candidate.requirement_text for candidate in candidates]

    assert "最高报价限价 (万元)：388.7963" in texts
    assert "响应文件有效期：90 日历天 (从响应截止之日起算)。" in texts
    assert "交易时间 (响应文件递交截止时间)：2024年12月30日14时0分" in texts
    assert "是否专门面向 中小企业：否" in texts
    assert any(
        candidate.item_type == "deadline" and candidate.risk_level == "high"
        for candidate in candidates
    )


def test_table_guard_keeps_short_value_fields_as_separate_items() -> None:
    chunk = SimpleNamespace(
        chunk_index=5,
        heading_path="PDF 第 4 页",
        table_json={"rows": [["工期要求或服务期限(日历天)", "180", "交易保证金", "无"]]},
    )
    existing = SimpleNamespace(
        requirement_text="工期要求或服务期限(日历天)：180；交易保证金；无",
        normalized_requirement="existing:combined",
    )

    candidates = _augment_section_candidates_from_table_guards(  # type: ignore[arg-type]
        [chunk],
        [existing],
        extraction_provider="deepseek:deepseek-v4-pro:table_guard",
    )

    assert any(candidate.requirement_text == "交易保证金：无" for candidate in candidates)


def test_text_guard_augments_commitment_basic_ability_requirements() -> None:
    chunk = SimpleNamespace(
        chunk_index=152,
        heading_path="PDF 第 118 页",
        content_text=(
            "格式3：\n诚信承诺书\n(一)基本能力方面\n"
            "1.具有独立承担民事责任的能力；\n"
            "2.具有良好的商业信誉和健全的财务会计制度；\n"
            "3.具有履行合同所必需的设备和专业技术能力；\n"
            "4.有依法缴纳税收和社会保障资金的良好记录；\n"
            "5.参加活动前三年内，在经营活动中没有重大违法记录；\n"
            "6.具有法律、行政法规规定的其他条件。"
        ),
    )

    candidates = _augment_section_candidates_from_text_guards(  # type: ignore[arg-type]
        [chunk],
        [],
        extraction_provider="deepseek:deepseek-v4-pro:text_guard",
    )
    texts = [candidate.requirement_text for candidate in candidates]

    assert "诚信承诺书：参加活动前三年内，在经营活动中没有重大违法记录" in texts
    assert len(candidates) == 6
    assert all(candidate.item_type == "qualification" for candidate in candidates)
    assert all(candidate.risk_level == "high" for candidate in candidates)
    assert all(candidate.explanation_json["source_quote"] for candidate in candidates)


def test_text_guard_skips_existing_commitment_requirement() -> None:
    chunk = SimpleNamespace(
        chunk_index=152,
        heading_path="PDF 第 118 页",
        content_text=(
            "诚信承诺书 (一)基本能力方面 "
            "5.参加活动前三年内，在经营活动中没有重大违法记录；"
        ),
    )
    existing = SimpleNamespace(
        requirement_text="参加活动前三年内，在经营活动中没有重大违法记录；",
        normalized_requirement="existing:major-violation",
    )

    candidates = _augment_section_candidates_from_text_guards(  # type: ignore[arg-type]
        [chunk],
        [existing],
        extraction_provider="deepseek:deepseek-v4-pro:text_guard",
    )

    assert candidates == [existing]


def test_text_guard_augments_epc_joint_venture_leader_requirement() -> None:
    chunks = [
        SimpleNamespace(
            chunk_index=26,
            heading_path="2.资格要求/2.2 须同时具备以下资质，资质证书处于有效期内：",
            content_text=(
                "2.3 拟任工程总承包项目负责人具备 市政公用工程二级及以上注册建造师或市政公用工程"
                "相关专业高级及以上技术职称，（可同时兼任本项目“施工项目负责人”，应同时满足施工项目负责人的"
                "资格要求），且不得同时在其他建筑工程项目中任职 （不得设置类似业绩要求；以联合体投标的，"
                "拟任工程总承包项目负责人须为联合体牵头人单位人员）；"
            ),
        ),
        SimpleNamespace(
            chunk_index=30,
            heading_path="2.资格要求/2.6 本次招标接受联合体投标",
            content_text="2.6 本次招标接受联合体投标，联合体投标的相关要求见投标人须知前附表",
        ),
        SimpleNamespace(
            chunk_index=20,
            heading_path="10.4 在建合同工程的认定及变更证明",
            content_text=(
                "1.对项目负责人（包括在工程总承包项目中担任施工负责人）“有在建合同工程”的认定标准："
                "（1）拟派项目负责人在投标截止时间尚有在其他在建合同工程中担任项目负责人的情形为“有在建合同工程”。"
                "（3）在建合同工程的时间界定：在建合同工程的开始时间为合同工程中标通知书发出日期，"
                "或者不通过招标方式的则以合同签订日期为开始时间，结束时间为该合同工程验收合格或合同解除日期）。"
            ),
        ),
    ]

    candidates = _augment_section_candidates_from_text_guards(  # type: ignore[arg-type]
        chunks,
        [],
        extraction_provider="deepseek:deepseek-v4-pro:text_guard",
    )

    assert any(
        candidate.requirement_text == "以联合体投标的，拟任工程总承包项目负责人须为联合体牵头人单位人员。"
        for candidate in candidates
    )
    assert any(candidate.requirement_text == "本项目接受联合体投标。" for candidate in candidates)
    assert any("认定为有在建合同工程" in candidate.requirement_text for candidate in candidates)
    assert any("在建合同工程时间界定" in candidate.requirement_text for candidate in candidates)
    assert all(candidate.item_type == "qualification" for candidate in candidates)


def test_text_guard_augments_contract_and_response_form_fields() -> None:
    chunks = [
        SimpleNamespace(
            chunk_index=58,
            heading_path="PDF 第 44 页",
            content_text=(
                "3.2.1 项目经理： 姓 名： ； 身份证号： ； "
                "建造师执业资格等级： ；"
            ),
        ),
        SimpleNamespace(
            chunk_index=59,
            heading_path="PDF 第 45 页",
            content_text=(
                "建造师注册证书号： ； 建造师执业印章号： ； "
                "安全生产考核合格证书号： ；"
            ),
        ),
        SimpleNamespace(
            chunk_index=119,
            heading_path="PDF 第 99 页",
            content_text="三、缺陷责任期 工程缺陷责任期为24个月，缺陷责任期自工程通过竣工验收之日起计算。",
        ),
        SimpleNamespace(
            chunk_index=13,
            heading_path="君山区城区燃气管网改造项目（EPC）/1.项目概况",
            content_text="1.3 工期要求： 270 ？天（日历日，下同）□月□年；",
        ),
        SimpleNamespace(
            chunk_index=18,
            heading_path="PDF 第 10 页",
            content_text=(
                "1.4.2 交易公告规定接受联合体承包的，联合体除应符合本章第 1.4.1 项和交易须知前 "
                "附表的要求外，还应遵守以下规定："
                "(1)联合体各方应按交易文件提供的格式签订联合体协议书，明确联合体牵头人和各方权 利义务；"
                "(2)由同一专业的单位组成的联合体，按照资质等级较低的单位确定资质等级；"
                "(3)联合体各方不得再以自己名义单独或参加其他联合体在同一项目中响应。"
            ),
        ),
        SimpleNamespace(
            chunk_index=46,
            heading_path="PDF 第 22 页",
            content_text="16、本工程暂列金额：200000 元",
        ),
        SimpleNamespace(
            chunk_index=23,
            heading_path="四、提交投标文件截止时间、开标时间和地点",
            content_text="地点： 网上开标，投标人应及时登录邢台市公共资源交易平台在线参与开标 。",
        ),
        SimpleNamespace(
            chunk_index=30,
            heading_path="PDF 第 22 页",
            content_text=(
                "第四章 评审办法 本项目采用综合评分法。"
                "2.改变竞标文件提供的工程量清单（含分部分项工程及措施项目、其他项目清单项目编 码、项目名称、计量单位、工程数量、项目特征描述）。"
                "3.改变竞标文件规定的暂定内容的；"
                "4.经评标委员会认定投标人的投标报价低于成本价的；"
                "5.投标人拒绝按评标委员会要求提供报价分析说明和证明材料的；"
                "6.工程量清单报价与工、料、机报价及对应的报价分析不相符的，或与拟建工程的施工 组织设计及施工方案明显不匹配的，经评标委员会书面质询，投标人不能说明理由或评标委员会认定其理由不成立的；"
            ),
        ),
        SimpleNamespace(
            chunk_index=40,
            heading_path="PDF 第 27 页",
            content_text="2.合同价格形式：\ue5e5固定单价合同 \ue5e5。",
        ),
        SimpleNamespace(
            chunk_index=89,
            heading_path="PDF 第 73 页",
            content_text="12. 合同价格、计量与支付 12.1 合同价格形式 1、单价合同。综合单价包含的风险范围：材料价格波动。",
        ),
        SimpleNamespace(
            chunk_index=92,
            heading_path="PDF 第 76 页",
            content_text="项目决算审计完成后支付至结算价的98.5%，剩余1.5%留作质量保证金（缴纳方式同投标担保）。质量保证金的返还详见工程质量保修书。",
        ),
        SimpleNamespace(
            chunk_index=99,
            heading_path="PDF 第 83 页",
            content_text=(
                "15.3.1 承包人提供质量保证金的方式 质量保证金采用以下第（3）种方式："
                "（1）质量保证金保函，保证金额为：；（2）1.5 %的工程款；"
                "（3）其他方式:缴纳形式同投标担保,保证金额为1.5%的工程款。"
            ),
        ),
        SimpleNamespace(
            chunk_index=102,
            heading_path="PDF 第 86 页",
            content_text="质量未达到招标文件要求或无法通过竣工验收的，承包人负责返修，直至达到要求，并扣罚全部质量履约保证金，且发包人保留采取其他制约措施的权力。",
        ),
        SimpleNamespace(
            chunk_index=77,
            heading_path="PDF 第 61 页",
            content_text=(
                "承包人提交详细施工组织设计的期限的约定：合同签订后14天内提供详细施工组织设计"
                "（施工方案）和进度计划 。承包人逾期未提交的，每逾期一日，发包人有权扣罚1000元工期"
                " 履约保证金，因此影响工程施工的，责任由承包人承担。"
            ),
        ),
        SimpleNamespace(
            chunk_index=79,
            heading_path="PDF 第 63 页",
            content_text="当上述措施仍无效或未采取时，以承包人违约论处，发包人可终止本合同，将承包人清退出场，罚没全部履约保证金，由此引起的工期延误责任及损失由承包人负责承担。如因承包人自身原因造成工期严重滞后，又无明显改进措施的，发包人有权要求承包人无条件退场，扣罚全部工期履约保证金，并承担由此产生的一切损失。",
        ),
        SimpleNamespace(
            chunk_index=111,
            heading_path="PDF 第 95 页",
            content_text="21.28总承包服务（配合）内容：甲方另行分包的如自来水、电力、通信等工程不列入本次招标范围，不计取总包管理费，但不能免除总包对专业承包工程应履行的配合义务，总包单位需预留水电管网等接口、管线打孔、补洞、保护费用由投标人承担，由投标人自行考虑相关费用在投标报价中，中标后招标人不再另行支付相关费用。",
        ),
        SimpleNamespace(
            chunk_index=135,
            heading_path="PDF 第 107 页",
            content_text=(
                "附件：10 支付担保格式 一、保证的范围及保证金额 "
                "3.我方保证的金额是主合同约定的工程款的2%，数额最高不超过人民币元（大写： ）。"
            ),
        ),
        SimpleNamespace(
            chunk_index=148,
            heading_path="PDF 第 114 页",
            content_text=(
                "格式1： 响应函 我方愿以人民币 (大 写) ： ，RMB：¥ 元，"
                "(大小写不一致的以大写金额为准) 的报价并按交易文件要求承包。"
                "本项目拟派项目负责人姓名： ，身份证 号： 。"
                "工期(服务期) 个日历天。"
                "一旦我方成为成交人，我方保证按交易文件要求向贵方递交经贵方认可的履约担保。"
                "在我方报价低于风险控制价的情况下，我方将按照规定以保函的形式提交成交价与风险控制 价之差额。"
            ),
        ),
        SimpleNamespace(
            chunk_index=150,
            heading_path="PDF 第 116 页",
            content_text=(
                "授权委托书 本授权委托书声明：我 (姓名) 系 (单位名称) 的法定代表人，"
                "现授权委 托 (姓名) 在 年 月 日至 年 月 日(代理时限)为我公司的代理人，"
                "以本公司的名义参加 (项目名称) 的交易活动。代理人在代理时间内参加交易活动过程中"
                "所签署的一切文件和处理与之相关的一切事务，本人均予以承认。代理人无权转委托。"
                "附 代理人身份证正面复印件粘贴处 代理人身份证背面复印件粘贴处"
            ),
        ),
        SimpleNamespace(
            chunk_index=102,
            heading_path="PDF 第 66 页",
            content_text="附件8：\n履约担保\n1. 担保金额人民币（大写） 元（¥ ）。2. 担保有效期自合同生效之日起。",
        ),
        SimpleNamespace(
            chunk_index=103,
            heading_path="PDF 第 67 页",
            content_text="附件9：\n预付款担保\n1. 担保金额人民币（大写） 元（¥ ）。2. 担保有效期自预付款支付给承包人起生效。",
        ),
        SimpleNamespace(
            chunk_index=104,
            heading_path="PDF 第 68 页",
            content_text="附件10:\n支付担保\n一、保证的范围及保证金额。二、保证的方式及保证期间 1. 我方保证的方式为：连带责任保证。",
        ),
    ]

    candidates = _augment_section_candidates_from_text_guards(  # type: ignore[arg-type]
        chunks,
        [],
        extraction_provider="deepseek:deepseek-v4-pro:text_guard",
    )
    texts = [candidate.requirement_text for candidate in candidates]

    assert "项目经理信息：须填写姓名、身份证号、建造师执业资格等级。" in texts
    assert "项目经理信息：须填写建造师注册证书号、建造师执业印章号、安全生产考核合格证书号。" in texts
    assert any("缺陷责任期：工程缺陷责任期为24个月" in text for text in texts)
    assert "工期要求：270天（日历日，下同）。" in texts
    assert "联合体承包：接受联合体的，须签订联合体协议书；同一专业按资质等级较低单位确定资质等级；联合体各方不得再以自己名义或参加其他联合体响应。" in texts
    assert "本工程暂列金额：200000元。" in texts
    assert "开标地点：网上开标，投标人应及时登录邢台市公共资源交易平台在线参与开标。" in texts
    assert "评审办法：本项目采用综合评分法。" in texts
    assert any("商务标废标条件：不得改变工程量清单或暂定内容" in text for text in texts)
    assert "合同价格形式：\ue5e5固定单价合同 \ue5e5。" in texts
    assert "合同价格形式：单价合同。" in texts
    assert "质量保证金：剩余1.5%留作质量保证金，返还详见工程质量保修书。" in texts
    assert "质量保证金：采用缴纳形式同投标担保，保证金额为1.5%的工程款。" in texts
    assert "履约保证金：质量未达到要求或无法通过竣工验收的，扣罚全部质量履约保证金。" in texts
    assert "工期履约保证金：承包人逾期未提交施工组织设计的，每逾期一日扣罚1000元。" in texts
    assert "履约保证金：承包人现场管理等未履行承诺且整改无效或未采取措施时，发包人可终止合同、清退出场并罚没全部履约保证金。" in texts
    assert "工期履约保证金：承包人自身原因造成工期严重滞后且无明显改进措施时，发包人可要求退场并扣罚全部工期履约保证金。" in texts
    assert "招标范围：甲方另行分包工程不列入本次招标范围，不计取总包管理费；总包仍需履行配合义务并承担相关预留、打孔补洞和保护费用。" in texts
    assert any("支付担保金额：我方保证的金额是主合同约定的工程款的2%" in text for text in texts)
    assert "响应函：须填写本项目拟派项目负责人姓名、身份证号。" in texts
    assert "响应函：须填写报价金额（人民币大写和RMB小写），大小写不一致以大写为准。" in texts
    assert "响应函：须填写工期(服务期)日历天。" in texts
    assert "响应函：成交后须递交履约担保；报价低于风险控制价的，须以保函形式提交成交价与风险控制价之差额。" in texts
    assert "授权委托书：须提供授权委托书，填写代理人信息和代理时限，并附代理人身份证复印件。" in texts
    assert "附件8履约担保格式：须填写担保金额、担保有效期，并由担保人盖章及法定代表人或委托代理人签字。" in texts
    assert "附件9预付款担保格式：须填写担保金额、担保有效期，并由担保人盖章及法定代表人或委托代理人签字。" in texts
    assert "附件10支付担保格式：须填写保证范围、保证金额、保证方式、保证期间及代偿安排。" in texts
    assert all(candidate.explanation_json["text_guard"] for candidate in candidates)


def test_section_plan_payload_compresses_long_pdf_pages() -> None:
    chunks = []
    long_body = (
        "本页包含招标文件正文、通用说明、合同条件、技术要求和响应文件格式。"
        "投标人应仔细阅读全部条款并按交易文件要求响应。"
    ) * 45
    for page_no in range(1, 127):
        rows = []
        if page_no == 4:
            rows = [
                ["项目名称", "明珠公寓老旧小区综合改造项目建安工程"],
                ["发包范围", "项目位于杭州市九堡明珠公寓，主要施工内容包括建筑立面及屋面整治提升等。"],
                ["资质要求", "具备施工总承包建筑工程施工总承包三级资质。"],
            ]
        chunks.append(
            SimpleNamespace(
                chunk_index=page_no,
                page_no=page_no,
                heading_path=f"PDF 第 {page_no} 页",
                content_text=(f"第{page_no}页\n第一章 交易公告\n" if page_no == 4 else f"第{page_no}页\n") + long_body,
                table_json={"rows": rows} if rows else None,
            )
        )

    payload = _page_payload(chunks)  # type: ignore[arg-type]
    payload_json = json.dumps(payload, ensure_ascii=False)
    raw_char_count = sum(len(chunk.content_text) for chunk in chunks)
    page_four = next(page for page in payload if page["page_no"] == 4)

    assert len(payload) == 126
    assert len(payload_json) < 80_000
    assert len(payload_json) < raw_char_count * 0.2
    assert all(len(page["text"]) <= 120 for page in payload)
    assert all(len(row) <= 120 for page in payload for row in page["table_rows"])
    assert "第一章 交易公告" in page_four["text"]
    assert any("发包范围" in row for row in page_four["table_rows"])


def test_section_plan_payload_uses_chunk_index_for_unpaginated_chunks() -> None:
    chunks = [
        SimpleNamespace(
            chunk_index=index,
            page_no=None,
            heading_path=f"Word chunk {index}",
            content_text=f"第 {index} 段资格要求文本。",
            table_json=None,
        )
        for index in range(1, 4)
    ]

    payload = _page_payload(chunks)  # type: ignore[arg-type]

    assert [page["page_no"] for page in payload] == [1, 2, 3]
    assert [page["chunk_indexes"] for page in payload] == [[1], [2], [3]]


def test_section_plan_splits_large_sections_for_extraction_budget() -> None:
    section = LLMDocumentSection(
        section_index=1,
        title="第五章 合同条款及格式",
        section_type="contract",
        start_page=1,
        end_page=42,
        confidence_score=0.8,
        evidence="模型识别为连续合同章节。",
    )
    chunks = [
        SimpleNamespace(
            chunk_index=page_no,
            page_no=page_no,
            heading_path=f"PDF 第 {page_no} 页",
            content_text=("合同条款正文，承包人应按要求履约。" * 90),
            table_json=None,
        )
        for page_no in range(1, 43)
    ]

    split_sections = _split_large_section_plan([section], chunks)  # type: ignore[arg-type]

    assert len(split_sections) > 1
    assert split_sections[0].start_page == 1
    assert split_sections[-1].end_page == 42
    assert [section.section_index for section in split_sections] == list(range(1, len(split_sections) + 1))
    assert all(
        current.end_page + 1 == next_section.start_page
        for current, next_section in zip(split_sections, split_sections[1:], strict=False)
    )
    assert all("长章节按页码拆分" in section.evidence for section in split_sections)
    assert all(section.title.startswith("第五章 合同条款及格式") for section in split_sections)


def test_document_section_plan_reuses_existing_when_force_replan_returns_empty_json(monkeypatch) -> None:
    settings.run_tasks_inline = False
    monkeypatch.setattr(settings, "llm_provider", "mock")
    monkeypatch.setattr(settings, "llm_api_key", "")

    section_plan_calls = 0

    def fake_chat_completion(db, **kwargs):  # noqa: ANN001, ARG001
        nonlocal section_plan_calls
        assert kwargs["prompt_version"] == "document_section_plan@1.1.0"
        section_plan_calls += 1
        if section_plan_calls == 1:
            user_content = kwargs["messages"][-1]["content"]
            pages = json.loads(user_content.rsplit("pages:\n", 1)[1])
            page_numbers = sorted(page["page_no"] for page in pages)
            content = {
                "sections": [
                    {
                        "section_index": 1,
                        "title": "资格要求测试章节",
                        "section_type": "announcement",
                        "start_page": page_numbers[0],
                        "end_page": page_numbers[-1],
                        "confidence_score": 0.9,
                        "evidence": "首次规划成功。",
                    }
                ]
            }
            response_content = json.dumps(content, ensure_ascii=False)
        else:
            response_content = ""
        return SimpleNamespace(
            content=response_content,
            provider="fake",
            model_name="unit-test",
            log_id=None,
            usage={},
        )

    monkeypatch.setattr("app.services.compliance_generation.chat_completion", fake_chat_completion)
    from app.worker import run_compliance_matrix_generation_task, run_document_parse_task

    monkeypatch.setattr(run_document_parse_task, "delay", lambda task_id: None)
    monkeypatch.setattr(run_compliance_matrix_generation_task, "delay", lambda task_id: None)

    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    filename = f"章节规划复用-{uuid4().hex}.docx"

    upload_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/upload",
        data={"doc_type": "tender", "title": "章节规划复用测试"},
        files={
            "file": (
                filename,
                build_qualification_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    document_payload = upload_response.json()

    parse_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document_payload['id']}/parse-tasks",
        json={"parser_name": "word-parser", "parser_version": "0.1.0"},
    )
    assert parse_response.status_code == 202
    with SessionLocal() as db:
        parse_result = execute_document_parse_task(db, UUID(parse_response.json()["task"]["id"]))
        assert parse_result["status"] == "succeeded"

    generate_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/generate",
        json={"document_version_id": document_payload["current_version_id"], "force": True},
    )
    assert generate_response.status_code == 202

    with SessionLocal() as db:
        task = db.get(AsyncTask, UUID(generate_response.json()["id"]))
        document = db.get(Document, UUID(document_payload["id"]))
        version = db.get(DocumentVersion, UUID(document_payload["current_version_id"]))
        assert task is not None
        assert document is not None
        assert version is not None
        chunks = db.scalars(
            select(DocumentChunk)
            .where(DocumentChunk.document_version_id == version.id)
            .order_by(DocumentChunk.chunk_index)
        ).all()

        initial_sections = ensure_document_section_plan(db, task, document, version, list(chunks))
        reused_sections = ensure_document_section_plan(
            db,
            task,
            document,
            version,
            list(chunks),
            force=True,
        )

        assert section_plan_calls == 2
        assert [section.id for section in reused_sections] == [section.id for section in initial_sections]
        db.refresh(task)
        assert task.output_json is not None
        assert task.output_json["progress_stage"] == "document_section_plan_reused"
        assert task.output_json["fallback_reason"] == "LLM_JSON_PARSE_FAILED"


def test_llm_retry_batches_overlap_adjacent_chunks_for_page_break_quotes() -> None:
    chunks = [
        SimpleNamespace(
            chunk_index=index,
            content_text="x" * 1000,
        )
        for index in range(1, 6)
    ]

    batches = _split_chunks_for_llm_retry(chunks)  # type: ignore[arg-type]

    assert len(batches) >= 2
    batch_indexes = [{chunk.chunk_index for chunk in batch} for batch in batches]
    for index in range(1, 5):
        assert any({index, index + 1}.issubset(indexes) for indexes in batch_indexes)


def test_llm_retry_recursive_batches_can_shrink_without_overlap() -> None:
    chunks = [
        SimpleNamespace(
            chunk_index=index,
            content_text="x" * 900,
        )
        for index in range(1, 3)
    ]

    batches = _split_chunks_for_llm_retry(chunks, include_adjacent_overlap=False)  # type: ignore[arg-type]

    assert [{chunk.chunk_index for chunk in batch} for batch in batches] == [{1}, {2}]


def test_llm_retry_treats_json_parse_failures_as_retriable() -> None:
    error = json.JSONDecodeError("Expecting value", "", 0)

    assert _is_retriable_llm_response_error(error)
    assert _llm_retry_reason(error) == "模型返回的 JSON 无法解析"


def test_cleanroom_notice_rule_fallback_atomizes_and_deduplicates_requirements() -> None:
    chunks = parse_docx_bytes(build_cleanroom_public_notice_docx_bytes())
    heading_by_text = {chunk.content_text: chunk.heading_path for chunk in chunks}
    assert heading_by_text["三、获取招标文件"].endswith("三、获取招标文件")
    assert "特定资格要求" not in (heading_by_text["三、获取招标文件"] or "")

    candidates = _rule_extract(chunks)  # type: ignore[arg-type]
    texts = [item.requirement_text for item in candidates]
    joined = "\n".join(texts)
    normalized_texts = [normalized_requirement_key(text) for text in texts]

    assert "净化运维服务托管项目公开招标公告" not in texts
    assert "3.本项目的特定资格要求" not in joined.replace("\n", "")
    assert not any("三、获取招标文件" in text and "特定资格要求" in text for text in texts)
    assert not any("2.2其他落实政府采购政策" in text for text in texts)
    assert not any("邢台市公共资源交易网" in text and item.item_type != "deadline" for item, text in zip(candidates, texts))
    assert len(normalized_texts) == len(set(normalized_texts))
    assert all(len(text) <= 180 for text in texts)

    assert any("建筑装修装饰工程专业承包贰级及以上资质" in text for text in texts)
    electronic_qualification = next(
        item for item in candidates if "电子与智能化工程专业承包贰级及以上资质" in item.requirement_text
    )
    assert electronic_qualification.item_type == "qualification"
    assert any("建筑机电安装工程专业承包叁级及以上资质" in text for text in texts)
    assert any("有效的安全生产许可证" in text for text in texts)
    assert any("拟派项目经理" in text and "注册建造师" in text for text in texts)
    assert any("安全生产考核" in text and "B 类" in text for text in texts)
    assert any("未担任其他在施建设工程项目" in text for text in texts)
    assert any("小微企业" in text for text in texts)
    assert any("不接受联合体投标" in text for text in texts)
    assert any("信用中国" in text and "没有资格参加" in text for text in texts)

    submit_deadlines = [
        item for item in candidates if item.item_type == "deadline" and "2026年06月12日" in item.requirement_text
    ]
    assert len(submit_deadlines) == 1

    qualification_count = sum(1 for item in candidates if item.item_type == "qualification")
    assert qualification_count >= 8


def test_rule_fallback_splits_chinese_numbered_notice_before_classification() -> None:
    source_text = (
        "说明 一、为规范上城区限额以下小额公共资源交易活动，保障市场主体的合法权益，"
        "根据有关规定，制定杭州市上城区小额交易文件示范文本2024年V1.0版。 "
        "二、本小额交易文件示范文本适用于杭州市上城区行政区域内采用竞标方式的"
        "限额以下小额工程建设项目，包括施工、勘察、设计、监理、全过程工程咨询项目。 "
        "三、有下划线和交易须知前附表空白部分，由发包人根据发包项目实际情况和国家有关"
        "法律法规规定进行填写，文字应采用斜体字；确实不需要填写内容的，用斜线标示。"
        "除可选择部分、下划线空白部分外，其他文字不得改动。 "
        "四、发包人委托代理机构组织交易的，代理服务费用原则上由发包人支付，可计入工程前期费用。 "
        "五、发包人应根据项目的实际情况合理选定评审办法，鼓励选用合理低价法和双随机信用法。 "
        "六、交易文件、工程量清单、图纸等技术资料的获取为网上免费下载形式。"
        "发包人应在交易公告及交易文件中明确获取方式及时 间，获取截止时间同响应文件提交截止时间一致。 "
        "七、采用综合评分法的工程项目，自交易文件发出之日起至响应文件提交截止时间"
        "不得少于10日（日历天），其他工程项目不得少于5个工作日。"
    )
    chunks = [ParsedWordChunk(chunk_index=1, heading_path="PDF 第 2 页", content_text=source_text)]

    candidates = _rule_extract(chunks)  # type: ignore[arg-type]
    texts = [item.requirement_text for item in candidates]

    assert len(candidates) > 1
    assert not any("一、" in text and "二、" in text and "七、" in text for text in texts)
    deadline_items = [item for item in candidates if item.item_type == "deadline"]
    assert any("不得少于10日" in item.requirement_text for item in deadline_items)
    early_marker_items = [
        item
        for item in candidates
        if any(marker in item.requirement_text for marker in ("一、", "二、", "三、", "四、", "五、"))
    ]
    assert not early_marker_items


def test_rule_fallback_uses_pdf_table_rows_and_filters_template_placeholders() -> None:
    chunk = SimpleNamespace(
        chunk_index=1,
        heading_path="PDF 第 4 页",
        content_text="PDF 表格整页文本包含资质要求、联合体、分包和截止时间。",
        table_json={
            "rows": [
                [
                    "资质要求",
                    "1.企业： (1)具备 具备 施工总承包 建筑工程施工总承包 三级 资质； "
                    "(2)自 / 年 / 月 / 日以来承接过 / 业绩。 "
                    "(3)发包人需要增加的、符合法律法规的其他要求。 "
                    "2.拟派项目负责人： 具有注册在潜在承包人单位的 建筑工程二级及以上 资格。 "
                    "☑如在响应截止日存在在其他任何在建合同工程的，不得以拟派项目负 责人的身份参加本次交易。 "
                    "□项目经理需配合办理发包备案手续，具体要求： / 。 "
                    "发包人需要增加的、符合法律法规的其他要求。 说明：多选，如有以上要求请标注“√”，若无标注“□”",
                ],
                [
                    "是否允许联合体 承包",
                    "不允许。 □允许，联合体承包的应提供联合体协议书。 说明：单选，选中项标注“√”，非选中项标注“□”",
                ],
                [
                    "是否允许分包",
                    "不允许。 \uf0a3允许。分包的工程内容： 分包企业应符合规定的资格要求。 说明：单选，选中项标注“√”，非选中项标注“□”",
                ],
                ["交易时间 (响应文件递交截 止时间)", "2024年12月30日14时0分"],
            ]
        },
    )

    candidates = _rule_extract([chunk])  # type: ignore[arg-type]
    texts = [item.requirement_text for item in candidates]
    compact_texts = [re.sub(r"\s+", "", text) for text in texts]

    assert any("施工总承包建筑工程施工总承包三级资质" in text for text in compact_texts)
    assert any("拟派项目负责人具有注册在潜在承包人单位的建筑工程二级及以上资格" in text for text in compact_texts)
    assert any("不得以拟派项目负责人的身份参加本次交易" in text for text in compact_texts)
    assert "是否允许联合体 承包：不允许。" in texts
    assert "是否允许分包：不允许。" in texts
    assert any(item.item_type == "deadline" and "2024年12月30日14时0分" in item.requirement_text for item in candidates)
    assert not any("/ 年 / 月 / 日" in text or "发包人需要增加" in text for text in texts)
    assert not any("□允许" in text or "\uf0a3允许" in text or text.startswith("PDF 第") for text in texts)


def test_junshan_rules_contextualize_and_deduplicate_low_information_fragments() -> None:
    chunks = [
        SimpleNamespace(
            chunk_index=1,
            heading_path="君山区城区燃气管网改造项目（EPC）/类似工程业绩要求",
            content_text="不要求",
            table_json=None,
        ),
        SimpleNamespace(
            chunk_index=2,
            heading_path="君山区城区燃气管网改造项目（EPC）/获取招标文件",
            content_text="2025 年 11 月 28 日17时",
            table_json=None,
        ),
        SimpleNamespace(
            chunk_index=3,
            heading_path="君山区城区燃气管网改造项目（EPC）/资格要求",
            content_text="以联合体投标的，拟任工程总承包项目负责人须为联合体牵头人单位人员",
            table_json=None,
        ),
        SimpleNamespace(
            chunk_index=4,
            heading_path="君山区城区燃气管网改造项目（EPC）/资格要求",
            content_text="以联合体投标的，拟任工程总承包项目负责人须为联合体牵头人单位人员。",
            table_json=None,
        ),
        SimpleNamespace(
            chunk_index=5,
            heading_path="君山区城区燃气管网改造项目（EPC）/资格要求",
            content_text="2.6 本次招标接受联合体投标，联合体投标的相关要求见投标人须知前附表",
            table_json=None,
        ),
        SimpleNamespace(
            chunk_index=6,
            heading_path="君山区城区燃气管网改造项目（EPC）/资格要求",
            content_text="本项目接受联合体投标。",
            table_json=None,
        ),
        SimpleNamespace(
            chunk_index=7,
            heading_path="君山区城区燃气管网改造项目（EPC）/项目概况",
            content_text="1.3 工期要求： 270 ？天（日历日，下同）□月□年；",
            table_json=None,
        ),
        SimpleNamespace(
            chunk_index=8,
            heading_path="君山区城区燃气管网改造项目（EPC）/项目概况",
            content_text="工期要求：270天（日历日，下同）。",
            table_json=None,
        ),
        SimpleNamespace(
            chunk_index=9,
            heading_path="君山区城区燃气管网改造项目（EPC）/项目概况",
            content_text="建设规模范围内的设计、施工总承包",
            table_json=None,
        ),
    ]

    candidates = _rule_extract(chunks)  # type: ignore[arg-type]
    texts = [candidate.requirement_text for candidate in candidates]

    similar_performance = next(item for item in candidates if item.requirement_text.startswith("类似工程业绩要求"))
    assert similar_performance.requirement_text == "类似工程业绩要求：不要求。"
    assert similar_performance.risk_level == "low"
    assert similar_performance.is_mandatory is False
    assert "2025 年 11 月 28 日17时" not in texts
    assert "建设规模范围内的设计、施工总承包" not in texts
    assert sum("联合体牵头人单位人员" in text for text in texts) == 1
    assert sum("接受联合体投标" in text for text in texts) == 1
    assert sum(text.startswith("工期要求") or "工期要求：" in text for text in texts) == 1


def test_rule_fallback_splits_pdf_numbered_clauses_without_long_prefix() -> None:
    chunk = SimpleNamespace(
        chunk_index=1,
        heading_path="PDF 第 46 页",
        content_text=(
            "发包人有权中止合同并罚没全部项目班子人员到位履约保证金，"
            "由此引起的工程损失费用由承包人承担。如因其项目经理履职不到位"
            "造成发包人损失的，承包人还应当另行予以赔偿。 "
            "3.2.3 承包人擅自更换项目经理的违约责任：继任项目经理不得低于"
            "原项目经理或技术负责人的注册执业资格、管理经验。 "
            "3.2.4 承包人无正当理由拒绝更换项目经理的违约责任：发包人有权"
            "要求承包人撤换工作不负责任的项目经理，承包人不得拒绝和拖延。"
        ),
        table_json=None,
    )

    candidates = _rule_extract([chunk])  # type: ignore[arg-type]
    texts = [item.requirement_text for item in candidates]

    assert any(text.startswith("3.2.3 承包人擅自更换项目经理") for text in texts)
    assert any(text.startswith("3.2.4 承包人无正当理由拒绝更换项目经理") for text in texts)
    assert not any("发包人有权中止合同" in text and "3.2.3" in text for text in texts)


def test_rule_fallback_splits_pdf_circled_and_chinese_parenthesized_markers() -> None:
    chunk = SimpleNamespace(
        chunk_index=1,
        heading_path="PDF 第 8 页",
        content_text="PDF 表格整页文本包含开标程序和电子交易说明。",
        table_json={
            "rows": [
                [
                    "5.2",
                    "开标程序 043607 （一）至开标时间，招标人代表应使用招标CA数字证书"
                    "对上传的电子加密标书进行解密。 （二）投标文件解密完成后，投标单位、"
                    "项目负责人、投标报价、工期及其他内容将在电子交易平台公开显示。"
                    "（三）招标人在开标会现场随机抽取 K1--权重系数，范围0.4～0.6；"
                    "K2--浮动系数，范围0.85～0.95。",
                ],
                [
                    "电子交易的说明",
                    "（2）交易前准备 ①注册账号：招必得平台注册填报企业信息并完成企业认证。 "
                    "②申领CA数字证书。 ③下载制作工具：☑下载安装最新版上城投标工具。",
                ],
                [
                    "响应文件的组成",
                    "最低价法项目 (1)响应函；(2)法定代表人身份证明；"
                    "☑综合评分法项目 (1)响应函；(2)法定代表人身份证明；"
                    "(3)授权委托书；说明：单选，选中项标注“√”，非选中项标注“□”",
                ],
            ]
        },
    )

    candidates = _rule_extract([chunk])  # type: ignore[arg-type]
    texts = [item.requirement_text for item in candidates]

    assert any("至开标时间" in text and "解密" in text for text in texts)
    assert any("投标文件解密完成后" in text for text in texts)
    assert any("K1--权重系数" in text for text in texts)
    assert any("K2--浮动系数" in text for text in texts)
    assert any("申领CA数字证书" in text for text in texts)
    assert any(text.startswith("响应文件的组成：综合评分法项目") for text in texts)
    assert not any("043607" in text for text in texts)
    assert not any("最低价法项目" in text and "综合评分法项目" in text for text in texts)
    assert all(len(text) < 260 for text in texts)


def test_compliance_matrix_generation_creates_items_from_word_chunks(monkeypatch) -> None:
    settings.run_tasks_inline = False
    monkeypatch.setattr(settings, "llm_provider", "mock")
    monkeypatch.setattr(settings, "llm_api_key", "")

    def fake_chat_completion(db, **kwargs):  # noqa: ANN001, ARG001
        prompt_version = kwargs["prompt_version"]
        if prompt_version == "document_section_plan@1.1.0":
            user_content = kwargs["messages"][-1]["content"]
            pages = json.loads(user_content.rsplit("pages:\n", 1)[1])
            end_page = max(page["page_no"] for page in pages)
            content = {
                "sections": [
                    {
                        "section_index": 1,
                        "title": "测试燃气项目招标文件",
                        "section_type": "announcement",
                        "start_page": 1,
                        "end_page": end_page,
                        "confidence_score": 0.92,
                        "evidence": "Word 解析文本无真实分页，按 chunk 顺序作为章节规划范围。",
                    }
                ]
            }
        elif prompt_version == "section_coverage_review@1.1.0":
            content = {"status": "passed", "issues": []}
        elif prompt_version == "compliance_extract_by_section@1.1.0":
            user_content = kwargs["messages"][-1]["content"]
            chunks = json.loads(user_content.rsplit("chunks:\n", 1)[1])

            def chunk_for(needle: str) -> dict | None:
                for chunk in chunks:
                    text = chunk.get("text") or ""
                    rows = "\n".join(chunk.get("table_rows") or [])
                    if needle in text or needle in rows:
                        return chunk
                return None

            def item(
                needle: str,
                requirement_text: str,
                item_type: str,
                *,
                risk_level: str = "high",
                is_mandatory: bool = True,
                quote: str | None = None,
            ) -> dict | None:
                chunk = chunk_for(quote or needle)
                if chunk is None:
                    return None
                return {
                    "source_chunk_index": chunk["chunk_index"],
                    "item_type": item_type,
                    "requirement_text": requirement_text,
                    "normalized_requirement": requirement_text,
                    "response_suggestion": "按招标文件要求准备响应材料。",
                    "risk_level": risk_level,
                    "is_mandatory": is_mandatory,
                    "classification_reason": "单元测试模拟模型分类。",
                    "split_reason": "按最小可审核要求拆分。",
                    "source_quote": quote or needle,
                    "review_hint": "核对来源原文。",
                    "needs_human_review": False,
                    "confidence_score": 0.88,
                }

            content = {
                "items": [
                    candidate
                    for candidate in [
                        item("招标范围", "1.4 招标范围：建设规模范围内的设计、施工总承包，包括验收及缺陷责任期服务。", "mandatory_response"),
                        item("保修要求", "1.6 保修要求：按国务院2000年279号令相关规定及合同约定进行施工质量保修。", "mandatory_response"),
                        item("缺陷责任期", "1.7 缺陷责任期：24个月。", "mandatory_response"),
                        item("标段划分", "1.8 标段划分：本项目分为一个标段。", "qualification"),
                        item("洁净设备", "1.9 技术响应要求：洁净设备应提供产品样本、检测报告和安装调试验收方案。", "technical_response"),
                        item("营业执照", "具有独立法人资格并依法取得企业营业执照，营业执照处于有效期。", "qualification"),
                        item("市政公用工程施工总承包贰级", "须具备市政公用工程施工总承包贰级及以上资质。", "qualification"),
                        item("安全生产许可证", "安全生产许可证处于有效期。", "qualification"),
                        item("项目负责人具备", "拟任施工项目负责人具备市政公用工程专业二级及以上注册建造师资格。", "qualification"),
                        item("联合体牵头人单位人员", "以联合体投标的，拟任工程总承包项目负责人须为联合体牵头人单位人员。", "qualification"),
                        item("接受联合体投标", "本次招标接受联合体投标。", "qualification"),
                        item("类似工程业绩要求", "类似工程业绩要求：不要求。", "qualification", risk_level="low", is_mandatory=False),
                        item("采用资格后审方式", "采用资格后审方式", "qualification"),
                        item("综合评估法", "本招标项目采用湘建监督[2024]34号文件规定的综合评估法。", "scoring"),
                        item("2025年12月11日9时30分", "电子投标文件递交的截止时间为2025年12月11日9时30分。", "deadline"),
                        item("招标投标监督机构", "本次招标项目招标投标监督机构为岳阳市君山区建设工程招投标管理办公室。", "reference_info", risk_level="low", is_mandatory=False),
                    ]
                    if candidate is not None
                ]
            }
        else:
            raise AssertionError(prompt_version)
        return SimpleNamespace(
            content=json.dumps(content, ensure_ascii=False),
            provider="fake",
            model_name="unit-test",
            log_id=None,
            usage={},
        )

    monkeypatch.setattr("app.services.compliance_generation.chat_completion", fake_chat_completion)
    from app.worker import run_compliance_matrix_generation_task, run_document_parse_task

    monkeypatch.setattr(run_document_parse_task, "delay", lambda task_id: None)
    monkeypatch.setattr(run_compliance_matrix_generation_task, "delay", lambda task_id: None)

    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    filename = f"资格要求-{uuid4().hex}.docx"

    upload_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/upload",
        data={"doc_type": "tender", "title": "资格要求测试"},
        files={
            "file": (
                filename,
                build_qualification_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    document = upload_response.json()

    parse_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document['id']}/parse-tasks",
        json={"parser_name": "word-parser", "parser_version": "0.1.0"},
    )
    assert parse_response.status_code == 202
    with SessionLocal() as db:
        parse_result = execute_document_parse_task(db, UUID(parse_response.json()["task"]["id"]))
        assert parse_result["status"] == "succeeded"

    generate_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/generate",
        json={"document_version_id": document["current_version_id"]},
    )
    assert generate_response.status_code == 202
    generation_task_payload = generate_response.json()

    with SessionLocal() as db:
        result = execute_compliance_matrix_generation_task(db, UUID(generation_task_payload["id"]))
        assert result["status"] == "succeeded"
        assert result["provider"] == "fake:unit-test"
        assert result["created_count"] >= 5

        task = db.get(AsyncTask, UUID(generation_task_payload["id"]))
        assert task is not None
        assert task.status == "succeeded"
        assert task.output_json is not None
        assert task.output_json["created_count"] >= 5
        assert task.output_json["section_count"] >= 1

        sections = db.scalars(
            select(DocumentSemanticSection)
            .where(
                DocumentSemanticSection.document_version_id == UUID(document["current_version_id"])
            )
            .order_by(DocumentSemanticSection.section_index)
        ).all()
        assert len(sections) == task.output_json["section_count"]
        assert sections[0].start_page == 1
        assert sections[-1].end_page >= sections[0].end_page
        quality_report = db.scalar(
            select(DocumentExtractionQualityReport)
            .where(DocumentExtractionQualityReport.document_version_id == UUID(document["current_version_id"]))
            .order_by(DocumentExtractionQualityReport.created_at.desc())
        )
        assert quality_report is not None
        assert quality_report.status == "passed"

        items = db.scalars(
            select(ComplianceItem).where(
                ComplianceItem.source_version_id == UUID(document["current_version_id"]),
                ComplianceItem.deleted_at.is_(None),
            )
        ).all()
        texts = "\n".join(item.requirement_text for item in items)
        assert "营业执照" in texts
        assert "市政公用工程施工总承包贰级" in texts
        assert "联合体投标" in texts
        assert "联合体牵头人单位人员" in texts
        assert "保修要求" in texts
        assert "缺陷责任期" in texts
        assert "洁净设备" in texts
        assert "采用资格后审方式" in texts
        assert "综合评估法" in texts
        assert any(item.item_type == "technical_response" for item in items)
        assert any(item.item_type == "deadline" for item in items)
        assert any(item.item_type == "scoring" for item in items)
        assert any(item.item_type == "reference_info" for item in items)
        assert all(not item.is_batch_confirm_allowed for item in items if item.risk_level == "high")
        high_risk_item = next(item for item in items if item.risk_level == "high")
        assert high_risk_item.explanation_json is not None
        assert high_risk_item.explanation_json["rule_reason"]
        assert high_risk_item.explanation_json["risk_reason"]
        assert high_risk_item.explanation_json["batch_confirm_reason"] == (
            "高风险项不允许批量确认，必须逐条核验来源和响应证据。"
        )

        repeated = execute_compliance_matrix_generation_task(db, UUID(generation_task_payload["id"]))
        assert repeated["status"] == "succeeded"
        assert repeated["updated_count"] >= 5

    generated_api_items = []
    for offset in range(0, 5_000, 500):
        matrix_response = client.get(
            f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items",
            params={"limit": 500, "offset": offset},
        )
        assert matrix_response.status_code == 200
        page_items = matrix_response.json()
        generated_api_items = [
            item
            for item in page_items
            if item["source_version_id"] == document["current_version_id"]
        ]
        if generated_api_items or len(page_items) < 500:
            break
    assert generated_api_items
    api_high_risk = next(item for item in generated_api_items if item["risk_level"] == "high")
    assert api_high_risk["source_content_text"]
    assert api_high_risk["source_chunk_index"] is not None
    assert api_high_risk["rule_explanation"]["risk_reason"]
    assert api_high_risk["rule_explanation"]["batch_confirm_reason"] == (
        "高风险项不允许批量确认，必须逐条核验来源和响应证据。"
    )

    blocked_bulk_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/bulk-confirm",
        json={"item_ids": [api_high_risk["id"]], "reason": "测试高风险项不可批量确认"},
    )
    assert blocked_bulk_response.status_code == 409


def test_compliance_matrix_generation_uses_fork_join_for_semantic_sections(monkeypatch) -> None:
    settings.run_tasks_inline = False
    monkeypatch.setattr(settings, "llm_provider", "mock")
    monkeypatch.setattr(settings, "llm_api_key", "")
    monkeypatch.setattr(settings, "matrix_fork_join_enabled", True)
    monkeypatch.setattr(settings, "matrix_fork_join_min_sections", 4)
    monkeypatch.setattr(settings, "matrix_fork_join_max_workers", 4)

    active_extract_calls = 0
    max_active_extract_calls = 0
    counter_lock = threading.Lock()
    progress_updates: list[dict[str, object]] = []

    from app.services import compliance_generation as compliance_generation_module

    real_update_matrix_task_progress = compliance_generation_module._update_matrix_task_progress

    def recording_update_matrix_task_progress(db, task, **kwargs):  # noqa: ANN001
        progress_updates.append(kwargs)
        return real_update_matrix_task_progress(db, task, **kwargs)

    monkeypatch.setattr(compliance_generation_module, "FORK_JOIN_PROGRESS_HEARTBEAT_SECONDS", 0.02)
    monkeypatch.setattr(
        compliance_generation_module,
        "_update_matrix_task_progress",
        recording_update_matrix_task_progress,
    )

    def fake_chat_completion(db, **kwargs):  # noqa: ANN001, ARG001
        nonlocal active_extract_calls, max_active_extract_calls
        prompt_version = kwargs["prompt_version"]
        if prompt_version == "document_section_plan@1.1.0":
            user_content = kwargs["messages"][-1]["content"]
            pages = json.loads(user_content.rsplit("pages:\n", 1)[1])
            page_numbers = sorted(page["page_no"] for page in pages)
            start_page = page_numbers[0]
            end_page = page_numbers[-1]
            span = max(1, (end_page - start_page + 1) // 4)
            ranges: list[tuple[int, int]] = []
            cursor = start_page
            for index in range(4):
                section_start = cursor
                section_end = end_page if index == 3 else min(end_page, section_start + span - 1)
                ranges.append((section_start, section_end))
                cursor = section_end + 1
            content = {
                "sections": [
                    {
                        "section_index": index,
                        "title": f"并发章节{index}",
                        "section_type": "announcement",
                        "start_page": section_start,
                        "end_page": section_end,
                        "confidence_score": 0.9,
                        "evidence": "并发测试章节。",
                    }
                    for index, (section_start, section_end) in enumerate(ranges, start=1)
                ]
            }
        elif prompt_version == "compliance_extract_by_section@1.1.0":
            with counter_lock:
                active_extract_calls += 1
                max_active_extract_calls = max(max_active_extract_calls, active_extract_calls)
            try:
                time.sleep(0.12)
                user_content = kwargs["messages"][-1]["content"]
                chunks = json.loads(user_content.rsplit("chunks:\n", 1)[1])
                chunk = next((item for item in chunks if (item.get("text") or "").strip()), chunks[0])
                source_text = " ".join((chunk.get("text") or "").split())
                source_quote = source_text[:80]
                content = {
                    "items": [
                        {
                            "source_chunk_index": chunk["chunk_index"],
                            "item_type": "mandatory_response",
                            "requirement_text": f"并发章节 {chunk['chunk_index']} 合规要求：{source_quote}",
                            "normalized_requirement": f"fork_join_requirement_{chunk['chunk_index']}",
                            "response_suggestion": "按并发测试来源准备响应。",
                            "risk_level": "medium",
                            "is_mandatory": True,
                            "source_quote": source_quote,
                            "confidence_score": 0.9,
                        }
                    ]
                }
            finally:
                with counter_lock:
                    active_extract_calls -= 1
        elif prompt_version == "section_coverage_review@1.1.0":
            content = {"status": "passed", "issues": []}
        else:
            raise AssertionError(prompt_version)
        return SimpleNamespace(
            content=json.dumps(content, ensure_ascii=False),
            provider="fake",
            model_name="unit-test",
            log_id=None,
            usage={},
        )

    monkeypatch.setattr("app.services.compliance_generation.chat_completion", fake_chat_completion)
    from app.worker import run_compliance_matrix_generation_task, run_document_parse_task

    monkeypatch.setattr(run_document_parse_task, "delay", lambda task_id: None)
    monkeypatch.setattr(run_compliance_matrix_generation_task, "delay", lambda task_id: None)

    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    filename = f"fork-join-{uuid4().hex}.docx"
    upload_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/upload",
        data={"doc_type": "tender", "title": "Fork Join 测试"},
        files={
            "file": (
                filename,
                build_qualification_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    document = upload_response.json()
    parse_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document['id']}/parse-tasks",
        json={"parser_name": "word-parser", "parser_version": "0.1.0"},
    )
    with SessionLocal() as db:
        assert execute_document_parse_task(db, UUID(parse_response.json()["task"]["id"]))["status"] == "succeeded"

    generate_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/generate",
        json={"document_version_id": document["current_version_id"]},
    )
    with SessionLocal() as db:
        result = execute_compliance_matrix_generation_task(db, UUID(generate_response.json()["id"]))
        assert result["status"] == "succeeded"
        task = db.get(AsyncTask, UUID(generate_response.json()["id"]))
        assert task is not None
        assert task.output_json is not None
        assert task.output_json["execution_mode"] == "fork_join"
        assert task.output_json["fork_join_max_workers"] == 4
        report = db.scalar(
            select(DocumentExtractionQualityReport)
            .where(DocumentExtractionQualityReport.document_version_id == UUID(document["current_version_id"]))
            .order_by(DocumentExtractionQualityReport.created_at.desc())
        )
        assert report is not None
        assert report.summary_json["execution_mode"] == "fork_join"
        assert report.summary_json["fork_join_max_workers"] == 4

    assert max_active_extract_calls >= 2
    assert any(
        (update.get("extra") or {}).get("fork_join_heartbeat")
        for update in progress_updates
        if update.get("stage") == "section_fork_join"
    )


def test_matrix_generation_blocks_when_model_is_not_configured(monkeypatch) -> None:
    settings.run_tasks_inline = False
    monkeypatch.setattr(settings, "llm_provider", "mock")
    monkeypatch.setattr(settings, "llm_api_key", "")

    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    filename = f"未配置模型-{uuid4().hex}.docx"

    upload_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/upload",
        data={"doc_type": "tender", "title": "未配置模型测试"},
        files={
            "file": (
                filename,
                build_qualification_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    document = upload_response.json()
    parse_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document['id']}/parse-tasks",
        json={"parser_name": "word-parser", "parser_version": "0.1.0"},
    )
    assert parse_response.status_code == 202
    with SessionLocal() as db:
        assert execute_document_parse_task(db, UUID(parse_response.json()["task"]["id"]))["status"] == "succeeded"

    generate_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/generate",
        json={"document_version_id": document["current_version_id"]},
    )
    assert generate_response.status_code == 202

    with SessionLocal() as db:
        result = execute_compliance_matrix_generation_task(db, UUID(generate_response.json()["id"]))
        assert result["status"] == "failed"
        task = db.get(AsyncTask, UUID(generate_response.json()["id"]))
        assert task is not None
        assert task.error_code == "LLM_NOT_CONFIGURED"
        items = db.scalars(
            select(ComplianceItem).where(
                ComplianceItem.source_version_id == UUID(document["current_version_id"]),
                ComplianceItem.deleted_at.is_(None),
            )
        ).all()
        assert items == []
        report = db.scalar(
            select(DocumentExtractionQualityReport)
            .where(DocumentExtractionQualityReport.document_version_id == UUID(document["current_version_id"]))
            .order_by(DocumentExtractionQualityReport.created_at.desc())
        )
        assert report is not None
        assert report.status == "blocked"
        assert report.issues_json[0]["code"] == "LLM_NOT_CONFIGURED"


def test_matrix_generation_blocks_when_source_quote_cannot_link_back(monkeypatch) -> None:
    settings.run_tasks_inline = False

    def fake_chat_completion(db, **kwargs):  # noqa: ANN001, ARG001
        prompt_version = kwargs["prompt_version"]
        if prompt_version == "document_section_plan@1.1.0":
            user_content = kwargs["messages"][-1]["content"]
            pages = json.loads(user_content.rsplit("pages:\n", 1)[1])
            end_page = max(page["page_no"] for page in pages)
            content = {
                "sections": [
                    {
                        "section_index": 1,
                        "title": "资格要求",
                        "section_type": "announcement",
                        "start_page": 1,
                        "end_page": end_page,
                        "confidence_score": 0.9,
                        "evidence": "测试段落。",
                    }
                ]
            }
        elif prompt_version == "compliance_extract_by_section@1.1.0":
            user_content = kwargs["messages"][-1]["content"]
            chunks = json.loads(user_content.rsplit("chunks:\n", 1)[1])
            chunk = next((item for item in chunks if "营业执照" in item["text"]), chunks[0])
            content = {
                "items": [
                    {
                        "source_chunk_index": chunk["chunk_index"],
                        "item_type": "qualification",
                        "requirement_text": "投标人须具有企业营业执照。",
                        "risk_level": "high",
                        "is_mandatory": True,
                        "source_quote": "这段摘录并不存在于来源 chunk",
                        "confidence_score": 0.9,
                    }
                ]
            }
        else:
            content = {"status": "passed", "issues": []}
        return SimpleNamespace(
            content=json.dumps(content, ensure_ascii=False),
            provider="fake",
            model_name="unit-test",
            log_id=None,
            usage={},
        )

    monkeypatch.setattr("app.services.compliance_generation.chat_completion", fake_chat_completion)
    client = TestClient(app)
    project_id, section_id = get_seed_project_and_section(client)
    filename = f"来源失败-{uuid4().hex}.docx"

    upload_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/upload",
        data={"doc_type": "tender", "title": "来源回链失败测试"},
        files={
            "file": (
                filename,
                build_qualification_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    document = upload_response.json()
    parse_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/documents/{document['id']}/parse-tasks",
        json={"parser_name": "word-parser", "parser_version": "0.1.0"},
    )
    with SessionLocal() as db:
        assert execute_document_parse_task(db, UUID(parse_response.json()["task"]["id"]))["status"] == "succeeded"

    generate_response = client.post(
        f"/api/v1/projects/{project_id}/sections/{section_id}/compliance-items/generate",
        json={"document_version_id": document["current_version_id"]},
    )
    with SessionLocal() as db:
        result = execute_compliance_matrix_generation_task(db, UUID(generate_response.json()["id"]))
        assert result["status"] == "failed"
        task = db.get(AsyncTask, UUID(generate_response.json()["id"]))
        assert task is not None
        assert task.error_code == "SOURCE_QUOTE_NOT_FOUND"
        assert db.scalars(
            select(ComplianceItem).where(
                ComplianceItem.source_version_id == UUID(document["current_version_id"]),
                ComplianceItem.deleted_at.is_(None),
            )
        ).all() == []
        report = db.scalar(
            select(DocumentExtractionQualityReport)
            .where(DocumentExtractionQualityReport.document_version_id == UUID(document["current_version_id"]))
            .order_by(DocumentExtractionQualityReport.created_at.desc())
        )
        assert report is not None
        assert report.status == "blocked"
        assert report.issues_json[0]["code"] == "SOURCE_QUOTE_NOT_FOUND"
