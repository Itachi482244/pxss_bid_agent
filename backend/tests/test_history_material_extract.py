from __future__ import annotations

import json
from datetime import date
from io import BytesIO
from pathlib import Path
from uuid import UUID, uuid4

import fitz
import pytest
from alembic import command
from alembic.config import Config
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import func, select

from app.core.config import settings
from app.db.session import SessionLocal
from app.main import app
from app.models import AsyncTask, EnterpriseMaterial, EnterpriseMaterialChunk, SourcePageImage, Tenant, User
from app.services.history_material_extract import ExtractedMaterialDraft
from app.services.history_material_extract import SourceTextBlock
from app.services.history_material_extract import create_pending_materials_from_extraction
from app.services.history_material_extract import execute_history_material_extract_task
from app.services.llm_gateway import LLMGatewayError, LLMResult
from app.services.history_material_extract import HistoryMaterialExtraction
from app.services.history_material_extract import parse_history_file_to_blocks
from app.services.ocr import OcrResult
from app.services.source_page_images import render_compressed_pdf_page_image
from scripts.seed_dev_data import seed


@pytest.fixture(scope="module", autouse=True)
def prepare_database() -> None:
    backend_dir = Path(__file__).resolve().parents[1]
    config = Config(str(backend_dir / "alembic.ini"))
    config.set_main_option("script_location", str(backend_dir / "migrations"))
    command.upgrade(config, "head")
    seed()


def make_history_docx(
    *,
    company_name: str = "测试建设有限公司",
    social_credit_code: str = "91310000TEST20260607",
) -> bytes:
    doc = DocxDocument()
    doc.add_heading("历史投标企业资料", level=1)
    doc.add_paragraph("营业执照")
    doc.add_paragraph(f"企业名称：{company_name}")
    doc.add_paragraph(f"统一社会信用代码：{social_credit_code}")
    doc.add_paragraph("资质名称：建筑工程施工总承包")
    doc.add_paragraph("资质等级：一级")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def make_scanned_pdf(page_count: int = 1) -> bytes:
    doc = fitz.open()
    for _ in range(page_count):
        doc.new_page(width=300, height=200)
    return doc.tobytes()


def make_extracted_license(
    *,
    source_file_name: str,
    source_sha256: str,
    certificate_no: str,
    valid_until: date | None = None,
    evidence_text: str | None = None,
) -> HistoryMaterialExtraction:
    evidence = evidence_text or (
        f"营业执照\n企业名称：测试建设有限公司\n统一社会信用代码：{certificate_no}\n"
        f"有效期至：{valid_until.isoformat() if valid_until else '长期'}"
    )
    block = SourceTextBlock(block_index=1, parser="word", heading_path="营业执照", content_text=evidence)
    draft = ExtractedMaterialDraft(
        material_type="license",
        name="测试建设有限公司营业执照",
        certificate_no=certificate_no,
        valid_until=valid_until,
        evidence_text=evidence,
        confidence=0.86,
        source_block_indexes=[1],
        structured_fields={"fallback_rule": "license.business_license"},
    )
    return HistoryMaterialExtraction(
        source_file_name=source_file_name,
        source_content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        source_file_size=len(evidence.encode("utf-8")),
        source_sha256=source_sha256,
        parser_summary={"parser": "word", "mode": "test", "block_count": 1},
        text_blocks=[block],
        drafts=[draft],
        extraction_method="local_rules",
        warnings=[],
    )


def demo_tenant_user(db):
    tenant = db.scalar(select(Tenant).where(Tenant.code == "demo"))
    user = db.scalar(select(User).where(User.external_id == "demo-admin"))
    assert tenant is not None
    assert user is not None
    return tenant, user


def test_parse_docx_history_file_to_source_blocks() -> None:
    blocks, summary, warnings = parse_history_file_to_blocks(
        data=make_history_docx(),
        filename="历史资料.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert summary["parser"] == "word"
    assert not warnings
    assert any("统一社会信用代码" in block.content_text for block in blocks)


def test_parse_pdf_fallbacks_to_ocr_when_text_empty(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakeOcrClient:
        provider = "fake"

        def recognize_image(self, image_bytes: bytes) -> OcrResult:
            assert image_bytes
            return OcrResult(
                text="营业执照\n统一社会信用代码：91310000OCRTEST001",
                confidence=0.91,
                provider="fake",
            )

    monkeypatch.setattr("app.services.history_material_extract.get_ocr_client", lambda: FakeOcrClient())

    blocks, summary, warnings = parse_history_file_to_blocks(
        data=make_scanned_pdf(),
        filename="扫描件.pdf",
        content_type="application/pdf",
    )

    assert summary["parser"] == "pdf_ocr"
    assert warnings
    assert blocks[0].page_no == 1
    assert blocks[0].ocr_confidence == 0.91
    assert "OCRTEST001" in blocks[0].content_text


def test_pdf_ocr_fork_join_returns_blocks_in_page_order(monkeypatch: pytest.MonkeyPatch) -> None:
    def fake_ocr_image_block(
        *,
        image_bytes: bytes,
        block_index: int,
        page_no: int | None,
        parser: str,
    ) -> SourceTextBlock:
        assert image_bytes
        return SourceTextBlock(
            block_index=block_index,
            page_no=page_no,
            parser=parser,
            heading_path=f"PDF 第 {page_no} 页",
            content_text=f"第 {page_no} 页 OCR 结果",
            ocr_confidence=0.88,
        )

    progress_updates: list[tuple[int, int]] = []
    monkeypatch.setattr("app.services.history_material_extract._ocr_image_block", fake_ocr_image_block)

    blocks, summary, warnings = parse_history_file_to_blocks(
        data=make_scanned_pdf(page_count=3),
        filename="扫描件.pdf",
        content_type="application/pdf",
        pdf_mode="ocr",
        ocr_progress_callback=lambda done, total: progress_updates.append((done, total)),
    )

    assert not warnings
    assert summary["parser"] == "pdf_ocr"
    assert [block.page_no for block in blocks] == [1, 2, 3]
    assert [block.content_text for block in blocks] == ["第 1 页 OCR 结果", "第 2 页 OCR 结果", "第 3 页 OCR 结果"]
    assert progress_updates[-1] == (3, 3)


def test_render_compressed_pdf_page_image_uses_readable_jpeg_preview(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(settings, "source_page_image_format", "jpeg")
    monkeypatch.setattr(settings, "source_page_image_max_width", 320)
    monkeypatch.setattr(settings, "source_page_image_jpeg_quality", 60)
    monkeypatch.setattr(settings, "source_page_image_render_scale", 2.0)

    rendered = render_compressed_pdf_page_image(make_scanned_pdf(), page_no=1)

    assert rendered.data.startswith(b"\xff\xd8")
    assert rendered.content_type == "image/jpeg"
    assert rendered.image_format == "jpeg"
    assert rendered.image_quality == 60
    assert rendered.width <= 600
    assert rendered.max_width == 600


def test_history_extract_api_creates_pending_materials_without_chunks(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    stored_objects: list[dict[str, object]] = []

    def fake_put_object_bytes(**kwargs: object) -> None:
        stored_objects.append(kwargs)

    def fake_chat_completion(*_: object, **__: object) -> None:
        raise LLMGatewayError("mocked not configured", code="LLM_NOT_CONFIGURED")

    monkeypatch.setattr("app.services.history_material_extract.chat_completion", fake_chat_completion)
    monkeypatch.setattr("app.api.v1.routes.enterprise.put_object_bytes", fake_put_object_bytes)
    client = TestClient(app)
    response = client.post(
        "/api/v1/enterprise/materials/history-extract",
        files={
            "file": (
                "历史投标资料.docx",
                make_history_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"data_level": "internal"},
    )

    assert response.status_code == 201
    payload = response.json()
    assert payload["draft_count"] >= 2
    assert payload["parser_summary"]["parser"] == "word"
    assert payload["extraction_method"] == "local_rules"
    assert stored_objects
    material_ids = [UUID(item["id"]) for item in payload["materials"]]
    assert {item["verification_status"] for item in payload["materials"]} == {"pending_confirm"}
    assert all(item["structured_fields"]["needs_human_confirm"] for item in payload["materials"])

    with SessionLocal() as db:
        chunk_count = db.scalar(
            select(func.count(EnterpriseMaterialChunk.id)).where(
                EnterpriseMaterialChunk.enterprise_material_id.in_(material_ids)
            )
        )
        assert chunk_count == 0


def test_history_extract_merges_duplicate_pending_material_sources_and_valid_until() -> None:
    certificate_no = f"91310000{uuid4().hex[:12].upper()}"

    with SessionLocal() as db:
        tenant, user = demo_tenant_user(db)
        first = create_pending_materials_from_extraction(
            db,
            tenant_id=tenant.id,
            actor_user_id=user.id,
            extraction=make_extracted_license(
                source_file_name="旧版营业执照.docx",
                source_sha256=f"old-{certificate_no}",
                certificate_no=certificate_no,
                valid_until=date(2027, 12, 31),
            ),
            bucket="test-bucket",
            object_key="history/old-license.docx",
            data_level="internal",
        )
        db.flush()
        material_id = first[0].id

        second = create_pending_materials_from_extraction(
            db,
            tenant_id=tenant.id,
            actor_user_id=user.id,
            extraction=make_extracted_license(
                source_file_name="新版营业执照.docx",
                source_sha256=f"new-{certificate_no}",
                certificate_no=certificate_no,
                valid_until=date(2029, 12, 31),
                evidence_text=f"营业执照\n企业名称：测试建设有限公司\n统一社会信用代码：{certificate_no}\n有效期至：2029-12-31\n复核备注：新版证照",
            ),
            bucket="test-bucket",
            object_key="history/new-license.docx",
            data_level="restricted",
        )
        db.flush()

        assert second[0].id == material_id
        saved = db.get(EnterpriseMaterial, material_id)
        assert saved is not None
        assert saved.valid_until == date(2029, 12, 31)
        assert saved.data_level == "restricted"
        assert saved.verification_status == "pending_confirm"
        assert "新版证照" in (saved.evidence_text or "")
        assert saved.structured_fields is not None
        assert saved.structured_fields["duplicate_merge"]["merge_count"] == 2
        assert [item["source_file_name"] for item in saved.structured_fields["source_files"]] == [
            "旧版营业执照.docx",
            "新版营业执照.docx",
        ]
        duplicate_count = db.scalar(
            select(func.count(EnterpriseMaterial.id)).where(
                EnterpriseMaterial.tenant_id == tenant.id,
                EnterpriseMaterial.certificate_no == certificate_no,
            )
        )
        assert duplicate_count == 1
        chunk_count = db.scalar(
            select(func.count(EnterpriseMaterialChunk.id)).where(
                EnterpriseMaterialChunk.enterprise_material_id == saved.id
            )
        )
        assert chunk_count == 0


def test_history_extract_flags_duplicate_of_confirmed_material_without_overwriting() -> None:
    certificate_no = f"91310000{uuid4().hex[:12].upper()}"

    with SessionLocal() as db:
        tenant, user = demo_tenant_user(db)
        confirmed = EnterpriseMaterial(
            tenant_id=tenant.id,
            material_type="license",
            name="已确认营业执照",
            certificate_no=certificate_no,
            valid_until=date(2027, 12, 31),
            data_level="internal",
            verification_status="confirmed",
            structured_fields={"manual_entry": True},
            evidence_text="已人工确认的营业执照",
            created_by=user.id,
            updated_by=user.id,
        )
        db.add(confirmed)
        db.flush()

        materials = create_pending_materials_from_extraction(
            db,
            tenant_id=tenant.id,
            actor_user_id=user.id,
            extraction=make_extracted_license(
                source_file_name="疑似重复营业执照.docx",
                source_sha256=f"dup-{certificate_no}",
                certificate_no=certificate_no,
                valid_until=date(2029, 12, 31),
            ),
            bucket="test-bucket",
            object_key="history/duplicate-license.docx",
            data_level="internal",
        )
        db.flush()

        duplicate = materials[0]
        assert duplicate.id != confirmed.id
        assert duplicate.verification_status == "pending_confirm"
        assert duplicate.valid_until == date(2029, 12, 31)
        assert duplicate.structured_fields is not None
        assert duplicate.structured_fields["duplicate_of_material_id"] == str(confirmed.id)
        assert duplicate.structured_fields["duplicate_review"]["status"] == "needs_human_review"
        db.refresh(confirmed)
        assert confirmed.valid_until == date(2027, 12, 31)
        assert confirmed.evidence_text == "已人工确认的营业执照"


def test_history_extract_drops_llm_blob_misclassified_as_performance(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    blob_name = (
        "的磋商邀请(采购代理编号：HNZJ(FW)-202603)，签字代表 胡扩军，项目负责人(姓名、职务)经正式授权并代表"
        "供应商 湖南远发建筑工程有限公司 提交响应文件正本一份，副本三份；响应文件电子文档：一份"
    )

    def fake_chat_completion(*_: object, **__: object) -> LLMResult:
        return LLMResult(
            content=json.dumps(
                {
                    "materials": [
                        {
                            "material_type": "performance",
                            "name": blob_name,
                            "evidence_text": blob_name,
                            "confidence": 0.61,
                            "source_block_indexes": [1],
                        }
                    ]
                },
                ensure_ascii=False,
            ),
            provider="fake",
            model_name="fake",
            complexity="complex",
            prompt_version="history-material-extract-v1",
            log_id=uuid4(),
            usage={},
        )

    monkeypatch.setattr("app.services.history_material_extract.chat_completion", fake_chat_completion)
    monkeypatch.setattr("app.api.v1.routes.enterprise.put_object_bytes", lambda **_: None)

    client = TestClient(app)
    response = client.post(
        "/api/v1/enterprise/materials/history-extract",
        files={
            "file": (
                "误抽历史资料.docx",
                make_history_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    payload = response.json()
    assert payload["draft_count"] >= 2
    assert all(blob_name not in item["name"] for item in payload["materials"])
    assert all(item["material_type"] != "performance" for item in payload["materials"])


def test_history_extract_route_is_not_shadowed_by_material_id_routes() -> None:
    client = TestClient(app)
    response = client.post("/api/v1/enterprise/materials/history-extract")

    assert response.status_code == 422
    assert response.status_code != 405


def test_history_extract_uses_dedicated_file_size_limit(monkeypatch: pytest.MonkeyPatch) -> None:
    data = b"x" * 8

    def fake_extract_history_material_drafts(*_: object, **__: object) -> HistoryMaterialExtraction:
        return HistoryMaterialExtraction(
            source_file_name="large.pdf",
            source_content_type="application/pdf",
            source_file_size=len(data),
            source_sha256="sha",
            parser_summary={"parser": "pdf_ocr", "mode": "test", "block_count": 0},
            text_blocks=[],
            drafts=[],
            extraction_method="local_rules",
            warnings=[],
        )

    monkeypatch.setattr("app.api.v1.routes.enterprise.MAX_FILE_BYTES", 4)
    monkeypatch.setattr("app.api.v1.routes.enterprise.HISTORY_MATERIAL_FILE_MAX_BYTES", 16)
    monkeypatch.setattr(
        "app.api.v1.routes.enterprise.extract_history_material_drafts",
        fake_extract_history_material_drafts,
    )
    monkeypatch.setattr("app.api.v1.routes.enterprise.put_object_bytes", lambda **_: None)

    client = TestClient(app)
    response = client.post(
        "/api/v1/enterprise/materials/history-extract",
        files={"file": ("large.pdf", data, "application/pdf")},
    )

    assert response.status_code == 201
    assert response.json()["source_file_size"] == len(data)


def test_history_extract_async_route_creates_task(monkeypatch: pytest.MonkeyPatch) -> None:
    stored_objects: list[dict[str, object]] = []
    dispatched_task_ids: list[object] = []

    monkeypatch.setattr(settings, "run_tasks_inline", True)
    monkeypatch.setattr("app.api.v1.routes.enterprise.put_object_bytes", lambda **kwargs: stored_objects.append(kwargs))
    monkeypatch.setattr(
        "app.api.v1.routes.enterprise._execute_history_extract_background",
        lambda task_id: dispatched_task_ids.append(task_id),
    )

    client = TestClient(app)
    response = client.post(
        "/api/v1/enterprise/materials/history-extract-tasks",
        files={
            "file": (
                "异步历史投标资料.docx",
                make_history_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"data_level": "internal"},
    )

    assert response.status_code == 202
    payload = response.json()
    assert payload["task_type"] == "history_material_extract"
    assert payload["status"] == "pending"
    assert payload["input_json"]["source_file_name"] == "异步历史投标资料.docx"
    assert stored_objects
    assert dispatched_task_ids


def test_execute_history_material_extract_task_creates_pending_materials(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    data = make_history_docx()

    def fake_chat_completion(*_: object, **__: object) -> None:
        raise LLMGatewayError("mocked not configured", code="LLM_NOT_CONFIGURED")

    monkeypatch.setattr("app.services.history_material_extract.chat_completion", fake_chat_completion)
    monkeypatch.setattr("app.services.history_material_extract.get_object_bytes", lambda **_: data)

    with SessionLocal() as db:
        tenant = db.scalar(select(Tenant).where(Tenant.code == "demo"))
        user = db.scalar(select(User).where(User.external_id == "demo-admin"))
        assert tenant is not None
        assert user is not None
        task = AsyncTask(
            tenant_id=tenant.id,
            project_id=None,
            section_id=None,
            task_type="history_material_extract",
            status="pending",
            idempotency_key=f"history-material-extract-test:{uuid4().hex}",
            progress=0,
            input_json={
                "source_file_name": "worker历史投标资料.docx",
                "source_file_size": len(data),
                "source_sha256": "sha",
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "bucket": "test-bucket",
                "object_key": "test-object",
                "pdf_mode": "auto",
                "data_level": "internal",
            },
            retry_count=0,
            max_retries=1,
            created_by=user.id,
        )
        db.add(task)
        db.commit()

        result = execute_history_material_extract_task(db, task.id)
        db.refresh(task)

        assert result["status"] == "succeeded"
        assert task.status == "succeeded"
        assert task.progress == 100
        assert task.output_json is not None
        assert task.output_json["draft_count"] >= 2
        material_ids = [UUID(item["id"]) for item in task.output_json["materials"]]
        materials = db.scalars(
            select(EnterpriseMaterial).where(EnterpriseMaterial.id.in_(material_ids))
        ).all()
        assert materials
        assert {material.verification_status for material in materials} == {"pending_confirm"}


def test_confirming_extracted_material_builds_chunks(monkeypatch: pytest.MonkeyPatch) -> None:
    def fake_chat_completion(*_: object, **__: object) -> None:
        raise LLMGatewayError("mocked not configured", code="LLM_NOT_CONFIGURED")

    monkeypatch.setattr("app.services.history_material_extract.chat_completion", fake_chat_completion)
    monkeypatch.setattr("app.api.v1.routes.enterprise.put_object_bytes", lambda **_: None)
    social_credit_code = f"91310000{uuid4().hex[:12].upper()}"
    client = TestClient(app)
    response = client.post(
        "/api/v1/enterprise/materials/history-extract",
        files={
            "file": (
                "待确认历史资料.docx",
                make_history_docx(social_credit_code=social_credit_code),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201
    material = next(item for item in response.json()["materials"] if item["certificate_no"] == social_credit_code)

    search_before_confirm = client.get(
        "/api/v1/enterprise/materials/search",
        params={"query": material["certificate_no"] or material["name"], "limit": 10},
    )
    assert search_before_confirm.status_code == 200
    assert all(item["id"] != material["id"] for item in search_before_confirm.json())

    confirm_response = client.patch(
        f"/api/v1/enterprise/materials/{material['id']}",
        json={"verification_status": "confirmed", "reason": "已核对历史文件来源和字段"},
    )

    assert confirm_response.status_code == 200
    assert confirm_response.json()["verification_status"] == "confirmed"
    with SessionLocal() as db:
        saved = db.get(EnterpriseMaterial, UUID(material["id"]))
        assert saved is not None
        chunks = db.scalars(
            select(EnterpriseMaterialChunk).where(
                EnterpriseMaterialChunk.enterprise_material_id == saved.id
            )
        ).all()
        assert chunks
        assert chunks[0].metadata_json["verification_status"] == "confirmed"

    search_after_confirm = client.get(
        "/api/v1/enterprise/materials/search",
        params={"query": material["certificate_no"] or material["name"], "limit": 10},
    )
    assert search_after_confirm.status_code == 200
    assert any(item["id"] == material["id"] for item in search_after_confirm.json())


def test_confirming_pdf_certificate_material_keeps_source_page_image_payload(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    stored_objects: list[dict[str, object]] = []
    certificate_no = f"91310000PDFIMAGE{uuid4().hex[:8].upper()}"

    class FakeOcrClient:
        provider = "fake"

        def recognize_image(self, image_bytes: bytes) -> OcrResult:
            assert image_bytes
            return OcrResult(
                text=f"营业执照\n企业名称：测试建设有限公司\n统一社会信用代码：{certificate_no}",
                confidence=0.93,
                provider="fake",
            )

    def fake_chat_completion(*_: object, **__: object) -> None:
        raise LLMGatewayError("mocked not configured", code="LLM_NOT_CONFIGURED")

    monkeypatch.setattr("app.services.history_material_extract.get_ocr_client", lambda: FakeOcrClient())
    monkeypatch.setattr("app.services.history_material_extract.chat_completion", fake_chat_completion)
    monkeypatch.setattr(settings, "source_page_image_format", "jpeg")
    monkeypatch.setattr(settings, "source_page_image_max_width", 640)
    monkeypatch.setattr(settings, "source_page_image_jpeg_quality", 60)
    monkeypatch.setattr(settings, "source_page_image_render_scale", 2.0)
    monkeypatch.setattr("app.api.v1.routes.enterprise.put_object_bytes", lambda **kwargs: stored_objects.append(kwargs))
    monkeypatch.setattr("app.services.source_page_images.put_object_bytes", lambda **kwargs: stored_objects.append(kwargs))

    client = TestClient(app)
    response = client.post(
        "/api/v1/enterprise/materials/history-extract",
        files={"file": ("营业执照扫描.pdf", make_scanned_pdf(), "application/pdf")},
        data={"pdf_mode": "ocr"},
    )

    assert response.status_code == 201
    material = response.json()["materials"][0]
    source_images = material["structured_fields"]["source_images"]
    assert source_images
    assert source_images[0]["page_no"] == 1
    assert source_images[0]["content_type"] == "image/jpeg"
    assert source_images[0]["image_quality"] == 60
    assert source_images[0]["width"] <= 640
    assert source_images[0]["source"] == "source_page_image_asset"
    assert source_images[0]["page_image_id"] == source_images[0]["id"]
    assert any(str(item["object_key"]).endswith(".jpg") for item in stored_objects)

    with SessionLocal() as db:
        asset = db.get(SourcePageImage, UUID(source_images[0]["page_image_id"]))
        assert asset is not None
        assert asset.image_format == "jpeg"
        assert asset.image_content_type == "image/jpeg"
        assert asset.image_quality == 60
        assert asset.max_width == 640
        assert asset.image_file_size > 0
        assert asset.width <= 640
        assert asset.source_kind == "history_upload"

    confirm_response = client.patch(
        f"/api/v1/enterprise/materials/{material['id']}",
        json={"verification_status": "confirmed", "reason": "已核对证照原页图片和 OCR 文本"},
    )
    assert confirm_response.status_code == 200

    with SessionLocal() as db:
        chunks = db.scalars(
            select(EnterpriseMaterialChunk).where(
                EnterpriseMaterialChunk.enterprise_material_id == UUID(material["id"])
            )
        ).all()
        assert chunks
        assert chunks[0].metadata_json["source_images"][0]["page_no"] == 1
        assert chunks[0].metadata_json["source_images"][0]["page_image_id"] == source_images[0]["page_image_id"]
        assert chunks[0].metadata_json["primary_source_image"]["content_type"] == "image/jpeg"
